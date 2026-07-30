#!/usr/bin/env python3
"""
Generate Research Protocol v1.1 — RecoveryLab
===============================================
Updated version incorporating external review feedback.

Key changes from v1.0:
  1. Product identity: "Benchmark Suite" is now a hypothesis, not a statement
  2. External validation: expanded beyond PhotoRec to multiple tool families
  3. Success criterion: 5% threshold replaced with empirical calibration framework
  4. RVS calibration: added plan for user surveys to validate value weights
  5. Added Section 11: Immediate Operational Objectives
"""

import sys
import os

# Add docx scripts to path
DOCX_SCRIPTS = os.path.join("/home/z/my-project/skills/docx", "scripts")
if DOCX_SCRIPTS not in sys.path:
    sys.path.insert(0, DOCX_SCRIPTS)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Color Palette — Deep Sea Academic ────────────────────────────────────────
P = {
    "primary": "#162032",
    "body": "#1C2A3D",
    "secondary": "#5B6B7D",
    "accent": "#8B7E5A",
    "surface": "#F5F7FA",
    "white": "#FFFFFF",
    "star_gold": "#C9A84C",
    "star_dim": "#B0B0B0",
    "red": "#C0392B",
    "green": "#27AE60",
    "orange": "#E67E22",
    "blue": "#2980B9",
    "hypothesis_bg": "#FFF8E7",
}

def c(hex_color):
    """Convert hex color to RGBColor."""
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.replace("#", "")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, highlight_col=None):
    """Add a professionally styled table with optional column highlighting."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = c(P["white"])
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, P["primary"])

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = c(P["body"])
            if row_idx % 2 == 1:
                set_cell_shading(cell, P["surface"])

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    return table


def add_callout_box(doc, title, text, color=P["accent"]):
    """Add a highlighted callout box with a title and body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = c(color)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = c(P["body"])
    run.italic = True


def build_research_protocol():
    """Build the complete Research Protocol v1.1 document."""
    doc = Document()

    # ─── Page Setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # ─── Default Font ─────────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = c(P["body"])

    # ─── COVER PAGE ────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    # Lab name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RECOVERYLAB")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"
    run.bold = True

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Research Protocol v1.1")
    run.font.size = Pt(28)
    run.font.color.rgb = c(P["primary"])
    run.font.name = "Calibri"
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Congelar la arquitectura, auditar la ciencia")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.italic = True

    # Version note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Revision incorporando auditoria externa")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    # Metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento vivo - Version 1.1")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clasificacion: INTERNO - Solo para uso del laboratorio")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["red"])
    run.font.name = "Calibri"

    # Change log
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Cambios respecto a v1.0")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"
    run.bold = True

    changes = [
        "Producto: 'ES Benchmark Suite' reformulado como hipotesis de negocio",
        "Validacion externa: expandida de PhotoRec a multiples familias de herramientas",
        "Criterio de exito: umbral 5% reemplazado por calibracion empirica",
        "RVS: agregado plan de calibracion con usuarios reales",
        "Nueva seccion: Objetivos Operativos Inmediatos",
    ]
    for change in changes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"  {change}")
        run.font.size = Pt(9)
        run.font.color.rgb = c(P["secondary"])
        run.font.name = "Calibri"

    # Page break
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1: PREGUNTA CENTRAL
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("1. Pregunta Central del Proyecto", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto ya no gira alrededor del Motor C. Gira alrededor de una pregunta:"
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # The central question — highlighted
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "\u00bfComo medir objetivamente la utilidad de una estrategia de recuperacion?"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta pregunta es mucho mas amplia que cualquier motor individual. Es mucho mas dificil de copiar. "
        "Y es la base sobre la cual se construye todo el proyecto. Si el laboratorio demuestra que sus resultados "
        "son reproducibles, comparables y resisten auditorias externas, habremos construido algo que vale "
        "mucho mas que un motor de recuperacion: una infraestructura de evaluacion objetiva que es "
        "extraordinariamente dificil de desarrollar y de replicar."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2: VARIABLES EXPERIMENTALES
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("2. Variables Experimentales", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    # 2.1 Independent Variable
    h2 = doc.add_heading("2.1 Variable Independiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable independiente es aquella que el experimentador controla. Solo una debe cambiar "
        "por experimento. Si cambian dos a la vez, las conclusiones se vuelven ambiguas. Este es "
        "el principio fundamental del control experimental: cada experimento debe aislar exactamente "
        "un factor para que la relacion causal sea interpretable. Cuando dos variables cambian "
        "simultaneamente, es imposible determinar cual de ellas produjo el efecto observado, y "
        "cualquier conclusion queda invalidada por esta confoundacion."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Independent variables table
    add_styled_table(
        doc,
        ["Variable Independiente", "Valores", "Experimento Asociado"],
        [
            ["Tipo de dano", "MFT parcial, head crash, sectores intermitentes, ruido aleatorio, etc.", "H4 Matrix"],
            ["Nivel de dano", "0%, 10%, 20%, ..., 100%", "Crossover Curve"],
            ["Formato de archivo", "JPEG, PNG, PDF (congelados)", "H5 Per-Format"],
            ["Estrategia", "Carving, MFT-First, Hybrid, Motor C", "H1.1 / H2"],
            ["Presupuesto de lectura", "0, 100, 500, 1000, 5000 sectores", "H1.7 Budget"],
        ],
        col_widths=[4.0, 6.0, 5.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Regla critica: ")
    run.bold = True
    run.font.color.rgb = c(P["red"])
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Solo UNA variable independiente por experimento. Si se cambia el formato Y el nivel de dano "
        "simultaneamente, no se puede atribuir la diferencia a ninguno de los dos. Esta regla no es "
        "negociable: es la base de la inferencia causal."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # 2.2 Dependent Variable
    h2 = doc.add_heading("2.2 Variable Dependiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable dependiente es lo que se mide. Debe elegirse ANTES de ejecutar el experimento, "
        "no despues de ver los resultados. Cambiar la definicion de exito a posteriori invalida "
        "cualquier conclusion. El RecoveryLab tiene multiples metricas disponibles, pero cada "
        "experimento debe declarar cual es la principal. La declaracion previa es la unica defensa "
        "contra el sesgo de confirmacion: si siempre es posible encontrar una metrica que favorezca "
        "nuestro resultado, ninguna conclusion es confiable."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Metrica", "Tipo", "Definicion", "Rango"],
        [
            ["Recovery Rate", "Conteo", "Fraccion de archivos recuperados (SHA-256 match)", "0.0-1.0"],
            ["RVS", "Valor", "Valor recuperado = suma(archivo_i.valor x peso_i) / total_valor", "0.0-1.0"],
            ["FQS", "Calidad", "Calidad funcional = suma(archivo_i.score x tamano_i) / total_tamano", "0.0-1.0"],
            ["Overall Utility", "Compuesto", "RVS x FQS - utilidad total de la recuperacion", "0.0-1.0"],
            ["Read Efficiency", "Eficiencia", "Lecturas utiles / Lecturas totales", "0.0-1.0"],
            ["Time to First File", "Tiempo", "Lecturas antes del primer archivo recuperado", "0-N"],
        ],
        col_widths=[3.0, 2.0, 6.5, 1.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Metrica principal declarada: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Overall Utility (RVS x FQS) es la metrica principal del laboratorio. "
        "Recovery Rate se mantiene como metrica secundaria para comparabilidad con herramientas externas. "
        "Cada experimento debe declarar explicitamente cual metrica es la principal antes de ejecutar."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3: CRITERIO DE EXITO (CORREGIDO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("3. Criterio de Exito", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El criterio de exito define cuando una estrategia se considera superior a otra. "
        "Este criterio debe declararse ANTES de ejecutar el experimento. Si se define a posteriori, "
        "siempre es posible encontrar un criterio que favorezca cualquier resultado. La definicion "
        "previa es la unica defensa contra el sesgo de confirmacion."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ── V1.1 CHANGE: Empirical threshold instead of fixed 5% ──
    add_callout_box(
        doc,
        "Cambio en v1.1: Umbral empirico en lugar de 5% fijo",
        "En la version 1.0 se establecia un umbral fijo del 5% de mejora en Overall Utility. "
        "Este numero era arbitrario: no surgia de datos experimentales, sino de una decision "
        "subjetiva. Un umbral de significancia debe derivarse de la variabilidad observada, "
        "el error de medicion y la repetibilidad de los experimentos. No de una intuicion.",
        P["orange"]
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Criterio de exito v1.1:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(
        "Una estrategia A se considera superior a una estrategia B unicamente si:"
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    criteria = [
        "Overall Utility (RVS x FQS) mejora por encima del umbral empirico (ver 3.1) sobre la estrategia B",
        "La diferencia es consistente en al menos 10 repeticiones con el mismo dataset",
        "La diferencia se mantiene en al menos 3 datasets distintos",
        "La mejora no se debe exclusivamente a una dimension (RVS o FQS) sino a ambas",
    ]
    for i, criterion in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f"{i}. {criterion}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Si solo se cumple el criterio 1 pero no los criterios 2-4, el resultado se clasifica como "
        "observacion aislada (1 estrella) y no como evidencia solida. La consistencia es tan importante "
        "como la magnitud de la mejora. Un resultado que no se replica es un resultado que no existe "
        "desde el punto de vista cientifico."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # 3.1 Umbral Empirico
    h2 = doc.add_heading("3.1 Calibracion del Umbral Empirico", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El umbral de significancia no debe ser un numero elegido arbitrariamente. Debe surgir de los "
        "datos del laboratorio. El procedimiento para calibrarlo es el siguiente: se ejecuta el mismo "
        "experimento (misma estrategia, mismo dataset, mismas condiciones) multiples veces, y se mide "
        "la variabilidad natural de Overall Utility. Si las fluctuaciones entre ejecuciones son del 3%, "
        "entonces una mejora del 3% no es significativa: esta dentro del ruido. El umbral debe ser "
        "al menos 2 veces la desviacion estandar de la variabilidad observada, para que una mejora "
        "declarada tenga una probabilidad baja de ser producto del azar."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Procedimiento de calibracion:")
    run.bold = True
    run.font.size = Pt(11)

    cal_steps = [
        "Ejecutar el experimento baseline (Carving sobre JPEG, dataset estandar) 30 veces seguidas.",
        "Registrar Overall Utility en cada ejecucion. Calcular media y desviacion estandar.",
        "Definir umbral = max(2 x desviacion_estandar, 0.01). Es decir: al menos 2 sigma, con un piso absoluto del 1%.",
        "Repetir para cada formato (PNG, PDF) y cada estrategia (Carving, MFT-First, Hybrid).",
        "El umbral final es el maximo entre todos los umbrales calculados. Esto garantiza que el umbral es conservador.",
        "Si la variabilidad es muy baja (determinista), el umbral se mantiene en el piso del 1%.",
    ]
    for i, step in enumerate(cal_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {step}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Estado actual: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Este procedimiento aun no se ha ejecutado. El umbral definitivo se determinara "
        "despues de las 30 ejecuciones baseline. Hasta entonces, se usa un umbral provisional "
        "del 3% (conservador, basado en la variabilidad observada en ejecuciones previas). "
        "Este numero provisional se reemplazara con el umbral calibrado tan pronto como "
        "los datos de variabilidad esten disponibles."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4: REGISTRO DE CONFIANZA
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("4. Registro de Confianza", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El registro de confianza es la forma mas transparente de comunicar la solidez cientifica "
        "de cada resultado. En lugar de porcentajes arbitrarios de confianza, utiliza un sistema "
        "de estrellas que mapea directamente al tipo de evidencia acumulada. Este sistema es "
        "mas honesto que cualquier numero: un lector puede ver inmediatamente que evidencia "
        "respalda un resultado y que tan lejos esta de ser validado externamente."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Star scale table
    add_styled_table(
        doc,
        ["Estrellas", "Evidencia Requerida", "Significado"],
        [
            ["*", "Observacion aislada (1 run, 1 dataset)", "Preliminar - no citar como conclusion"],
            ["**", "Repetido 10+ veces (determinista, mismo dataset)", "Estable - pero no generalizado"],
            ["***", "Repetido con datasets distintos", "Generalizable - pero no validado externamente"],
            ["****", "Validado con herramientas externas (multiples familias, ver Seccion 7.2)", "Robusto - comparable con el estado del arte"],
            ["*****", "Validado con hardware real", "Definitivo - predictivo del mundo real"],
        ],
        col_widths=[2.0, 6.0, 5.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Reglas del registro: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Las estrellas solo pueden subir (acumulacion de evidencia), nunca bajar. "
        "Si aparece evidencia contradictoria, el resultado se marca como CONTESTADO pero "
        "no pierde estrellas. El lector debe decidir como interpretar la contradiccion. "
        "Ademas, es fundamental que cada estrella este vinculada a un registro de los "
        "experimentos que la soportan: no basta con decir que algo se repitio 10 veces; "
        "debe haber un registro de las 10 ejecuciones con sus resultados individuales."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # Current state table
    h2 = doc.add_heading("4.1 Estado Actual del Registro", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Resultado", "Estrellas", "Resumen"],
        [
            ["H1.1", "*** (3/5)", "MFT-First supera a Carving consistentemente. 100/100 escenarios."],
            ["H1.2", "* (1/5)", "Umbral de cambio de estrategia no determinado."],
            ["H2", "** (2/5)", "Crossover existe pero es artefacto del carving limitado."],
            ["H4", "* (1/5)", "Matriz preliminar con 3 celdas de muchas."],
            ["H5", "* (1/5)", "Sin experimento sistematico por formato."],
            ["H6", "** (2/5)", "FunctionalValidator implementado, 19/19 tests."],
            ["H7", "* (1/5)", "RVS implementado, sin validacion de usuarios (ver Seccion 5.1)."],
            ["H8", "* (1/5)", "Crossover al 95% es artefacto, no descubrimiento."],
            ["RVS", "** (2/5)", "4 dimensiones implementadas, sin calibracion externa (ver Seccion 5.1)."],
            ["FQS", "** (2/5)", "5 niveles funcionales, 19/19 tests."],
            ["WFS", "* (1/5)", "Concepto definido (RVS x FQS), aun no integrado en Judge."],
            ["BLOCKER-001", "*** (3/5)", "Resuelto: comparaciones previas invalidas."],
            ["H1.5", "* (1/5)", "Gaps: sin fragmentacion, sin INDX, sin jerarquia."],
            ["H1.6", "** (2/5)", "100 ejecuciones, resultados deterministas."],
            ["H1.7", "* (1/5)", "Motor C apenas supera MFT-First (5% support)."],
        ],
        col_widths=[2.5, 2.0, 10.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Resumen: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "8 resultados a 1 estrella, 5 a 2 estrellas, 2 a 3 estrellas, 0 a 4 estrellas, "
        "0 a 5 estrellas. Esto es honesto: la mayoria de los resultados son preliminares. "
        "No hay resultados validados externamente ni con hardware real. El laboratorio "
        "esta en una fase temprana de acumulacion de evidencia. El objetivo de la Fase A "
        "es elevar las hipotesis principales (H1.1, H6, H7) a al menos 3 estrellas."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 5: DECOMPOSICION DE METRICAS
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("5. Decomposicion de Metricas", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La metrica compuesta anterior (WFS) mezclaba dos dimensiones distintas: "
        "que archivos se recuperaron y que tan bien se recuperaron. La separacion "
        "es esencial porque un motor puede ganar por dos razones muy diferentes, "
        "y la interpretacion cambia completamente segun cual sea. Un motor que recupera "
        "la tesis pero con pixeles corruptos tiene alto RVS y bajo FQS; un motor que "
        "recupera perfectamente 200 thumbnails tiene bajo RVS y alto FQS. Sin esta "
        "separacion, ambas victorias se ven identicas, y eso es cientificamente inaceptable."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Decomposition formula
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Overall Utility = RVS (Valor) x FQS (Calidad)")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.bold = True

    # RVS definition
    h2 = doc.add_heading("5.1 RVS - Recovery Value Score (Que se recupero)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El RVS mide el valor de lo recuperado, ponderado por la importancia de cada archivo. "
        "No es lo mismo recuperar la tesis que 200 thumbnails. El RVS incorpora cuatro dimensiones: "
        "valor intrinseco del archivo, probabilidad de reemplazo, tiempo de recreacion, e impacto "
        "emocional de la perdida. Un motor que recupera la tesis pero pierde 200 thumbnails tiene "
        "mayor RVS que uno que hace lo contrario, porque la tesis tiene mas valor, es mas dificil "
        "de reemplazar, requiere mas tiempo para recrear, y su perdida tiene mayor impacto emocional."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Tipo de Archivo", "Valor", "Reemplazo", "Recreacion", "Emocional", "RVS Peso"],
        [
            ["tesis.docx", "100", "0.05", "500h", "100", "100"],
            ["sqlite.db", "95", "0.02", "200h", "50", "95"],
            ["foto familiar", "90", "0.00", "0h", "100", "90"],
            ["video vacaciones", "70", "0.00", "0h", "80", "70"],
            ["PSD proyecto", "65", "0.10", "40h", "30", "65"],
            ["RAW foto", "60", "0.00", "0h", "60", "60"],
            ["MP4 descargado", "20", "0.90", "1h", "5", "20"],
            ["Linux ISO", "2", "1.00", "0.5h", "0", "2"],
            ["thumbnail", "1", "1.00", "0h", "0", "1"],
        ],
        col_widths=[3.0, 1.5, 1.5, 2.0, 2.0, 1.5],
    )

    # ── V1.1 CHANGE: RVS calibration warning ──
    add_callout_box(
        doc,
        "Limitacion reconocida (v1.1): RVS necesita calibracion con usuarios reales",
        "La tabla de valores anterior es una asignacion razonable hecha por el laboratorio. "
        "Pero sigue siendo una decision interna. No ha sido calibrada con comportamiento real "
        "de usuarios. Ver Seccion 5.3 para el plan de calibracion.",
        P["orange"]
    )

    # FQS definition
    h2 = doc.add_heading("5.2 FQS - Functional Quality Score (Que tan bien se recupero)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El FQS mide la calidad funcional de la recuperacion. No es lo mismo recuperar un "
        "JPEG bit-perfect que uno con 2 pixeles corruptos. El FQS reemplaza el binario "
        "SHA-256 coincide/no coincide con un espectro de calidad funcional que refleja "
        "mejor la experiencia real del usuario. Un archivo que abre y funciona tiene valor "
        "incluso si su checksum no coincide exactamente. La ventaja del FQS es que captura "
        "gradaciones de calidad que el binario pasa por alto: un JPEG con 2 pixeles malos "
        "no es lo mismo que un JPEG completamente corrupto, pero ambos son 'fallo' en el "
        "modelo binario."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Nivel", "Score", "Definicion", "Ejemplo"],
        [
            ["FULL", "1.0", "SHA-256 coincide exactamente", "Copia bit-perfect del original"],
            ["FUNCTIONAL", "0.8", "Archivo funciona con dano menor", "JPEG con 2 pixeles corruptos"],
            ["PARTIAL", "0.5", "Funciona parcialmente, perdida de datos", "MP4 que se reproduce hasta la mitad"],
            ["DEGRADED", "0.2", "Contenido accesible pero daniado", "DOCX que abre pero perdio imagenes"],
            ["FAILED", "0.0", "Completamente inutilizable", "Datos aleatorios sin estructura"],
        ],
        col_widths=[2.5, 1.5, 5.0, 5.0],
    )

    # Diagnostic interpretation
    h2 = doc.add_heading("5.3 Diagnostico: Por que gano un motor?", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La separacion RVS x FQS permite diagnosticar por que un motor gano. "
        "Un motor puede ganar por dos razones fundamentalmente distintas, y la "
        "interpretacion cambia completamente segun cual sea la dominante."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["RVS", "FQS", "Diagnostico", "Interpretacion"],
        [
            ["Alto", "Alto", "STRONG", "Recupero archivos importantes con buena calidad"],
            ["Alto", "Bajo", "VALUE-DRIVEN", "Recupero archivos importantes pero con mala calidad"],
            ["Bajo", "Alto", "QUALITY-DRIVEN", "Recupero archivos bien pero no eran importantes"],
            ["Bajo", "Bajo", "WEAK", "Fallo en ambas dimensiones"],
        ],
        col_widths=[2.0, 2.0, 3.0, 7.0],
    )

    # ── V1.1 NEW: RVS Calibration Plan ──
    h2 = doc.add_heading("5.4 Plan de Calibracion de RVS con Usuarios Reales", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La tabla de valores de RVS (Seccion 5.1) es una asignacion razonable hecha por el "
        "laboratorio, pero sigue siendo una decision interna. No ha sido validada con comportamiento "
        "real de usuarios. Esto es el punto mas debil del protocolo actual: los pesos de RVS "
        "determinan que estrategia se considera superior, y si esos pesos estan mal calibrados, "
        "todas las conclusiones del laboratorio se basan en una premisa incorrecta."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run(
        "Para que el RVS deje de ser un modelo disenado por el laboratorio y pase a estar "
        "calibrado con comportamiento real, se necesita una encuesta de calibracion. La pregunta "
        "central de la encuesta es simple: 'Si solo pudieras recuperar uno de estos archivos, "
        "cual elegirias?' La respuesta a esta pregunta, repetida con suficiente diversidad de "
        "usuarios, produce una ranking de valor que reemplaza la tabla actual."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Poblaciones objetivo para la encuesta:")
    run.bold = True
    run.font.size = Pt(11)

    survey_pops = [
        "Fotografos profesionales: valoran RAW y fotos familiares, desprecian ISO y thumbnails.",
        "Estudios juridicos: valoran documentos y bases de datos, desprecian thumbnails.",
        "Empresas de tecnologia: valoran bases de datos y codigo, desprecian ISO.",
        "Usuarios domesticos: valoran fotos y videos familiares, desprecian archivos de sistema.",
        "Estudiantes: valoran tesis y documentos, desprecian thumbnails.",
    ]
    for pop in survey_pops:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(pop)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Metodologia de la encuesta:")
    run.bold = True
    run.font.size = Pt(11)

    survey_methods = [
        "Presentar pares de archivos y pedir al usuario que elija cual recuperaria primero.",
        "Usar el metodo de comparacion por pares (Bradley-Terry) para obtener un ranking continuo.",
        "Minimo 30 respuestas por poblacion para obtener estimaciones estables.",
        "Comparar los pesos obtenidos con la tabla actual. Si difieren significativamente, actualizar la tabla y recalcular resultados previos.",
        "Publicar los datos de la encuesta como parte del Benchmark Suite para transparencia.",
    ]
    for method in survey_methods:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(method)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Estado: ")
    run.bold = True
    run.font.color.rgb = c(P["orange"])
    run.font.size = Pt(11)
    run2 = p.add_run(
        "La encuesta no se ha realizado. Es una actividad de la Fase B, no de la Fase A. "
        "Hasta que se complete, los pesos de RVS deben tratarse como supuestos, no como "
        "hechos. Cada resultado que dependa de RVS debe incluir un descargo: 'Los pesos de "
        "valor son asignaciones del laboratorio, no estan calibrados con usuarios reales.'"
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 6: AUDITORIA DE HIPOTESIS
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("6. Auditoria de Hipotesis", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cada hipotesis debe declarar explicitamente su variable independiente, variable dependiente, "
        "y criterio de exito antes de ejecutar el experimento. Si estos no estan definidos, "
        "la hipotesis no es testeable y debe ser reformulada. Una hipotesis sin criterio de "
        "refutacion no es una hipotesis cientifica: es una opinion. El protocolo exige que cada "
        "hipotesis defina explicitamente que resultado la refutaria, no solo que resultado la "
        "confirmaria."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Hipotesis", "Variable Independiente", "Variable Dependiente", "Criterio de Exito", "Confiabilidad"],
        [
            ["H1.1", "Estrategia (Carving vs MFT-First)", "Overall Utility", "MFT-First > Carving por umbral empirico+ en 3+ datasets", "***"],
            ["H1.2", "Confianza en metadatos (0-100%)", "Estrategia optima", "Cambio de estrategia en umbral definido", "*"],
            ["H2", "Nivel de dano MFT (0-100%)", "Recovery Rate", "Crossover observable con carving completo", "**"],
            ["H4", "Tipo de dano", "Mejor estrategia", "Matriz completa con >80% celdas llenas", "*"],
            ["H5", "Formato de archivo", "Recovery Rate por formato", "Diferencia >10% entre formatos", "*"],
            ["H6", "Nivel de dano en archivo", "FQS", "Espectro no-binario con >2 niveles", "**"],
            ["H7", "Tipo de archivo recuperado", "RVS", "Tesis > 200 thumbnails (RVS calibrado)", "*"],
            ["H8", "Motor de carving (limitado vs completo)", "Crossover point", "Cambio significativo en crossover", "*"],
        ],
        col_widths=[1.5, 3.5, 2.5, 4.0, 1.5],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 7: HOJA DE RUTA POR FASES (CORREGIDA)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("7. Hoja de Ruta por Fases", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    # Phase A
    h2 = doc.add_heading("7.1 Fase A - Ahora Mismo (Congelar)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    phase_a_items = [
        "Congelar nuevas funcionalidades. No agregar nuevos formatos (MP4, DOCX, SQLite) al motor de carving.",
        "Consolidar el protocolo experimental. Este documento es el paso 1.",
        "Ejecutar experimentos por formato con JPEG, PNG y PDF unicamente.",
        "Repetir los experimentos suficientes veces para verificar estabilidad (minimo 10 repeticiones, 3 datasets).",
        "Completar la integracion de Overall Utility = RVS x FQS en el Judge.",
        "Calibrar el umbral empirico de significancia (30 ejecuciones baseline, ver Seccion 3.1).",
        "Exigir ratio 1:1 de codigo de recuperacion vs codigo de validacion.",
    ]
    for item in phase_a_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Phase B — CORREGIDO
    h2 = doc.add_heading("7.2 Fase B - Validacion Externa (Corregida)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Validar el laboratorio contra herramientas reales. Pero no basta con comparar contra "
        "una sola herramienta. PhotoRec representa carving puro, pero el laboratorio evalua "
        "multiples familias de estrategias: MFT-first, hybridas, orquestadores, presupuestos "
        "de lectura. Para que la validacion externa sea significativa, necesitamos comparar "
        "contra herramientas que representen CADA familia de estrategias, no solo una."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ── V1.1 CHANGE: Expanded validation tool families ──
    add_styled_table(
        doc,
        ["Familia de Estrategia", "Herramienta Representativa", "Que valida"],
        [
            ["Carving puro", "PhotoRec / TestDisk", "Si nuestro carving es comparable al estado del arte en carving"],
            ["MFT-first", "R-Studio / ReclaiMe", "Si nuestras estrategias MFT son comparables a herramientas comerciales"],
            ["Hybrida", "DMDE (Disk Drill)", "Si nuestras estrategias hibridas son comparables a herramientas que combinan MFT + carving"],
            ["Orquestador", "UFS Explorer", "Si nuestro Motor C es comparable a herramientas que adaptan estrategia al caso"],
        ],
        col_widths=[3.5, 4.0, 6.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Cada herramienta externa se ejecuta sobre los mismos datasets de JPEG/PNG/PDF. "
        "Se comparan Recovery Rate, RVS y FQS. Si las conclusiones del laboratorio coinciden "
        "con el comportamiento de las herramientas de la misma familia, los resultados ganan "
        "la cuarta estrella. Si no coinciden, el laboratorio tiene un problema de validez que "
        "debe resolverse antes de continuar. La comparacion con multiples familias es esencial: "
        "si solo validamos contra carving (PhotoRec), estamos validando solo una cara del "
        "problema, y las conclusiones sobre estrategias MFT-first o hibridas quedarian sin "
        "validacion externa."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    phase_b_items = [
        "Ejecutar PhotoRec sobre los mismos datasets de JPEG/PNG/PDF (carving baseline).",
        "Ejecutar R-Studio o ReclaiMe sobre los mismos datasets (MFT-first baseline).",
        "Ejecutar DMDE sobre los mismos datasets (hybrid baseline).",
        "Comparar Recovery Rate, RVS y FQS de cada herramienta vs los motores del laboratorio.",
        "Si los resultados no son consistentes con la herramienta de la misma familia, investigar y documentar las discrepancias.",
        "Ejecutar la encuesta de calibracion de RVS (ver Seccion 5.4).",
        "Publicar los resultados de la comparacion como parte del Benchmark Suite.",
    ]
    for item in phase_b_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Phase C
    h2 = doc.add_heading("7.3 Fase C - Expansion Controlada", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Solo entonces incorporar un formato nuevo (por ejemplo MP4) y exigirle el mismo "
        "estandar de calidad y pruebas que ya se alcanzo con JPEG, PNG y PDF. Cada nuevo "
        "formato debe pasar por el mismo proceso: parser impecable, 19+ tests de validacion, "
        "experimentos por formato, y estabilidad verificada. La regla es simple: un parser "
        "excelente de JPEG tiene muchisimo mas valor cientifico que diez parsers aceptables."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 8: FORMATOS CONGELADOS
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("8. Formatos Congelados - Referencia Dorada", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Los tres parsers actuales (JPEG, PNG, PDF) tienen 19/19 tests pasados. "
        "Esto vale oro. Antes de agregar un nuevo formato, estos tres deben convertirse "
        "en una referencia absoluta. La calidad de los parsers existentes es la base "
        "sobre la cual se construye la credibilidad del laboratorio. Un parser con 19/19 "
        "tests no es solo un parser que funciona: es un parser que se puede usar como "
        "patron de referencia para calibrar otros parsers y para validar que los "
        "experimentos de recuperacion no estan sesgados por defectos del parser."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Formato", "Parser", "Tests", "Estado", "Accion"],
        [
            ["JPEG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PNG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PDF", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["MP4", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["DOCX", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["SQLite", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["CR2/NEF", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["TXT", "Imposible", "0/0", "IMPOSIBLE", "Sin firma ni footer - no es carvable"],
        ],
        col_widths=[2.0, 3.0, 1.5, 2.5, 5.0],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 9: REGLA DE ORO
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("9. Regla de Oro", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    # Highlighted box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Por cada 500 lineas nuevas de codigo de recuperacion, "
        "agregar al menos 500 lineas de validacion, pruebas y metricas."
    )
    run.font.size = Pt(13)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta regla no es una sugerencia. Es la unica defensa contra la tentacion de agregar "
        "funcionalidades sin validarlas. La calidad del laboratorio es su activo mas valioso. "
        "Un motor con 3 parsers impecables y 19/19 tests es mas cientificamente valioso que "
        "un motor con 17 parsers mediocres y tests insuficientes. La presion para agregar "
        "formatos es real, pero ceder a ella destruye la base sobre la cual se construye "
        "la credibilidad del proyecto. Si realmente se mantiene esta disciplina durante uno "
        "o dos anos, RecoveryLab podria terminar siendo mucho mas confiable que muchos "
        "proyectos open source donde el codigo crece mucho mas rapido que las pruebas."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 10: PRODUCTO (CORREGIDO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("10. Producto del Proyecto", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    # ── V1.1 CHANGE: Product as hypothesis, not statement ──
    add_callout_box(
        doc,
        "Cambio en v1.1: Producto como hipotesis de negocio, no como declaracion",
        "En la version 1.0 se afirmaba: 'El producto ya no es un motor sino un Benchmark Suite.' "
        "Esto puede ser cierto, pero todavia es una hipotesis de negocio. No esta demostrado que "
        "exista suficiente demanda por una plataforma de benchmarking para recuperacion de datos. "
        "La redaccion correcta es una pregunta, no una afirmacion.",
        P["orange"]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Estamos investigando si el verdadero activo competitivo del proyecto termina siendo "
        "el RecoveryLab Benchmark Suite: una plataforma objetiva para evaluar estrategias de "
        "recuperacion de datos. Esta hipotesis de negocio se basa en la observacion de que "
        "medir objetivamente la utilidad de una estrategia de recuperacion es un problema "
        "mas amplio y mas dificil de replicar que construir un motor individual. Pero no "
        "esta demostrado que exista suficiente demanda por esta plataforma. La validacion "
        "de esta hipotesis de negocio es parte del trabajo futuro, no una conclusion actual."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run(
        "La diferencia entre decir 'nuestro producto ES el Benchmark Suite' y 'estamos investigando "
        "si el verdadero activo competitivo termina siendo el Benchmark Suite' es pequena en "
        "redaccion pero enorme en metodologia. La primera afirmacion cierra la puerta a "
        "evidencia contradictoria. La segunda la deja abierta. La honestidad metodologica "
        "exige la segunda formulacion."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Implicaciones practicas:")
    run.bold = True
    run.font.size = Pt(11)

    product_implications = [
        "El Benchmark Suite se desarrolla como hipotesis, no como producto confirmado.",
        "Si la evidencia muestra que la demanda no existe, el proyecto pivota sin crisis de identidad.",
        "La credibilidad del laboratorio es el activo real, independientemente de la forma que tome el producto.",
        "No se toman decisiones de arquitectura basadas en la suposicion de que el Benchmark Suite es el producto final.",
    ]
    for imp in product_implications:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(imp)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 11: OBJETIVOS OPERATIVOS INMEDIATOS (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("11. Objetivos Operativos Inmediatos", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Durante las proximas semanas, no se escribira una sola linea nueva del motor. "
        "El trabajo se concentra exclusivamente en cuatro objetivos. Estos objetivos "
        "representan la transicion desde un laboratorio con resultados preliminares hacia "
        "un laboratorio con evidencia solida. La credibilidad del laboratorio vale mas que "
        "agregar capacidades nuevas."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Objective 1
    h2 = doc.add_heading("11.1 Referencia Dorada: JPEG, PNG, PDF", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Conseguir que JPEG, PNG y PDF sean una referencia dorada absolutamente confiable. "
        "Esto significa que los tres parsers deben pasar no solo los 19 tests existentes, sino "
        "que deben ser robustos ante edge cases, datasets adversarios, y condiciones de dano "
        "extremo. La referencia dorada es el estandar contra el cual se mide todo lo demas: "
        "si los parsers base no son confiables, ningun experimento que dependa de ellos puede "
        "ser confiable. Este objetivo es el prerequisito de todos los demas."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    obj1_items = [
        "Verificar que los 19/19 tests pasan consistentemente en 30 ejecuciones consecutivas.",
        "Agregar tests de edge cases: archivos vacios, archivos de 1 byte, archivos con headers corruptos.",
        "Agregar tests adversarios: archivos que empiezan con la firma de un formato pero son otro.",
        "Documentar las limitaciones conocidas de cada parser.",
    ]
    for item in obj1_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Objective 2
    h2 = doc.add_heading("11.2 Alcanzar 3 Estrellas en Hipotesis Principales", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ejecutar todos los experimentos suficientes veces para alcanzar al menos 3 estrellas "
        "en las hipotesis principales, antes de abrir nuevas lineas de investigacion. Las "
        "hipotesis principales son H1.1 (MFT-First vs Carving), H6 (FQS no-binario) y H7 "
        "(RVS). Actualmente H1.1 esta en 3 estrellas, pero H6 y H7 estan en 1-2 estrellas. "
        "Elevarlas a 3 estrellas requiere ejecutar experimentos con multiples datasets distintos, "
        "no solo repeticiones del mismo dataset."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Hipotesis", "Estrellas Actuales", "Estrellas Objetivo", "Accion Requerida"],
        [
            ["H1.1", "***", "*** (mantener)", "Ejecutar en 3+ datasets distintos para confirmar"],
            ["H6", "**", "***", "Ejecutar FQS en 3+ datasets distintos"],
            ["H7", "*", "***", "Calibrar RVS con encuesta (ver 5.4) + ejecutar en 3+ datasets"],
            ["H2", "**", "***", "Repetir con carving completo (H8 resuelto)"],
        ],
        col_widths=[2.0, 2.5, 2.5, 7.0],
    )

    # Objective 3
    h2 = doc.add_heading("11.3 Validacion con Herramientas Externas", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Validar los resultados contra herramientas externas usando exactamente los mismos "
        "datasets. No solo PhotoRec, sino herramientas que representen cada familia de "
        "estrategias (ver Seccion 7.2). La validacion con multiples familias es esencial "
        "porque cada familia de herramientas opera bajo supuestos distintos, y si nuestros "
        "resultados solo coinciden con carving, eso solo valida nuestro carving, no nuestras "
        "estrategias MFT-first o hibridas."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    obj3_items = [
        "Preparar los datasets en formato compatible con cada herramienta externa.",
        "Ejecutar PhotoRec (carving), R-Studio (MFT-first), y DMDE (hybrid) sobre los mismos datasets.",
        "Comparar Recovery Rate, RVS y FQS de cada herramienta vs los motores del laboratorio.",
        "Documentar discrepancias y investigar sus causas.",
    ]
    for item in obj3_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Objective 4
    h2 = doc.add_heading("11.4 Primeros Datasets del Mundo Real", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Buscar los primeros datasets provenientes del mundo real, aunque sean pocos. "
        "Aqui empieza la transicion desde un laboratorio bien disenado hacia evidencia con "
        "mayor validez externa. Los datasets sinteticos son necesarios para el control "
        "experimental, pero no son suficientes: un laboratorio que solo funciona con datos "
        "sinteticos puede tener sesgos invisibles que solo aparecen con datos reales. "
        "El objetivo no es reemplazar los datasets sinteticos, sino complementarlos con "
        "algunos datos reales que permitan detectar sesgos del laboratorio."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    obj4_items = [
        "Identificar fuentes de imagenes de disco reales (anonymized, de dominio publico).",
        "Crear al menos 1 dataset real con ground truth verificable.",
        "Ejecutar los mismos experimentos sobre el dataset real y comparar con resultados sinteticos.",
        "Si los resultados difieren significativamente, investigar si el laboratorio tiene sesgos que solo aparecen con datos reales.",
    ]
    for item in obj4_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Si estos cuatro objetivos se cumplen, recien entonces tendria sentido volver a hablar "
        "de Motor C, IA, aprendizaje adaptativo o nuevos formatos. En este punto del proyecto, "
        "la credibilidad del laboratorio vale mas que agregar capacidades nuevas."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.italic = True

    # ─── Save ──────────────────────────────────────────────────────────────
    output_path = "/home/z/my-project/download/Research_Protocol_v1.1.docx"
    doc.save(output_path)
    print(f"Research Protocol v1.1 saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build_research_protocol()
    print(f"Generated: {path}")
