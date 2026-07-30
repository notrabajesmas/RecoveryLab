#!/usr/bin/env python3
"""
Generate Research Constitution — RecoveryLab
=============================================
A separate, short document (≤2 pages) containing only principles.
Not a protocol — a constitution. Defines how the lab behaves, not how to run experiments.

Based on the external auditor's final review (round 5):
> "Lo unico que todavia echo en falta no es una seccion del protocolo.
>  Es un documento completamente separado. Lo llamaria Research Constitution.
>  No tendria mas de dos paginas. Y contendria unicamente principios."
"""

import sys
import os

DOCX_SCRIPTS = os.path.join("/home/z/my-project/skills/docx", "scripts")
if DOCX_SCRIPTS not in sys.path:
    sys.path.insert(0, DOCX_SCRIPTS)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Color Palette ────────────────────────────────────────────────────────────
P = {
    "primary": "#162032",
    "body": "#1C2A3D",
    "secondary": "#5B6B7D",
    "accent": "#8B7E5A",
    "surface": "#F5F7FA",
    "white": "#FFFFFF",
    "red": "#C0392B",
    "green": "#27AE60",
    "blue": "#1A5276",
    "gold": "#C9A84C",
}

def c(hex_color):
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.replace("#", "")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def build_constitution():
    doc = Document()

    # ─── Page Setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = c(P["body"])

    # ═══════════════════════════════════════════════════════════════════════
    # COVER AREA (compact — this is a 2-page document)
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RECOVERYLAB")
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Research Constitution")
    run.font.size = Pt(26)
    run.font.color.rgb = c(P["primary"])
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Los principios que gobiernan el comportamiento del laboratorio")
    run.font.size = Pt(13)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.italic = True

    # Horizontal rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Este documento no explica como hacer experimentos.\n"
        "Explica como se comporta el laboratorio."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f"Version 1.0 | {datetime.datetime.now().strftime('%Y-%m-%d')} | Vinculado al Research Protocol v1.5")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # THE 8 PRINCIPLES
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("Los Ocho Principios", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Estos principios son la constitucion del laboratorio. No describen procedimientos: "
        "describen valores. No dicen como hacer un experimento: dicen como pensar sobre los "
        "experimentos. Cuando un procedimiento del protocolo entra en conflicto con un "
        "principio, el principio prevalece. Cuando una decision no esta cubierta por el "
        "protocolo, los principios guian la decision. Son la ultima instancia de apelacion."
    )
    run.font.size = Pt(11)

    principles = [
        {
            "number": "I",
            "title": "La evidencia tiene prioridad sobre la intuicion",
            "body": (
                "Cuando la intuicion y la evidencia entran en conflicto, la evidencia prevalece. "
                "Si todos 'sienten' que Motor C es mejor pero los datos no lo confirman, "
                "la conclusion es: los datos no lo confirman. La intuicion es una hipotesis, "
                "no una conclusion. La historia de la ciencia esta llena de intuiciones "
                "elegantes que resultaron ser incorrectas. La evidencia es la unica defensa "
                "contra el sesgo de confirmacion."
            )
        },
        {
            "number": "II",
            "title": "Ningun claim puede adelantarse a su evidencia",
            "body": (
                "No se publica una conclusion antes de tener la evidencia que la respalda. "
                "El Evidence Gate (5 niveles) existe exactamente para esto: el lenguaje "
                "debe reflejar honestamente la fuerza de la evidencia acumulada. Un claim "
                "en nivel OBSERVED no puede usar el lenguaje de un claim en nivel "
                "EXTERNALLY_VALIDATED. Adelantarse a la evidencia es la forma mas "
                "comun de autodefensa en la ciencia."
            )
        },
        {
            "number": "III",
            "title": "Las hipotesis no se modifican durante una fase",
            "body": (
                "Las hipotesis congeladas (H1.1-H8) no se reescriben. Si una hipotesis "
                "resulta incorrecta, se marca como REFUTADA. Si aparece una idea nueva, "
                "nace como H9, H10, etc. Modificar una hipotesis para que encaje con los "
                "resultados es la forma mas sutil de fraude cientifico: no se falsan datos, "
                "se falsa la pregunta para que la respuesta siempre sea correcta."
            )
        },
        {
            "number": "IV",
            "title": "Todo experimento debe poder reproducirse",
            "body": (
                "Si un experimento no se puede reproducir, no es un experimento: es una "
                "anecdota. La reproducibilidad es el criterio de demarcacion entre ciencia "
                "y anecdota. El script run_all.py es el contrato de reproducibilidad: "
                "cualquier persona debe poder ejecutarlo y obtener los mismos resultados. "
                "Si no se puede, el resultado no pertenece al laboratorio."
            )
        },
        {
            "number": "V",
            "title": "Los resultados negativos tienen el mismo valor que los positivos",
            "body": (
                "Un experimento que refuta una hipotesis es tan valioso como uno que la "
                "confirma. Saber que algo no funciona es informacion. No publicar resultados "
                "negativos sesga el cuerpo de conocimiento y hace que otros repitan los "
                "mismos errores. El Evidence Ledger registra todos los experimentos, "
                "incluidos los que produjeron resultados inconvenientes."
            )
        },
        {
            "number": "VI",
            "title": "Reducir deuda de evidencia tiene prioridad sobre agregar funcionalidades",
            "body": (
                "El proyecto ya tiene suficientes modulos. Lo que le falta es evidencia. "
                "Cada nuevo documento, modulo o algoritmo debe responder a una pregunta: "
                "Reduce una deuda de evidencia identificada? Si la respuesta es no, "
                "probablemente no pertenece a la Fase A. El Evidence Debt Registry "
                "es la unica lista de trabajo pendiente del proyecto."
            )
        },
        {
            "number": "VII",
            "title": "La complejidad es un costo cientifico",
            "body": (
                "Un protocolo con 27 secciones es mas dificil de seguir que uno con 10. "
                "Un laboratorio con 12 modulos es mas dificil de auditar que uno con 5. "
                "La complejidad no es solo un costo de ingenieria: es un costo cientifico "
                "porque reduce la probabilidad de que el equipo siga obedeciendo su propio "
                "protocolo. Si el equipo no sigue el protocolo, el protocolo es inutil. "
                "El proyecto no morira por errores. Morira por complejidad."
            )
        },
        {
            "number": "VIII",
            "title": "Toda interpretacion debe poder rastrearse hasta una observacion",
            "body": (
                "Ninguna interpretacion flota libre. Cada conclusion debe estar vinculada "
                "a una observacion especifica, y esa observacion debe estar vinculada a un "
                "experimento, y ese experimento debe estar en el Evidence Ledger. La cadena "
                "de trazabilidad es: claim -> experimento -> commit -> codigo. Si un eslabon "
                "falta, la interpretacion no tiene fundamento. La separacion de observacion "
                "y explicacion (6ta regla sagrada) es la manifestacion practica de este principio."
            )
        },
    ]

    for pr in principles:
        # Principle number + title
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"Principio {pr['number']}: ")
        run.font.size = Pt(13)
        run.font.color.rgb = c(P["blue"])
        run.bold = True
        run.font.name = "Calibri"
        run2 = p.add_run(pr['title'])
        run2.font.size = Pt(13)
        run2.font.color.rgb = c(P["primary"])
        run2.bold = True
        run2.font.name = "Calibri"

        # Body
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(pr['body'])
        run.font.size = Pt(11)
        run.font.color.rgb = c(P["body"])

    # ─── Footer area ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("_" * 60)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Este documento es complementario al Research Protocol v1.5 (Frozen for Phase A).\n"
        "El Protocol define procedimientos. La Constitucion define principios.\n"
        "En caso de conflicto, los principios prevalecen."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.italic = True

    # ─── Save ──────────────────────────────────────────────────────────────
    output_path = "/home/z/my-project/download/Research_Constitution_v1.0.docx"
    doc.save(output_path)
    print(f"Research Constitution saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build_constitution()
    print(f"Generated: {path}")
