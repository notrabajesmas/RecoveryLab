#!/usr/bin/env python3
"""
Generate Research Protocol v1.3 — RecoveryLab
===============================================
Third revision: Evidence Gate, three-level structure, meta-rule.

Key changes from v1.2:
  1. Added Section 14: Evidence Gate (language enforcement)
  2. Added Section 15: Three-Level Data Architecture (/data, /analysis, /claims)
  3. Added meta-rule: "No agregar una sola caracteristica nueva si no aumenta la calidad de la evidencia"
  4. KPI dashboard is now purely research-based, not code-based
"""

import sys, os
DOCX_SCRIPTS = os.path.join("/home/z/my-project/skills/docx", "scripts")
if DOCX_SCRIPTS not in sys.path:
    sys.path.insert(0, DOCX_SCRIPTS)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

P = {
    "primary": "#162032", "body": "#1C2A3D", "secondary": "#5B6B7D",
    "accent": "#8B7E5A", "surface": "#F5F7FA", "white": "#FFFFFF",
    "red": "#C0392B", "green": "#27AE60", "orange": "#E67E22", "blue": "#2980B9",
}

def c(hex_color):
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.replace("#", "")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = c(P["white"])
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, P["primary"])
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = c(P["body"])
            if row_idx % 2 == 1:
                set_cell_shading(cell, P["surface"])
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    return table

def add_callout(doc, title, text, color=P["accent"]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = c(color)
    run.font.name = "Calibri"
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = c(P["body"])
    run.italic = True


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = c(P["body"])

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RECOVERYLAB")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Research Protocol v1.3")
    run.font.size = Pt(28)
    run.font.color.rgb = c(P["primary"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Congelar la arquitectura, auditar la ciencia")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.italic = True

    # ── Meta-Rule Hero ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("META-REGLA")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("No agregar una sola caracteristica nueva\nsi no aumenta la calidad de la evidencia.")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["red"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = c(P["accent"])

    # KPI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("KPI: Resultados con 3+ estrellas: 2 / 15")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tablero de investigacion, no de software")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])

    # Metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(f"Documento vivo - Version 1.3 | {datetime.datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clasificacion: INTERNO")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["red"])

    # Changes
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Cambios respecto a v1.2")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    for change in [
        "Nueva Seccion 14: Evidence Gate (control de lenguaje)",
        "Nueva Seccion 15: Arquitectura de tres niveles (/data, /analysis, /claims)",
        "Meta-regla: no agregar features sin aumentar la evidencia",
        "KPI dashboard: ahora es de investigacion, no de software",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"  {change}")
        run.font.size = Pt(9)
        run.font.color.rgb = c(P["secondary"])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTIONS 1-12: Same as v1.2 (abbreviated — key content preserved)
    # ═══════════════════════════════════════════════════════════════════════

    # Section 1
    h = doc.add_heading("1. Pregunta Central", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])
    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto gira alrededor de una pregunta: "
        "\u00bfComo medir objetivamente la utilidad de una estrategia de recuperacion? "
        "Si el laboratorio demuestra que sus resultados son reproducibles, comparables y "
        "resisten auditorias externas, habremos construido algo que vale mucho mas que un "
        "motor de recuperacion: una infraestructura de evaluacion objetiva."
    )
    run.font.size = Pt(11)

    # Section 2
    h = doc.add_heading("2. Variables Experimentales", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])
    p = doc.add_paragraph()
    run = p.add_run(
        "Solo UNA variable independiente por experimento. Las variables independientes son: "
        "tipo de dano, nivel de dano, formato de archivo, estrategia, presupuesto de lectura. "
        "La metrica principal declarada es Overall Utility (RVS x FQS). "
        "Recovery Rate es metrica secundaria para comparabilidad con herramientas externas."
    )
    run.font.size = Pt(11)

    # Section 3
    h = doc.add_heading("3. Criterio de Exito", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])
    p = doc.add_paragraph()
    run = p.add_run(
        "Una estrategia A se considera superior a B unicamente si: "
        "(1) Overall Utility mejora por encima del umbral empirico (ver 3.1), "
        "(2) la diferencia es consistente en 10+ repeticiones, "
        "(3) se mantiene en 3+ datasets distintos, "
        "(4) no se debe exclusivamente a una dimension (RVS o FQS). "
        "El umbral provisional es 3% hasta la calibracion empirica (30 ejecuciones baseline)."
    )
    run.font.size = Pt(11)

    # Section 4
    h = doc.add_heading("4. Registro de Confianza", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])
    add_styled_table(
        doc,
        ["Estrellas", "Evidencia", "Significado"],
        [
            ["*", "Observacion aislada", "Preliminar"],
            ["**", "Repetido 10+ veces", "Estable"],
            ["***", "Datasets distintos", "Generalizable"],
            ["****", "Validado externamente", "Robusto"],
            ["*****", "Hardware real", "Definitivo"],
        ],
        col_widths=[2.0, 5.0, 6.0],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("KPI principal: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Resultados con 3+ estrellas. Actualmente 2/15. Objetivo Fase A: 5/15. "
        "Este es el verdadero cuello de botella del proyecto."
    )
    run2.font.size = Pt(11)

    # Section 5
    h = doc.add_heading("5. Decomposicion de Metricas", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Overall Utility = RVS (Valor) x FQS (Calidad)")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(
        "RVS mide que tan valioso era lo recuperado. FQS mide que tan utilizable quedo. "
        "La separacion permite diagnosticar por que gano un motor (VALUE-DRIVEN vs QUALITY-DRIVEN). "
        "RVS necesita calibracion con usuarios reales (Seccion 5.4): los pesos actuales son "
        "asignaciones del laboratorio, no datos calibrados."
    )
    run.font.size = Pt(11)

    # Sections 6-12 (abbreviated)
    for sec_num, sec_title in [
        (6, "Auditoria de Hipotesis"), (7, "Hoja de Ruta por Fases"),
        (8, "Formatos Congelados"), (9, "Regla de Oro"),
        (10, "Producto del Proyecto"), (11, "Objetivos Operativos Inmediatos"),
        (12, "Threats to Validity"),
    ]:
        h = doc.add_heading(f"{sec_num}. {sec_title}", level=1)
        for run in h.runs: run.font.color.rgb = c(P["primary"])
        p = doc.add_paragraph()
        run = p.add_run(f"[Ver Research Protocol v1.2 para el contenido completo de esta seccion. "
                        f"Los cambios en v1.3 estan en las Secciones 13, 14 y 15.]")
        run.font.size = Pt(10)
        run.font.color.rgb = c(P["secondary"])
        run.italic = True

    # Section 13
    h = doc.add_heading("13. Hypothesis Set v1.0 (Frozen)", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])
    p = doc.add_paragraph()
    run = p.add_run(
        "Las hipotesis congeladas (H1.1-H8) no se reescriben. Si una hipotesis resulta incorrecta, "
        "se marca como REFUTADA. Si aparece una idea nueva, nace como H9, H10, etc. "
        "H3 fue eliminada y documentada. Su contenido sustantivo vive en H4."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 14: EVIDENCE GATE (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("14. Evidence Gate (Control de Lenguaje)", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Gate es el cambio mas importante del protocolo. Es un mecanismo que controla "
        "el lenguaje permitido para describir un resultado, basandose en la cantidad y calidad de "
        "evidencia acumulada. Nadie puede escribir frases como 'Motor C demuestra...' hasta pasar "
        "por el gate. Mientras no llegue al cuarto casillero, queda prohibido escribir 'demuestra'. "
        "Solo: 'es consistente con' o 'observamos'. Parece un detalle de lenguaje, pero cambia "
        "completamente el rigor del proyecto."
    )
    run.font.size = Pt(11)

    # Gate levels
    h2 = doc.add_heading("14.1 Niveles del Evidence Gate", level=2)
    for run in h2.runs: run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Nivel", "Casillero", "Lenguaje permitido", "Lenguaje prohibido"],
        [
            ["1 - OBSERVED", "[X] observado",
             "observamos, es consistente con, aparece",
             "demuestra, prueba, confirma, establece"],
            ["2 - REPEATED", "[X] repetido",
             "es estable, se repite, es consistente en repeticiones",
             "demuestra, prueba, confirma, establece"],
            ["3 - REPRODUCIBLE", "[X] reproducible",
             "la evidencia sugiere, es reproducible, los datos indican",
             "demuestra, prueba, confirma, establece"],
            ["4 - EXTERNALLY VALIDATED", "[X] validado externamente",
             "demuestra, validado externamente, es robusto",
             "(ninguno — lenguaje completo permitido)"],
            ["5 - HARDWARE VALIDATED", "[X] validado en hardware real",
             "confirmado, predictivo del mundo real, definitivo",
             "(ninguno — lenguaje completo permitido)"],
        ],
        col_widths=[3.0, 3.0, 4.0, 4.0],
    )

    # Claim system
    h2 = doc.add_heading("14.2 Sistema de Claims", level=2)
    for run in h2.runs: run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cada afirmacion del proyecto se registra como un CLAIM con su propia ficha. "
        "Cada CLAIM tiene: evidencia vinculada (experimentos), amenazas vinculadas (threats), "
        "y el proximo experimento necesario. El Evidence Gate se implementa en el modulo "
        "evidence_gate.py y genera fichas automaticamente en el directorio /claims/. "
        "El sistema es mucho mas dificil de enganar que la narrativa suelta: cada claim "
        "esta vinculada a datos concretos, y el lenguaje esta restringido por el nivel "
        "de evidencia acumulada."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Claim", "Titulo", "Nivel", "Hipotesis", "Proximo paso"],
        [
            ["CLAIM-001", "MFT-First > Carving", "OBSERVED", "H1.1", "3+ datasets distintos"],
            ["CLAIM-002", "FQS no es binario", "OBSERVED", "H6", "3+ datasets distintos"],
            ["CLAIM-003", "Tesis > thumbnails (RVS)", "OBSERVED", "H7", "Encuesta de calibracion"],
            ["CLAIM-004", "Crossover 95% es artefacto", "OBSERVED", "H8", "Carving completo"],
            ["CLAIM-005", "Parsers referencia dorada", "OBSERVED", "-", "30 ejecuciones + edge cases"],
        ],
        col_widths=[2.0, 3.0, 2.0, 2.0, 4.0],
    )

    # Language enforcement
    h2 = doc.add_heading("14.3 Reglas de Lenguaje", level=2)
    for run in h2.runs: run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run("Las siguientes palabras estan PROHIBIDAS para claims en niveles 1-3:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["red"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(
        "demuestra / demonstrates | prueba / proves | confirma / confirms | "
        "establece / establishes | garantiza / guarantees | verifica / verifies"
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("En su lugar, usar:")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(
        "observamos | es consistente con | la evidencia sugiere | es reproducible | "
        "los datos indican | es estable en repeticiones"
    )
    run.font.size = Pt(11)

    add_callout(
        doc,
        "Por que importa el lenguaje?",
        "Si se permite escribir 'demuestra' con una sola observacion, se pierde la disciplina. "
        "El lenguaje influye en como pensamos: decir 'demuestra' cierra la puerta a la refutacion. "
        "Decir 'es consistente con' la mantiene abierta. El Evidence Gate obliga a que el lenguaje "
        "refleje honestamente la fuerza de la evidencia. No es un detalle cosmético: es la diferencia "
        "entre un laboratorio que se convence a si mismo y uno que se mantiene honesto.",
        P["red"]
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 15: THREE-LEVEL ARCHITECTURE (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("15. Arquitectura de Tres Niveles", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ahora mismo todavia hay pequenas mezclas entre observacion e interpretacion. "
        "Para evitar que el laboratorio se engane a si mismo, se separa absolutamente todo "
        "en tres niveles. Cada nivel tiene reglas estrictas sobre que puede contener. "
        "Es muchisimo mas dificil enganarse asi."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Directorio", "Contenido", "Reglas", "Prohibido"],
        [
            ["/data", "Solo datos crudos", "CSV, JSON, imagenes de disco, resultados de experimentos",
             "Nada de conclusiones, nada de interpretacion, nada de hipotesis"],
            ["/analysis", "Solo scripts estadisticos", "Analisis, graficos, calculos de significancia",
             "No motores, no parsers, no hipotesis, no claims"],
            ["/claims", "Solo fichas de claims", "Cada claim con su evidencia, amenazas y proximo experimento",
             "No datos, no codigo, no motores"],
        ],
        col_widths=[2.0, 3.0, 5.0, 4.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Principio: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Los datos no mienten. Las interpretaciones si pueden. Separando datos de interpretaciones, "
        "cualquier persona puede verificar los datos sin necesidad de aceptar las interpretaciones. "
        "Y cualquier persona puede desafiar las interpretaciones sin necesidad de cuestionar los datos. "
        "Esta separacion es la base de la reproducibilidad."
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 16: META-REGLA (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("16. Meta-Regla", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(
        "No agregar una sola caracteristica nueva\nsi no aumenta la calidad de la evidencia."
    )
    run.font.size = Pt(16)
    run.font.color.rgb = c(P["red"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta regla, por si sola, probablemente mantenga al proyecto en el camino cientifico "
        "que el protocolo intenta establecer. El proyecto ya tiene suficientes modulos: "
        "Dataset Builder, Corruptor, Judge, Carving, Motor MFT, Motor C, RVS, FQS, Protocol, "
        "Confidence Registry, Evidence Gate. Eso ya es muchisimo. El tablero ya no mide lineas "
        "de codigo, features o modulos. Mide unicamente la calidad de la evidencia."
    )
    run.font.size = Pt(11)

    # Research KPI Dashboard
    h2 = doc.add_heading("16.1 Tablero de Investigacion", level=2)
    for run in h2.runs: run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["KPI", "Actual", "Objetivo Fase A", "Objetivo Fase B"],
        [
            ["Resultados 3+ estrellas", "2/15", "5/15", "8/15"],
            ["Resultados 4+ estrellas", "0/15", "0/15", "3/15"],
            ["Hipotesis principales 3+", "1/4", "4/4", "4/4"],
            ["Amenazas mitigadas", "4/19", "8/19", "12/19"],
            ["Amenazas nuevas documentadas", "0", "Todas", "Todas"],
            ["Claims con evidencia suficiente", "0/5", "3/5", "5/5"],
            ["Umbral empirico calibrado", "No", "Si", "Si"],
            ["RVS calibrado con usuarios", "No", "No", "Si"],
        ],
        col_widths=[4.5, 2.5, 3.5, 3.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Este es un tablero de investigacion, no de software. No mide lineas de codigo, "
        "features o modulos. Mide la calidad de la evidencia. Si el numero de resultados "
        "con 3+ estrellas no sube, el proyecto no avanza, sin importar cuantas lineas de "
        "codigo se escriban."
    )
    run.font.size = Pt(11)
    run.italic = True

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 17: RVS CALIBRATION EXPERIMENT
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("17. Experimento de Calibracion de RVS", level=1)
    for run in h.runs: run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Toda la tesis del RVS descansa en una idea: recuperar una tesis vale mas que recuperar "
        "miniaturas. Eso parece obvio. Pero todavia no fue medido. Este experimento convierte "
        "esa intuicion en datos. No para validar el software. Para validar el modelo. "
        "El experimento muestra pares de archivos y pide a los usuarios que elijan cual "
        "preferirian recuperar. Despues se ajusta el RVS con esos datos. Ese experimento "
        "puede terminar siendo tan importante como cualquier benchmark."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("17.1 Diseno del Experimento", level=2)
    for run in h2.runs: run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Parametro", "Valor"],
        [
            ["ID", "EXP-RVS-CAL"],
            ["Tipo", "HUMANO (no software)"],
            ["Metodo", "Bradley-Terry pairwise comparison"],
            ["Pares de archivos", "12 pares (tesis vs fotos, RAW vs MP4, SQLite vs DOCX, etc.)"],
            ["Poblaciones", "5: fotografos, juridicos, tecnologia, domesticos, estudiantes"],
            ["Respuestas minimas", "30 por poblacion (150 total)"],
            ["Pregunta", "Si solo pudieras recuperar uno de estos archivos, cual elegirias?"],
            ["Output", "Pesos RVS calibrados por poblacion"],
            ["Estado", "DISSENADO (sin datos reales todavia)"],
        ],
        col_widths=[4.0, 9.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Los resultados de este experimento podrian mostrar que los pesos de RVS necesitan "
        "ser especificos por poblacion: un fotografo valora RAW mas que un estudiante, un "
        "estudio juridico valora contratos mas que un usuario domestico. Si es asi, el RVS "
        "deberia ser parametrico, no universal. Este experimento es el que puede transformar "
        "al RVS de una metrica del laboratorio a una metrica calibrada con comportamiento real."
    )
    run.font.size = Pt(11)

    # ─── Save ──────────────────────────────────────────────────────────────
    output_path = "/home/z/my-project/download/Research_Protocol_v1.3.docx"
    doc.save(output_path)
    print(f"Research Protocol v1.3 saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build()
    print(f"Generated: {path}")
