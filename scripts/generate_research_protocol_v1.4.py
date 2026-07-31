#!/usr/bin/env python3
"""
Generate Research Protocol v1.4 — RecoveryLab
===============================================
Fourth revision incorporating third round of external review.

Key changes from v1.3:
  1. Section 16: Separation of Observation and Explanation (6th sacred rule)
  2. Section 17: Evidence Ledger (per-experiment traceability)
  3. Section 18: Judge API Freeze (RVS/FQS/Overall Utility locked during Phase A)
  4. Section 19: Reproducibility Contract (run_all.py)
  5. Section 20: Reproducible Claims Ratio (RCR) as primary KPI
  6. Phase A strict regimen: only 4 activities for several weeks
  7. Cover page updated with RCR and 6 sacred rules
  8. Full content from v1.2 sections 1-13 preserved
"""

import sys
import os

DOCX_SCRIPTS = os.path.join("/home/z/my-project/skills/docx", "scripts")
if DOCX_SCRIPTS not in sys.path:
    sys.path.insert(0, DOCX_SCRIPTS)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Color Palette ────────────────────────────────────────────────────────────
P = {
    "primary": "#162032",
    "body": "#1C2A3D",
    "secondary": "#5B6B7D",
    "accent": "#8B7E5A",
    "surface": "#F5F7FA",
    "white": "#FFFFFF",
    "star_gold": "#C9A84C",
    "star_dim": "#B0B0B0",
    "red": "#C0392B",
    "green": "#27AE60",
    "orange": "#E67E22",
    "blue": "#2980B9",
    "threat_open": "#E74C3C",
    "threat_mitigated": "#F39C12",
    "threat_resolved": "#27AE60",
}

def c(hex_color):
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.replace("#", "")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = c(P["white"])
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, P["primary"])
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = c(P["body"])
            if row_idx % 2 == 1:
                set_cell_shading(cell, P["surface"])
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    return table

def add_callout_box(doc, title, text, color=P["accent"]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = c(color)
    run.font.name = "Calibri"
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = c(P["body"])
    run.italic = True


def build_research_protocol():
    doc = Document()

    # ─── Page Setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = c(P["body"])

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RECOVERYLAB")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Research Protocol v1.4")
    run.font.size = Pt(28)
    run.font.color.rgb = c(P["primary"])
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("El proyecto que necesita madurar, no reinventarse")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Revision incorporando auditoria externa (ronda 3)")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"

    # ── Meta-Rule ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("META-REGLA")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("No agregar una sola caracteristica nueva si no aumenta la calidad de la evidencia.")
    run.font.size = Pt(13)
    run.font.color.rgb = c(P["red"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    # ── KPI Heroes ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("KPI PRIMARIO")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Reproducible Claims Ratio (RCR)")
    run.font.size = Pt(16)
    run.font.color.rgb = c(P["blue"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Claims totales: 5 | Reproducibles: 0 | RCR = 0%")
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Objetivo Fase A: RCR >= 60%")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])

    # ── Secondary KPI ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Resultados con 3+ estrellas: 2 / 15")
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    # ── 6 Sacred Rules ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Las 6 Reglas Sagradas")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    rules = [
        "1. Congelar hipotesis — H1-H8 no se reescriben",
        "2. KPI de evidencia — no de codigo",
        "3. Tres niveles — /data, /analysis, /claims",
        "4. Evidence Gate — lenguaje controlado por nivel de evidencia",
        "5. RVS con usuarios — calibracion con datos reales",
        "6. Separar observacion de explicacion — lo mas estricto",
    ]
    for rule in rules:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(rule)
        run.font.size = Pt(9)
        run.font.color.rgb = c(P["body"])

    # Metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f"Documento vivo - Version 1.4 | {datetime.datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clasificacion: INTERNO - Solo para uso del laboratorio")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["red"])

    # Change log
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Cambios respecto a v1.3")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    changes = [
        "Nueva Seccion 16: Separation of Observation and Explanation (6ta regla sagrada)",
        "Nueva Seccion 17: Evidence Ledger (trazabilidad por experimento)",
        "Nueva Seccion 18: Judge API Freeze (RVS/FQS/Overall Utility congelados en Fase A)",
        "Nueva Seccion 19: Reproducibility Contract (run_all.py)",
        "Nueva Seccion 20: Reproducible Claims Ratio (RCR) como KPI primario",
        "Regimen estricto de Fase A: solo 4 actividades durante varias semanas",
        "KPI primario cambia de ★★★ a RCR (Reproducible Claims Ratio)",
    ]
    for change in changes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"  {change}")
        run.font.size = Pt(9)
        run.font.color.rgb = c(P["secondary"])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1: PREGUNTA CENTRAL
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("1. Pregunta Central del Proyecto", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto ya no gira alrededor del Motor C. Gira alrededor de una pregunta:"
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "\u00bfComo medir objetivamente la utilidad de una estrategia de recuperacion?"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta pregunta es mucho mas amplia que cualquier motor individual. Es mucho mas dificil de copiar. "
        "Y es la base sobre la cual se construye todo el proyecto. Si el laboratorio demuestra que sus resultados "
        "son reproducibles, comparables y resisten auditorias externas, habremos construido algo que vale "
        "mucho mas que un motor de recuperacion: una infraestructura de evaluacion objetiva que es "
        "extraordinariamente dificil de desarrollar y de replicar."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto ha pasado por tres etapas. La primera fue construir un recuperador mejor. La segunda fue "
        "descubrir que antes hacia falta construir un laboratorio. La tercera, la actual, es descubrir que antes "
        "del laboratorio hace falta construir un sistema que garantice que las conclusiones son confiables. "
        "Este protocolo es el documento vivo que define ese sistema. Cada seccion agrega una capa de rigor: "
        "desde las variables experimentales hasta la reproducibilidad completa, desde el Evidence Gate hasta "
        "la separacion estricta de observacion y explicacion."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2: VARIABLES EXPERIMENTALES
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("2. Variables Experimentales", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    h2 = doc.add_heading("2.1 Variable Independiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable independiente es aquella que el experimentador controla. Solo una debe cambiar "
        "por experimento. Si cambian dos a la vez, las conclusiones se vuelven ambiguas. Este es "
        "el principio fundamental del control experimental: cada experimento debe aislar exactamente "
        "un factor para que la relacion causal sea interpretable. Cuando dos variables cambian "
        "simultaneamente, es imposible determinar cual de ellas produjo el efecto observado, y "
        "cualquier conclusion queda invalidada por esta confoundacion."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Variable Independiente", "Valores", "Experimento Asociado"],
        [
            ["Tipo de dano", "MFT parcial, head crash, sectores intermitentes, ruido aleatorio", "H4 Matrix"],
            ["Nivel de dano", "0%, 10%, 20%, ..., 100%", "Crossover Curve"],
            ["Formato de archivo", "JPEG, PNG, PDF (congelados)", "H5 Per-Format"],
            ["Estrategia", "Carving, MFT-First, Hybrid, Motor C", "H1.1 / H2"],
            ["Presupuesto de lectura", "0, 100, 500, 1000, 5000 sectores", "H1.7 Budget"],
        ],
        col_widths=[4.0, 6.0, 5.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Regla critica: ")
    run.bold = True
    run.font.color.rgb = c(P["red"])
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Solo UNA variable independiente por experimento. Si se cambia el formato Y el nivel de dano "
        "simultaneamente, no se puede atribuir la diferencia a ninguno de los dos. Esta regla no es "
        "negociable: es la base de la inferencia causal."
    )
    run2.font.size = Pt(11)

    h2 = doc.add_heading("2.2 Variable Dependiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable dependiente es lo que se mide. Debe elegirse ANTES de ejecutar el experimento, "
        "no despues de ver los resultados. Cambiar la definicion de exito a posteriori invalida "
        "cualquier conclusion. El RecoveryLab tiene multiples metricas disponibles, pero cada "
        "experimento debe declarar cual es la principal. La declaracion previa es la unica defensa "
        "contra el sesgo de confirmacion: si siempre es posible encontrar una metrica que favorezca "
        "nuestro resultado, ninguna conclusion es confiable."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Metrica", "Tipo", "Definicion", "Rango"],
        [
            ["Recovery Rate", "Conteo", "Fraccion de archivos recuperados (SHA-256 match)", "0.0-1.0"],
            ["RVS", "Valor", "Valor recuperado = suma(archivo_i.valor x peso_i) / total_valor", "0.0-1.0"],
            ["FQS", "Calidad", "Calidad funcional = suma(archivo_i.score x tamano_i) / total_tamano", "0.0-1.0"],
            ["Overall Utility", "Compuesto", "RVS x FQS - utilidad total de la recuperacion", "0.0-1.0"],
            ["Read Efficiency", "Eficiencia", "Lecturas utiles / Lecturas totales", "0.0-1.0"],
            ["Time to First File", "Tiempo", "Lecturas antes del primer archivo recuperado", "0-N"],
        ],
        col_widths=[3.0, 2.0, 6.5, 1.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Metrica principal declarada: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Overall Utility (RVS x FQS) es la metrica principal del laboratorio. "
        "Recovery Rate se mantiene como metrica secundaria para comparabilidad con herramientas externas. "
        "Cada experimento debe declarar explicitamente cual metrica es la principal antes de ejecutar."
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3: CRITERIO DE EXITO
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("3. Criterio de Exito", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El criterio de exito define cuando una estrategia se considera superior a otra. "
        "Este criterio debe declararse ANTES de ejecutar el experimento. Si se define a posteriori, "
        "siempre es posible encontrar un criterio que favorezca cualquier resultado. La definicion "
        "previa es la unica defensa contra el sesgo de confirmacion."
    )
    run.font.size = Pt(11)

    add_callout_box(
        doc,
        "Cambio en v1.1: Umbral empirico en lugar de 5% fijo",
        "En la version 1.0 se establecia un umbral fijo del 5% de mejora en Overall Utility. "
        "Este numero era arbitrario: no surgia de datos experimentales, sino de una decision "
        "subjetiva. Un umbral de significancia debe derivarse de la variabilidad observada, "
        "el error de medicion y la repetibilidad de los experimentos.",
        P["orange"]
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Criterio de exito v1.4:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run("Una estrategia A se considera superior a una estrategia B unicamente si:")
    run.font.size = Pt(11)

    criteria = [
        "Overall Utility (RVS x FQS) mejora por encima del umbral empirico (ver Seccion 3.1) sobre la estrategia B",
        "La diferencia es consistente en al menos 10 repeticiones con el mismo dataset",
        "La diferencia se mantiene en al menos 3 datasets distintos",
        "La mejora no se debe exclusivamente a una dimension (RVS o FQS) sino a ambas",
    ]
    for i, criterion in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f"{i}. {criterion}")
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Si solo se cumple el criterio 1 pero no los criterios 2-4, el resultado se clasifica como "
        "observacion aislada (1 estrella) y no como evidencia solida. La consistencia es tan importante "
        "como la magnitud de la mejora. Un resultado que no se replica es un resultado que no existe "
        "desde el punto de vista cientifico."
    )
    run.font.size = Pt(11)

    # 3.1 Umbral Empirico
    h2 = doc.add_heading("3.1 Calibracion del Umbral Empirico", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El umbral de significancia no debe ser un numero elegido arbitrariamente. Debe surgir de los "
        "datos del laboratorio. El procedimiento para calibrarlo es el siguiente: se ejecuta el mismo "
        "experimento (misma estrategia, mismo dataset, mismas condiciones) multiples veces, y se mide "
        "la variabilidad natural de Overall Utility. Si las fluctuaciones entre ejecuciones son del 3%, "
        "entonces una mejora del 3% no es significativa: esta dentro del ruido. El umbral debe ser "
        "al menos 2 veces la desviacion estandar de la variabilidad observada, para que una mejora "
        "declarada tenga una probabilidad baja de ser producto del azar."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Procedimiento de calibracion:")
    run.bold = True
    run.font.size = Pt(11)

    cal_steps = [
        "Ejecutar el experimento baseline (Carving sobre JPEG, dataset estandar) 30 veces seguidas.",
        "Registrar Overall Utility en cada ejecucion. Calcular media y desviacion estandar.",
        "Definir umbral = max(2 x desviacion_estandar, 0.01). Es decir: al menos 2 sigma, con un piso absoluto del 1%.",
        "Repetir para cada formato (PNG, PDF) y cada estrategia (Carving, MFT-First, Hybrid).",
        "El umbral final es el maximo entre todos los umbrales calculados. Esto garantiza que el umbral es conservador.",
        "Si la variabilidad es muy baja (determinista), el umbral se mantiene en el piso del 1%.",
    ]
    for i, step in enumerate(cal_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {step}")
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Estado actual: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Este procedimiento aun no se ha ejecutado. El umbral definitivo se determinara "
        "despues de las 30 ejecuciones baseline. Hasta entonces, se usa un umbral provisional "
        "del 3% (conservador, basado en la variabilidad observada en ejecuciones previas)."
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4: REGISTRO DE CONFIANZA
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("4. Registro de Confianza", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El registro de confianza es la forma mas transparente de comunicar la solidez cientifica "
        "de cada resultado. En lugar de porcentajes arbitrarios de confianza, utiliza un sistema "
        "de estrellas que mapea directamente al tipo de evidencia acumulada. Este sistema es "
        "mas honesto que cualquier numero: un lector puede ver inmediatamente que evidencia "
        "respalda un resultado y que tan lejos esta de ser validado externamente."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Estrellas", "Evidencia Requerida", "Significado"],
        [
            ["*", "Observacion aislada (1 run, 1 dataset)", "Preliminar - no citar como conclusion"],
            ["**", "Repetido 10+ veces (determinista, mismo dataset)", "Estable - pero no generalizado"],
            ["***", "Repetido con datasets distintos", "Generalizable - pero no validado externamente"],
            ["****", "Validado con herramientas externas (multiples familias)", "Robusto - comparable con el estado del arte"],
            ["*****", "Validado con hardware real", "Definitivo - predictivo del mundo real"],
        ],
        col_widths=[2.0, 6.0, 5.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Reglas del registro: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Las estrellas solo pueden subir (acumulacion de evidencia), nunca bajar. "
        "Si aparece evidencia contradictoria, el resultado se marca como CONTESTADO pero "
        "no pierde estrellas. El lector debe decidir como interpretar la contradiccion. "
        "Ademas, cada estrella debe estar vinculada a un registro de los experimentos que la soportan."
    )
    run2.font.size = Pt(11)

    # ── ★★★ KPI ──
    h2 = doc.add_heading("4.1 KPI de Progreso: Resultados con 3+ Estrellas", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El indicador mas importante del proyecto no es el Recovery Rate, ni siquiera el Overall Utility. "
        "Es la cantidad de resultados que alcanzan 3 estrellas o mas. Este numero refleja la solidez "
        "real del laboratorio: no cuantos resultados tenemos, sino cuantos estan lo suficientemente "
        "validados como para ser citados como conclusiones. Un resultado con 3 estrellas ha sido "
        "repetido con datasets distintos y es generalizable. Un resultado con 1-2 estrellas es "
        "preliminar y no debe usarse como base para decisiones arquitectonicas."
    )
    run.font.size = Pt(11)

    # KPI Dashboard
    add_styled_table(
        doc,
        ["KPI", "Actual", "Objetivo Fase A", "Objetivo Fase B"],
        [
            ["Resultados con 3+ estrellas", "2 / 15", "5 / 15", "8 / 15"],
            ["Resultados con 4+ estrellas", "0 / 15", "0 / 15", "3 / 15"],
            ["Hipotesis principales con 3+", "1 / 4 (H1.1)", "4 / 4", "4 / 4"],
            ["Umbral empirico calibrado", "No", "Si", "Si"],
            ["RVS calibrado con usuarios", "No", "No", "Si"],
            ["RCR (Reproducible Claims Ratio)", "0%", ">= 60%", ">= 80%"],
        ],
        col_widths=[5.0, 3.0, 3.0, 3.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Antes de ampliar el motor, antes de agregar IA y antes de incorporar nuevos formatos, "
        "el objetivo es que varias hipotesis importantes lleguen a 3 estrellas mediante repeticion, "
        "multiples datasets y validacion consistente. A partir de ahi, el laboratorio empezara a "
        "tener una base experimental solida sobre la cual construir el resto del proyecto."
    )
    run.font.size = Pt(11)
    run.italic = True

    # Current state table
    h2 = doc.add_heading("4.2 Estado Actual del Registro", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Resultado", "Estrellas", "Resumen"],
        [
            ["H1.1", "*** (3/5)", "MFT-First supera a Carving consistentemente. 100/100 escenarios."],
            ["H1.2", "* (1/5)", "Umbral de cambio de estrategia no determinado."],
            ["H2", "** (2/5)", "Crossover existe pero es artefacto del carving limitado."],
            ["H4", "* (1/5)", "Matriz preliminar con 3 celdas de muchas."],
            ["H5", "* (1/5)", "Sin experimento sistematico por formato."],
            ["H6", "** (2/5)", "FunctionalValidator implementado, 19/19 tests."],
            ["H7", "* (1/5)", "RVS implementado, sin validacion de usuarios (ver Seccion 5.4)."],
            ["H8", "* (1/5)", "Crossover al 95% es artefacto, no descubrimiento."],
            ["RVS", "** (2/5)", "4 dimensiones implementadas, sin calibracion externa."],
            ["FQS", "** (2/5)", "5 niveles funcionales, 19/19 tests."],
            ["WFS", "* (1/5)", "Concepto definido (RVS x FQS), aun no integrado en Judge."],
            ["BLOCKER-001", "*** (3/5)", "Resuelto: comparaciones previas invalidas."],
            ["H1.5", "* (1/5)", "Gaps: sin fragmentacion, sin INDX, sin jerarquia."],
            ["H1.6", "** (2/5)", "100 ejecuciones, resultados deterministas."],
            ["H1.7", "* (1/5)", "Motor C apenas supera MFT-First (5% support)."],
        ],
        col_widths=[2.5, 2.0, 10.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Resumen: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "8 resultados a 1 estrella, 5 a 2 estrellas, 2 a 3 estrellas, 0 a 4 estrellas, "
        "0 a 5 estrellas. La mayoria de los resultados son preliminares. "
        "El cuello de botella real es: cuantos resultados alcanzan 3 estrellas o mas?"
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 5: DECOMPOSICION DE METRICAS
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("5. Decomposicion de Metricas", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La metrica compuesta anterior (WFS) mezclaba dos dimensiones distintas: "
        "que archivos se recuperaron y que tan bien se recuperaron. La separacion "
        "es esencial porque un motor puede ganar por dos razones muy diferentes, "
        "y la interpretacion cambia completamente segun cual sea. Un motor que recupera "
        "la tesis pero con pixeles corruptos tiene alto RVS y bajo FQS; un motor que "
        "recupera perfectamente 200 thumbnails tiene bajo RVS y alto FQS. Sin esta "
        "separacion, ambas victorias se ven identicas, y eso es cientificamente inaceptable."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Overall Utility = RVS (Valor) x FQS (Calidad)")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.bold = True

    # RVS
    h2 = doc.add_heading("5.1 RVS - Recovery Value Score (Que se recupero)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El RVS mide el valor de lo recuperado, ponderado por la importancia de cada archivo. "
        "No es lo mismo recuperar la tesis que 200 thumbnails. El RVS incorpora cuatro dimensiones: "
        "valor intrinseco del archivo, probabilidad de reemplazo, tiempo de recreacion, e impacto "
        "emocional de la perdida. Un motor que recupera la tesis pero pierde 200 thumbnails tiene "
        "mayor RVS que uno que hace lo contrario, porque la tesis tiene mas valor, es mas dificil "
        "de reemplazar, requiere mas tiempo para recrear, y su perdida tiene mayor impacto emocional."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Tipo de Archivo", "Valor", "Reemplazo", "Recreacion", "Emocional", "RVS Peso"],
        [
            ["tesis.docx", "100", "0.05", "500h", "100", "100"],
            ["sqlite.db", "95", "0.02", "200h", "50", "95"],
            ["foto familiar", "90", "0.00", "0h", "100", "90"],
            ["video vacaciones", "70", "0.00", "0h", "80", "70"],
            ["PSD proyecto", "65", "0.10", "40h", "30", "65"],
            ["RAW foto", "60", "0.00", "0h", "60", "60"],
            ["MP4 descargado", "20", "0.90", "1h", "5", "20"],
            ["Linux ISO", "2", "1.00", "0.5h", "0", "2"],
            ["thumbnail", "1", "1.00", "0h", "0", "1"],
        ],
        col_widths=[3.0, 1.5, 1.5, 2.0, 2.0, 1.5],
    )

    add_callout_box(
        doc,
        "Limitacion reconocida: RVS necesita calibracion con usuarios reales",
        "La tabla de valores anterior es una asignacion razonable hecha por el laboratorio. "
        "Pero sigue siendo una decision interna. No ha sido calibrada con comportamiento real "
        "de usuarios. Ver Seccion 5.4 para el plan de calibracion. Ademas, el Judge API esta "
        "congelado durante la Fase A (ver Seccion 18): si los pesos de RVS necesitan cambiar, "
        "se crea Judge v1.1 y se reejecutan los experimentos afectados. Nunca se mezclan versiones.",
        P["orange"]
    )

    # FQS
    h2 = doc.add_heading("5.2 FQS - Functional Quality Score (Que tan bien se recupero)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El FQS mide la calidad funcional de la recuperacion. No es lo mismo recuperar un "
        "JPEG bit-perfect que uno con 2 pixeles corruptos. El FQS reemplaza el binario "
        "SHA-256 coincide/no coincide con un espectro de calidad funcional que refleja "
        "mejor la experiencia real del usuario. Un archivo que abre y funciona tiene valor "
        "incluso si su checksum no coincide exactamente. La ventaja del FQS es que captura "
        "gradaciones de calidad que el binario pasa por alto: un JPEG con 2 pixeles malos "
        "no es lo mismo que un JPEG completamente corrupto, pero ambos son 'fallo' en el "
        "modelo binario."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Nivel", "Score", "Definicion", "Ejemplo"],
        [
            ["FULL", "1.0", "SHA-256 coincide exactamente", "Copia bit-perfect del original"],
            ["FUNCTIONAL", "0.8", "Archivo funciona con dano menor", "JPEG con 2 pixeles corruptos"],
            ["PARTIAL", "0.5", "Funciona parcialmente, perdida de datos", "MP4 que se reproduce hasta la mitad"],
            ["DEGRADED", "0.2", "Contenido accesible pero daniado", "DOCX que abre pero perdio imagenes"],
            ["FAILED", "0.0", "Completamente inutilizable", "Datos aleatorios sin estructura"],
        ],
        col_widths=[2.5, 1.5, 5.0, 5.0],
    )

    # RVS Calibration Plan
    h2 = doc.add_heading("5.3 Diagnostico: Por que gano un motor?", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La separacion RVS x FQS permite diagnosticar por que un motor gano. "
        "Un motor puede ganar por dos razones fundamentalmente distintas, y la "
        "interpretacion cambia completamente segun cual sea la dominante."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["RVS", "FQS", "Diagnostico", "Interpretacion"],
        [
            ["Alto", "Alto", "STRONG", "Recupero archivos importantes con buena calidad"],
            ["Alto", "Bajo", "VALUE-DRIVEN", "Recupero archivos importantes pero con mala calidad"],
            ["Bajo", "Alto", "QUALITY-DRIVEN", "Recupero archivos bien pero no eran importantes"],
            ["Bajo", "Bajo", "WEAK", "Fallo en ambas dimensiones"],
        ],
        col_widths=[2.0, 2.0, 3.0, 7.0],
    )

    # RVS Calibration Plan
    h2 = doc.add_heading("5.4 Plan de Calibracion de RVS con Usuarios Reales", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La tabla de valores de RVS (Seccion 5.1) es una asignacion razonable hecha por el "
        "laboratorio, pero sigue siendo una decision interna. No ha sido validada con comportamiento "
        "real de usuarios. Esto es el punto mas debil del protocolo actual: los pesos de RVS "
        "determinan que estrategia se considera superior, y si esos pesos estan mal calibrados, "
        "todas las conclusiones del laboratorio se basan en una premisa incorrecta."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        "Para que el RVS deje de ser un modelo disenado por el laboratorio y pase a estar "
        "calibrado con comportamiento real, se necesita una encuesta de calibracion. La pregunta "
        "central de la encuesta es simple: 'Si solo pudieras recuperar uno de estos archivos, "
        "cual elegirias?' La respuesta a esta pregunta, repetida con suficiente diversidad de "
        "usuarios, produce una ranking de valor que reemplaza la tabla actual."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Poblaciones objetivo para la encuesta:")
    run.bold = True
    run.font.size = Pt(11)

    survey_pops = [
        "Fotografos profesionales: valoran RAW y fotos familiares, desprecian ISO y thumbnails.",
        "Estudios juridicos: valoran documentos y bases de datos, desprecian thumbnails.",
        "Empresas de tecnologia: valoran bases de datos y codigo, desprecian ISO.",
        "Usuarios domesticos: valoran fotos y videos familiares, desprecian archivos de sistema.",
        "Estudiantes: valoran tesis y documentos, desprecian thumbnails.",
    ]
    for pop in survey_pops:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(pop)
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Metodologia de la encuesta:")
    run.bold = True
    run.font.size = Pt(11)

    survey_methods = [
        "Presentar pares de archivos y pedir al usuario que elija cual recuperaria primero.",
        "Usar el metodo de comparacion por pares (Bradley-Terry) para obtener un ranking continuo.",
        "Minimo 30 respuestas por poblacion para obtener estimaciones estables.",
        "Comparar los pesos obtenidos con la tabla actual. Si difieren significativamente, actualizar la tabla y recalcular resultados previos.",
        "Publicar los datos de la encuesta como parte del Benchmark Suite para transparencia.",
    ]
    for method in survey_methods:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(method)
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Estado: ")
    run.bold = True
    run.font.color.rgb = c(P["orange"])
    run.font.size = Pt(11)
    run2 = p.add_run(
        "La encuesta no se ha realizado. Es una actividad de la Fase A (regimen estricto). "
        "Hasta que se complete, los pesos de RVS deben tratarse como supuestos, no como "
        "hechos. Cada resultado que dependa de RVS debe incluir un descargo: 'Los pesos de "
        "valor son asignaciones del laboratorio, no estan calibrados con usuarios reales.'"
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 6: AUDITORIA DE HIPOTESIS
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("6. Auditoria de Hipotesis", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cada hipotesis debe declarar explicitamente su variable independiente, variable dependiente, "
        "y criterio de exito antes de ejecutar el experimento. Si estos no estan definidos, "
        "la hipotesis no es testeable y debe ser reformulada. Una hipotesis sin criterio de "
        "refutacion no es una hipotesis cientifica: es una opinion. El protocolo exige que cada "
        "hipotesis defina explicitamente que resultado la refutaria, no solo que resultado la "
        "confirmaria."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Hipotesis", "Variable Independiente", "Variable Dependiente", "Criterio de Exito", "Confiabilidad"],
        [
            ["H1.1", "Estrategia (Carving vs MFT-First)", "Overall Utility", "MFT-First > Carving por umbral empirico+ en 3+ datasets", "***"],
            ["H1.2", "Confianza en metadatos (0-100%)", "Estrategia optima", "Cambio de estrategia en umbral definido", "*"],
            ["H2", "Nivel de dano MFT (0-100%)", "Recovery Rate", "Crossover observable con carving completo", "**"],
            ["H4", "Tipo de dano", "Mejor estrategia", "Matriz completa con >80% celdas llenas", "*"],
            ["H5", "Formato de archivo", "Recovery Rate por formato", "Diferencia >10% entre formatos", "*"],
            ["H6", "Nivel de dano en archivo", "FQS", "Espectro no-binario con >2 niveles", "**"],
            ["H7", "Tipo de archivo recuperado", "RVS", "Tesis > 200 thumbnails (RVS calibrado)", "*"],
            ["H8", "Motor de carving (limitado vs completo)", "Crossover point", "Cambio significativo en crossover", "*"],
        ],
        col_widths=[1.5, 3.5, 2.5, 4.0, 1.5],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 7: HOJA DE RUTA POR FASES
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("7. Hoja de Ruta por Fases", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    h2 = doc.add_heading("7.1 Fase A - Regimen Estricto (Ahora Mismo)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_callout_box(
        doc,
        "REGLA DE FASE A: No escribir ni un motor nuevo durante varias semanas.",
        "Solo se hacen cuatro cosas. Si despues de esas cuatro cosas los resultados siguen "
        "sosteniendose, recien entonces se vuelve a desarrollar algoritmos de recuperacion. "
        "Dejar de preguntarse 'que mas puedo construir?' y empezar a preguntarse 'que evidencia "
        "necesito para que otra persona crea lo que encontre?'. Esa es la pregunta correcta "
        "para la etapa en la que esta el proyecto ahora.",
        P["red"]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "La Fase A tiene un regimen estricto de cuatro actividades. No se escriben nuevos motores, "
        "no se agregan nuevos formatos, no se construyen nuevas funcionalidades. El unico trabajo "
        "es consolidar la evidencia existente y hacerla reproducible. La credibilidad del laboratorio "
        "vale mas que cualquier capacidad nueva."
    )
    run.font.size = Pt(11)

    phase_a_items = [
        "1. Ejecutar las 30 corridas baseline para obtener la variabilidad real del sistema (Seccion 3.1).",
        "2. Ejecutar la calibracion de RVS con usuarios reales (Seccion 5.4 y Seccion 21).",
        "3. Validar los parsers JPEG/PNG/PDF contra herramientas externas (Seccion 7.2).",
        "4. Convertir todos los experimentos en reproducibles de principio a fin (Seccion 19).",
    ]
    for item in phase_a_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Si despues de estas cuatro actividades los resultados siguen sosteniendose, recien "
        "ahi se vuelve a desarrollar algoritmos de recuperacion. En otras palabras: dejar de "
        "preguntarse 'que mas puedo construir?' y empezar a preguntarse 'que evidencia necesito "
        "para que otra persona crea lo que encontre?'. Esa es la pregunta correcta para la etapa "
        "en la que esta el proyecto ahora."
    )
    run.font.size = Pt(11)
    run.italic = True

    h2 = doc.add_heading("7.2 Fase B - Validacion Externa", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Validar el laboratorio contra herramientas reales. Pero no basta con comparar contra "
        "una sola herramienta. PhotoRec representa carving puro, pero el laboratorio evalua "
        "multiples familias de estrategias: MFT-first, hybridas, orquestadores, presupuestos "
        "de lectura. Para que la validacion externa sea significativa, necesitamos comparar "
        "contra herramientas que representen CADA familia de estrategias, no solo una."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Familia de Estrategia", "Herramienta Representativa", "Que valida"],
        [
            ["Carving puro", "PhotoRec / TestDisk", "Si nuestro carving es comparable al estado del arte"],
            ["MFT-first", "R-Studio / ReclaiMe", "Si nuestras estrategias MFT son comparables a herramientas comerciales"],
            ["Hybrida", "DMDE (Disk Drill)", "Si nuestras estrategias hibridas son comparables"],
            ["Orquestador", "UFS Explorer", "Si nuestro Motor C es comparable a herramientas que adaptan estrategia"],
        ],
        col_widths=[3.5, 4.0, 6.5],
    )

    h2 = doc.add_heading("7.3 Fase C - Expansion Controlada", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Solo entonces incorporar un formato nuevo (por ejemplo MP4) y exigirle el mismo "
        "estandar de calidad y pruebas que ya se alcanzo con JPEG, PNG y PDF. Cada nuevo "
        "formato debe pasar por el mismo proceso: parser impecable, 19+ tests de validacion, "
        "experimentos por formato, y estabilidad verificada. La regla es simple: un parser "
        "excelente de JPEG tiene muchisimo mas valor cientifico que diez parsers aceptables."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 8: FORMATOS CONGELADOS
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("8. Formatos Congelados - Referencia Dorada", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Los tres parsers actuales (JPEG, PNG, PDF) tienen 19/19 tests pasados. "
        "Esto vale oro. Antes de agregar un nuevo formato, estos tres deben convertirse "
        "en una referencia absoluta. La calidad de los parsers existentes es la base "
        "sobre la cual se construye la credibilidad del laboratorio. Un parser con 19/19 "
        "tests no es solo un parser que funciona: es un parser que se puede usar como "
        "patron de referencia para calibrar otros parsers y para validar que los "
        "experimentos de recuperacion no estan sesgados por defectos del parser."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Formato", "Parser", "Tests", "Estado", "Accion"],
        [
            ["JPEG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PNG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PDF", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["MP4", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["DOCX", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["SQLite", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["CR2/NEF", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["TXT", "Imposible", "0/0", "IMPOSIBLE", "Sin firma ni footer - no es carvable"],
        ],
        col_widths=[2.0, 3.0, 1.5, 2.5, 5.0],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 9: REGLA DE ORO
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("9. Regla de Oro", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Por cada 500 lineas nuevas de codigo de recuperacion, "
        "agregar al menos 500 lineas de validacion, pruebas y metricas."
    )
    run.font.size = Pt(13)
    run.font.color.rgb = c(P["accent"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta regla no es una sugerencia. Es la unica defensa contra la tentacion de agregar "
        "funcionalidades sin validarlas. La calidad del laboratorio es su activo mas valioso. "
        "Un motor con 3 parsers impecables y 19/19 tests es mas cientificamente valioso que "
        "un motor con 17 parsers mediocres y tests insuficientes. Si realmente se mantiene esta "
        "disciplina durante uno o dos anos, RecoveryLab podria terminar siendo mucho mas confiable "
        "que muchos proyectos open source donde el codigo crece mucho mas rapido que las pruebas."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 10: PRODUCTO
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("10. Producto del Proyecto", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    add_callout_box(
        doc,
        "Producto como hipotesis de negocio, no como declaracion",
        "En la version 1.0 se afirmaba: 'El producto ya no es un motor sino un Benchmark Suite.' "
        "Esto puede ser cierto, pero todavia es una hipotesis de negocio. No esta demostrado que "
        "exista suficiente demanda por una plataforma de benchmarking para recuperacion de datos. "
        "La redaccion correcta es una pregunta, no una afirmacion.",
        P["orange"]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Estamos investigando si el verdadero activo competitivo del proyecto termina siendo "
        "el RecoveryLab Benchmark Suite: una plataforma objetiva para evaluar estrategias de "
        "recuperacion de datos. Esta hipotesis de negocio se basa en la observacion de que "
        "medir objetivamente la utilidad de una estrategia de recuperacion es un problema "
        "mas amplio y mas dificil de replicar que construir un motor individual. Pero no "
        "esta demostrado que exista suficiente demanda por esta plataforma. La validacion "
        "de esta hipotesis de negocio es parte del trabajo futuro, no una conclusion actual."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 11: OBJETIVOS OPERATIVOS INMEDIATOS
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("11. Objetivos Operativos Inmediatos", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Durante las proximas semanas, no se escribira una sola linea nueva del motor. "
        "El trabajo se concentra exclusivamente en cuatro objetivos. La credibilidad del "
        "laboratorio vale mas que agregar capacidades nuevas."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.1 30 Corridas Baseline", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ejecutar el experimento baseline (Carving sobre JPEG, dataset estandar) 30 veces seguidas "
        "para obtener la variabilidad real del sistema. Sin esta variabilidad, el umbral de "
        "significancia es arbitrario. Las 30 corridas baseline son el cimiento sobre el cual "
        "se construye todo el edicion estadistico del protocolo. Sin ellas, no se puede "
        "calibrar el umbral empirico (Seccion 3.1), no se puede calcular la potencia estadistica, "
        "y no se puede distinguir una mejora real del ruido. Cada corrida debe documentarse "
        "en el Evidence Ledger (Seccion 17) con su ID, fecha, seed, commit y resultados."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.2 Calibracion de RVS con Usuarios", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ejecutar la encuesta de calibracion de RVS (Seccion 5.4 y Seccion 21). Este es el "
        "eslabon mas debil del protocolo: los pesos de RVS determinan que estrategia se "
        "considera superior, y si esos pesos estan mal calibrados, todas las conclusiones "
        "del laboratorio se basan en una premisa incorrecta. La encuesta debe alcanzar al "
        "menos 30 respuestas por poblacion en 5 poblaciones distintas (150 respuestas total)."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.3 Validacion de Parsers contra Herramientas Externas", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Validar los parsers JPEG/PNG/PDF contra herramientas externas usando exactamente los "
        "mismos datasets. No solo PhotoRec, sino herramientas que representen cada familia de "
        "estrategias (ver Seccion 7.2). La validacion con multiples familias es esencial "
        "porque cada familia opera bajo supuestos distintos. Si los resultados del laboratorio "
        "coinciden con las herramientas de la misma familia, los claims ganan el nivel "
        "EXTERNALLY_VALIDATED del Evidence Gate."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.4 Convertir Todo en Reproducible", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cualquier persona debe poder ejecutar 'git clone ... && python run_all.py' y obtener "
        "exactamente los mismos datasets, los mismos CSV, las mismas figuras, y el mismo "
        "registro de claims, sin intervencion humana. Ese script termina siendo casi tan "
        "importante como el laboratorio. Ver Seccion 19: Reproducibility Contract para "
        "los detalles completos del contrato de reproducibilidad."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 12: THREATS TO VALIDITY
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("12. Threats to Validity (Amenazas a la Validez)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta seccion es una practica estandar en trabajos cientificos. Su proposito es explicitar "
        "los factores que podrian invalidar las conclusiones del laboratorio. No es una admission "
        "de debilidad: es una demostracion de rigor. Un laboratorio que no reconoce sus amenazas "
        "a la validez es un laboratorio que no las ha pensado. Cada vez que aparezca una amenaza "
        "nueva, no se elimina: se agrega aqui junto con su estado (ABIERTA, MITIGADA o RESUELTA). "
        "Con el tiempo esta seccion puede convertirse en una de las mas importantes del protocolo."
    )
    run.font.size = Pt(11)

    # 12.1 Internal Validity
    h2 = doc.add_heading("12.1 Validez Interna", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La validez interna pregunta: existe algun sesgo dentro del propio RecoveryLab que podria "
        "explicar los resultados observados? Es la pregunta mas incomoda porque implica que el "
        "laboratorio podria estar enganandose a si mismo. Las amenazas a la validez interna son "
        "las mas peligrosas porque son invisibles desde adentro: el experimentador no detecta sus "
        "propios sesgos."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Motores conocen ground truth",
             "Los motores podrian acceder a informacion que no tendrian en un escenario real",
             "MITIGADA",
             "Motor A (Carving) no lee MFT. Motor B (MFT-First) lee MFT legible pero no el ground truth directo."],
            ["Datasets favorecen una estrategia",
             "Los datasets sinteticos podrian estar disenados de forma que favorezcan MFT-First o Carving",
             "ABIERTA",
             "Verificar distribucion de dano. Usar multiples generadores. Fase B: comparar con datasets reales."],
            ["Parser sesgado",
             "El parser de carving podria tener bugs que lo hagan peor de lo que seria un carving real",
             "MITIGADA",
             "19/19 tests de carving impecable. Fase B: comparar con PhotoRec."],
            ["Confirmacion de hipotesis",
             "El experimentador podria disenar experimentos que favorezcan sus hipotesis",
             "MITIGADA",
             "Criterio de exito declarado antes de ejecutar. Umbral empirico. Registro de confianza publico."],
            ["Carving limitado",
             "El carving actual solo detecta 3 formatos. Esto infla artificialmente el crossover point.",
             "ABIERTA",
             "H8 registrado. Congelado hasta Fase C. Crossover actual marcado como artefacto."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    # 12.2 External Validity
    h2 = doc.add_heading("12.2 Validez Externa", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Solo datasets sinteticos",
             "Todos los resultados provienen de imagenes de disco generadas artificialmente.",
             "ABIERTA",
             "Objetivo 11.4: buscar primeros datasets reales. Fase B: validacion con herramientas externas."],
            ["Solo NTFS",
             "El laboratorio solo evalua NTFS. Otros filesystems quedan fuera.",
             "ABIERTA",
             "Registrar como limitacion. No expandir filesystems hasta que NTFS sea solido."],
            ["Solo dano simulado",
             "Los modelos de dano son simulaciones. El dano real puede tener patrones diferentes.",
             "ABIERTA",
             "Fase B: validacion con herramientas que se usan en escenarios reales. Objetivo 5 estrellas: hardware real."],
            ["Tamano de imagen limitado",
             "Las imagenes de disco son de 10 MB. Los discos reales son de 100 GB a varios TB.",
             "ABIERTA",
             "Registrar como limitacion. Evaluar escalabilidad en futuras fases."],
            ["Solo carving de 3 formatos",
             "JPEG, PNG, PDF. Los discos reales contienen decenas de formatos.",
             "ABIERTA",
             "Congelado hasta Fase C. Parser excelente de 3 formatos > parser mediocre de 17."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    # 12.3 Statistical Validity
    h2 = doc.add_heading("12.3 Validez Estadistica", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Repeticiones insuficientes",
             "Muchos resultados tienen solo 1-3 ejecuciones. No hay potencia estadistica.",
             "MITIGADA",
             "H1.1: 100 ejecuciones. H1.6: 100 ejecuciones. Objetivo: minimo 10 repeticiones por resultado."],
            ["Sin potencia estadistica formal",
             "No se ha calculado la potencia estadistica de ningun experimento.",
             "ABIERTA",
             "Calcular potencia despues de las 30 ejecuciones baseline (Seccion 3.1)."],
            ["Variabilidad desconocida",
             "No se conoce la variabilidad natural de las metricas.",
             "MITIGADA",
             "Seccion 3.1: calibracion empirica del umbral. Umbral provisional de 3% hasta calibracion."],
            ["Multiples comparaciones",
             "Al comparar multiples estrategias, la probabilidad de falsos positivos aumenta.",
             "ABIERTA",
             "Registrar como amenaza. Considerar correccion de Bonferroni o FDR en futuras analisis."],
            ["Resultados deterministas",
             "Los resultados son deterministas (misma semilla = mismo resultado).",
             "ABIERTA",
             "Usar multiples semillas y datasets. La consistencia entre datasets distintos es la que importa."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    # 12.4 Construct Validity
    h2 = doc.add_heading("12.4 Validez Constructiva", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Overall Utility no representa utilidad real",
             "RVS x FQS es una formula matematica. No esta demostrado que represente la utilidad percibida.",
             "ABIERTA",
             "Fase B: encuesta de calibracion. Si los usuarios no coinciden, la metrica debe reformularse."],
            ["RVS no representa valor real",
             "Los pesos de RVS son asignados por el laboratorio, no calibrados con usuarios.",
             "ABIERTA",
             "Seccion 5.4: plan de calibracion con Bradley-Terry. Descargo explicito en cada resultado que use RVS."],
            ["FQS no representa calidad funcional real",
             "Los 5 niveles son definidos por el laboratorio. Un usuario podria discrepar.",
             "ABIERTA",
             "Encuesta de calibracion para FQS (simultanea con la de RVS)."],
            ["Recuperacion binaria vs funcional",
             "Los umbrales de FQS (0.8, 0.5, 0.2) son arbitrarios.",
             "ABIERTA",
             "Calibrar los umbrales de FQS con usuarios reales."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Resumen de amenazas: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "4 amenazas MITIGADAS, 12 ABIERTAS, 0 RESUELTAS. La mayoria de las amenazas abiertas "
        "requieren validacion externa (Fase B) o datos reales (Fase B/C). Esto es esperado en "
        "un laboratorio en Fase A. La honestidad sobre estas amenazas es mas valiosa que "
        "ignorarlas: un lector que ve las amenazas puede evaluar la confianza que merecen los "
        "resultados. Un lector que no las ve no puede."
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 13: HYPOTHESIS SET v1.0 (FROZEN)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("13. Hypothesis Set v1.0 (Frozen)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "A lo largo del proyecto, las hipotesis fueron reformuladas multiples veces. Esto fue util "
        "durante la exploracion, pero llega un punto donde la reformulacion continua destruye la "
        "trazabilidad del historial cientifico. Si una hipotesis se reescribe cada vez que los "
        "resultados no la favorecen, nunca puede ser refutada, y un experimento que no puede ser "
        "refutado no es un experimento. Por eso se declara el conjunto actual de hipotesis como "
        "Hypothesis Set v1.0 (Frozen). Las hipotesis congeladas no se reescriben. Si aparece una "
        "idea nueva, nace como H9, H10, etc. Si una hipotesis resulta incorrecta, se marca como "
        "REFUTADA, no se modifica retroactivamente."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Reglas del congelamiento:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])

    frozen_rules = [
        "Las hipotesis congeladas (H1.1 a H8) no se reescriben. Su formulacion es definitiva.",
        "Si una hipotesis resulta incorrecta, se marca como REFUTADA con la evidencia que la refuta.",
        "Si aparece una idea nueva, se registra como H9, H10, H11, etc. No se fusiona con una existente.",
        "Si una hipotesis necesita mayor granularidad, se crean sub-hipotesis (H9.1, H9.2).",
        "El historial de cada hipotesis se mantiene en el registro.",
        "Cualquier cambio posterior al congelamiento se documenta con fecha, razon y aprobacion.",
    ]
    for i, rule in enumerate(frozen_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    # Current frozen set
    h2 = doc.add_heading("13.1 Conjunto Congelado Actual", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["ID", "Hipotesis (formulacion congelada)", "Estado", "Estrellas"],
        [
            ["H1.1", "Priorizar metadatos recuperables reduce significativamente el costo de adquisicion cuando los metadatos son suficientemente confiables", "ACTIVA", "***"],
            ["H1.2", "Cuando la confianza en metadatos baja de un umbral, la estrategia optima cambia de priorizacion a hibrida", "ACTIVA", "*"],
            ["H1.5", "Los gaps actuales (sin fragmentacion, sin INDX, sin jerarquia) limitan la capacidad de evaluacion", "ACTIVA", "*"],
            ["H1.6", "Los resultados son deterministas: misma semilla produce mismo resultado", "ACTIVA", "**"],
            ["H1.7", "Motor C (orquestador) supera a MFT-First cuando el presupuesto de lectura es limitado", "ACTIVA", "*"],
            ["H2", "Existe una frontera observable donde la estrategia optima cambia segun el estado del medio", "CONTESTADA", "**"],
            ["H4", "La estrategia optima depende del tipo de dano (Damage x Strategy Matrix)", "ACTIVA", "*"],
            ["H5", "La eficacia de recuperacion varia significativamente por formato de archivo", "ACTIVA", "*"],
            ["H6", "La recuperacion funcional no es binaria: existe un espectro de calidad funcional (FQS)", "ACTIVA", "**"],
            ["H7", "El RVS predice la satisfaccion del usuario", "ACTIVA", "*"],
            ["H8", "El crossover al 95% MFT damage es un artefacto del carving limitado", "ACTIVA", "*"],
        ],
        col_widths=[1.5, 8.0, 2.0, 1.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Nota: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "H3 fue eliminada del conjunto congelado. Su contenido sustantivo vive ahora en H4. "
        "No se reescribe: se elimina y se explica por que."
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 14: EVIDENCE GATE
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("14. Evidence Gate (Control de Lenguaje)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Gate es un mecanismo que controla el lenguaje permitido para describir "
        "un resultado, basandose en la cantidad y calidad de evidencia acumulada. Nadie puede "
        "escribir frases como 'Motor C demuestra...' hasta pasar por el gate. Mientras no "
        "llegue al cuarto casillero, queda prohibido escribir 'demuestra'. Solo: 'es consistente "
        "con' o 'observamos'. Parece un detalle de lenguaje, pero cambia completamente el "
        "rigor del proyecto. El Evidence Gate obliga a que el lenguaje refleje honestamente "
        "la fuerza de la evidencia."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("14.1 Niveles del Evidence Gate", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Nivel", "Casillero", "Lenguaje permitido", "Lenguaje prohibido"],
        [
            ["1 - OBSERVED", "[X] observado", "observamos, es consistente con, aparece", "demuestra, prueba, confirma, establece"],
            ["2 - REPEATED", "[X] repetido", "es estable, se repite, es consistente en repeticiones", "demuestra, prueba, confirma, establece"],
            ["3 - REPRODUCIBLE", "[X] reproducible", "la evidencia sugiere, es reproducible, los datos indican", "demuestra, prueba, confirma, establece"],
            ["4 - EXTERNALLY VALIDATED", "[X] validado externamente", "demuestra, validado externamente, es robusto", "(ninguno)"],
            ["5 - HARDWARE VALIDATED", "[X] validado en hardware real", "confirmado, predictivo del mundo real, definitivo", "(ninguno)"],
        ],
        col_widths=[3.0, 3.0, 4.0, 4.0],
    )

    h2 = doc.add_heading("14.2 Sistema de Claims", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cada afirmacion del proyecto se registra como un CLAIM con su propia ficha. "
        "Cada CLAIM tiene: evidencia vinculada (experimentos), amenazas vinculadas (threats), "
        "y el proximo experimento necesario. El sistema es mucho mas dificil de enganar que "
        "la narrativa suelta: cada claim esta vinculada a datos concretos, y el lenguaje "
        "esta restringido por el nivel de evidencia acumulada."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Claim", "Titulo", "Nivel", "Hipotesis", "Proximo paso"],
        [
            ["CLAIM-001", "MFT-First > Carving", "OBSERVED", "H1.1", "3+ datasets distintos"],
            ["CLAIM-002", "FQS no es binario", "OBSERVED", "H6", "3+ datasets distintos"],
            ["CLAIM-003", "Tesis > thumbnails (RVS)", "OBSERVED", "H7", "Encuesta de calibracion"],
            ["CLAIM-004", "Crossover 95% es artefacto", "OBSERVED", "H8", "Carving completo"],
            ["CLAIM-005", "Parsers referencia dorada", "OBSERVED", "-", "30 ejecuciones + edge cases"],
        ],
        col_widths=[2.0, 3.0, 2.0, 2.0, 4.0],
    )

    h2 = doc.add_heading("14.3 Reglas de Lenguaje", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run("Las siguientes palabras estan PROHIBIDAS para claims en niveles 1-3:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["red"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(
        "demuestra / demonstrates | prueba / proves | confirma / confirms | "
        "establece / establishes | garantiza / guarantees | verifica / verifies"
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("En su lugar, usar:")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(
        "observamos | es consistente con | la evidencia sugiere | es reproducible | "
        "los datos indican | es estable en repeticiones"
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 15: THREE-LEVEL ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("15. Arquitectura de Tres Niveles", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Para evitar que el laboratorio se engane a si mismo, se separa absolutamente todo "
        "en tres niveles. Cada nivel tiene reglas estrictas sobre que puede contener. "
        "Es muchisimo mas dificil enganarse asi. Los datos no mienten. Las interpretaciones "
        "si pueden. Separando datos de interpretaciones, cualquier persona puede verificar "
        "los datos sin necesidad de aceptar las interpretaciones. Y cualquier persona puede "
        "desafiar las interpretaciones sin necesidad de cuestionar los datos. Esta separacion "
        "es la base de la reproducibilidad."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Directorio", "Contenido", "Reglas", "Prohibido"],
        [
            ["/data", "Solo datos crudos", "CSV, JSON, imagenes de disco, resultados de experimentos",
             "Nada de conclusiones, nada de interpretacion, nada de hipotesis"],
            ["/analysis", "Solo scripts estadisticos", "Analisis, graficos, calculos de significancia",
             "No motores, no parsers, no hipotesis, no claims"],
            ["/claims", "Solo fichas de claims", "Cada claim con su evidencia, amenazas y proximo experimento",
             "No datos, no codigo, no motores"],
        ],
        col_widths=[2.0, 3.0, 5.0, 4.0],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 16: SEPARATION OF OBSERVATION AND EXPLANATION (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("16. Separation of Observation and Explanation (6ta Regla Sagrada)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    add_callout_box(
        doc,
        "Esta regla es mas estricta que el Evidence Gate.",
        "Incluso usando 'observamos...' todavia es facil colar interpretacion. "
        "Por ejemplo: 'Observamos que Motor C elige correctamente la estrategia.' Eso ya interpreta. "
        "La observacion pura seria: 'En 27/30 ejecuciones Motor C selecciono carving cuando el dano "
        "MFT supero X%.' Despues, en otra seccion distinta: 'Esto es consistente con la hipotesis "
        "de que...' Esta separacion es muy usada en ciencia experimental y reduce muchisimo el "
        "sesgo del investigador.",
        P["red"]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Gate (Seccion 14) controla el lenguaje segun el nivel de evidencia. Pero incluso "
        "respetando el gate, es facil colar interpretacion en la observacion. La frase 'Observamos que "
        "Motor C elige correctamente la estrategia' parece inocente, pero 'correctamente' ya es una "
        "interpretacion: presupone que existe una estrategia correcta y que el motor la encontro. La "
        "observacion pura seria: 'En 27/30 ejecuciones Motor C selecciono carving cuando el dano MFT "
        "supero X%.' Ni mas ni menos. Despues, en otra seccion distinta, se puede escribir: 'Esto es "
        "consistente con la hipotesis de que el motor selecciona la estrategia apropiada cuando el dano "
        "MFT supera el umbral.' La separacion entre observacion y explicacion es una practica estandar "
        "en ciencia experimental y reduce dramaticamente el sesgo del investigador."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("16.1 La Regla", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Toda afirmacion experimental debe separarse en dos partes:\n"
        "1. OBSERVACION (datos puros, sin interpretacion)\n"
        "2. EXPLICACION (interpretacion, siempre en otro parrafo)"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])
    run.bold = True

    h2 = doc.add_heading("16.2 Ejemplos", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Tipo", "Texto", "Valido?"],
        [
            ["Observacion con interpretacion", "Observamos que Motor C elige correctamente la estrategia", "NO"],
            ["Observacion pura", "En 27/30 ejecuciones Motor C selecciono carving cuando el dano MFT supero X%", "SI"],
            ["Explicacion separada", "Esto es consistente con la hipotesis de que el motor adapta su estrategia al dano", "SI"],
            ["Observacion con interpretacion", "MFT-First supera a Carving porque aprovecha los metadatos", "NO"],
            ["Observacion pura", "MFT-First recupero 85% de archivos vs 6.7% de Carving en 100 ejecuciones con MFT intacto", "SI"],
            ["Explicacion separada", "La diferencia es consistente con la hipotesis de que los metadatos facilitan la recuperacion", "SI"],
        ],
        col_widths=[3.5, 8.0, 1.5],
    )

    h2 = doc.add_heading("16.3 Principios de la Observacion Pura", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    obs_principles = [
        "Solo numeros, conteos, porcentajes y medidas. Ningun adjetivo calificativo ('correctamente', 'mejor', 'apropiadamente').",
        "Ningun porque. La observacion no explica por que algo ocurrio, solo que ocurrio.",
        "Ningun termino evaluativo. 'Elige correctamente' presupone que existe una eleccion correcta. 'Selecciono carving' no presupone nada.",
        "Siempre cuantificar. 'En 27/30 ejecuciones' es observacion. 'Generalmente selecciona' es interpretacion.",
        "Despues de la observacion, en otro parrafo, se puede escribir: 'Esto es consistente con la hipotesis de que...'",
        "La explicacion siempre debe estar vinculada a una hipotesis. No se explica sin hipotesis.",
    ]
    for i, principle in enumerate(obs_principles, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {principle}")
        run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 17: EVIDENCE LEDGER (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("17. Evidence Ledger (Registro de Evidencia)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Ledger es un registro extremadamente simple de cada experimento ejecutado. "
        "No es un PDF. Es un archivo de texto con campos fijos que documenta exactamente que se "
        "ejecuto, cuando, con que datos, con que version del codigo, y que claims afecta. "
        "Despues ningun claim puede citar una evidencia que no exista en ese ledger. Esto hace "
        "la trazabilidad casi perfecta: cualquier persona puede verificar que un claim se basa "
        "en un experimento real, y puede reproducir ese experimento si tiene el commit, la seed "
        "y el dataset."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("17.1 Formato del Evidence Ledger", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Evidence ID: EXP-0031\n"
        "Fecha: 2026-08-02\n"
        "Dataset: dataset_0042\n"
        "Seed: 18472\n"
        "Motor: Motor C\n"
        "Commit: 7c913d2\n"
        "Resultados: Overall Utility = 0.73, RVS = 0.85, FQS = 0.86\n"
        "Claims afectados: CLAIM-002, CLAIM-004\n"
        "Threats: T03"
    )
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    run.font.color.rgb = c(P["body"])

    h2 = doc.add_heading("17.2 Reglas del Evidence Ledger", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    ledger_rules = [
        "Cada experimento ejecutado genera exactamente una entrada en el ledger. No se omiten experimentos, incluso si los resultados son inconvenientes.",
        "Ningun claim puede citar una evidencia que no exista en el ledger. Si el ledger no tiene la entrada, la evidencia no existe.",
        "Los campos Evidence ID, Fecha, Dataset, Seed, Motor y Commit son obligatorios. Si falta alguno, la entrada es invalida.",
        "El campo Commit debe apuntar a un commit real del repositorio. Esto permite reproducir exactamente el experimento.",
        "El campo Claims afectados vincula el experimento con los claims que soporta o refuta. Un claim sin evidencia en el ledger es un claim sin fundamento.",
        "El campo Threats vincula el experimento con las amenazas a la validez que podrian afectarlo.",
        "El ledger se almacena en /data/evidence_ledger.csv (o .json) y es de solo lectura una vez escrito. No se editan entradas, solo se agregan.",
    ]
    for i, rule in enumerate(ledger_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    h2 = doc.add_heading("17.3 Trazabilidad Completa", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Ledger crea una cadena de trazabilidad que va desde el claim hasta el experimento, "
        "desde el experimento hasta el commit, y desde el commit hasta el codigo exacto que se ejecuto. "
        "Esto significa que cualquier persona puede, en principio, clonar el repositorio en el commit "
        "indicado, ejecutar el experimento con la seed y el dataset indicados, y obtener los mismos "
        "resultados. Si no los obtiene, hay un problema de reproducibilidad que debe investigarse. "
        "El ledger es la herramienta mas poderosa del laboratorio para la integridad cientifica: "
        "no permite que un claim flote sin ancla, y no permite que un experimento se pierda en el olvido."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 18: JUDGE API FREEZE (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("18. Judge API Freeze (Congelamiento del Judge)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ahora que existen RVS, FQS y Overall Utility, el Judge tiene una API definida. "
        "Si la forma de medir cambia a mitad de los experimentos, comparar resultados antiguos "
        "con nuevos se vuelve mucho mas dificil. Por eso se declara el Judge API v1.0 como "
        "congelado durante toda la Fase A. Si alguna metrica necesita evolucionar, se crea "
        "Judge v1.1 y se vuelven a ejecutar los experimentos afectados. Nunca se mezclan "
        "resultados de versiones distintas del Judge."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("18.1 Judge API v1.0 (Frozen)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Componente", "Version", "Definicion", "Estado"],
        [
            ["RVS", "v1.0", "4 dimensiones: valor intrinseco, reemplazo, recreacion, emocional", "CONGELADO"],
            ["FQS", "v1.0", "5 niveles: FULL(1.0), FUNCTIONAL(0.8), PARTIAL(0.5), DEGRADED(0.2), FAILED(0.0)", "CONGELADO"],
            ["Overall Utility", "v1.0", "RVS x FQS", "CONGELADO"],
            ["Confidence Registry", "v1.0", "5 estrellas: observacion aislada a hardware validado", "CONGELADO"],
            ["Evidence Gate", "v1.0", "5 niveles: OBSERVED a HARDWARE_VALIDATED", "CONGELADO"],
        ],
        col_widths=[3.0, 1.5, 6.0, 2.5],
    )

    h2 = doc.add_heading("18.2 Reglas del Congelamiento", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    judge_rules = [
        "Durante la Fase A, el Judge API v1.0 no cambia. Ni los pesos de RVS, ni los umbrales de FQS, ni la formula de Overall Utility.",
        "Si se descubre que una metrica necesita cambiar (por ejemplo, los pesos de RVS despues de la calibracion con usuarios), se crea Judge API v1.1.",
        "Al crear Judge v1.1, todos los experimentos previos que usaron v1.0 deben reejecutarse con v1.1. Los resultados de v1.0 se archivan pero no se mezclan con los de v1.1.",
        "Nunca se comparan resultados de versiones distintas del Judge. Si un experimento se ejecuto con v1.0 y otro con v1.1, no son comparables.",
        "El cambio de version se documenta con fecha, razon y los experimentos afectados. El Evidence Ledger (Seccion 17) registra la version del Judge usada en cada experimento.",
        "La calibracion de RVS con usuarios (Seccion 5.4) probablemente requiera Judge v1.1. Cuando eso ocurra, se reejecutan todos los experimentos. Esto es un costo, no un problema: es el costo de la integridad cientifica.",
    ]
    for i, rule in enumerate(judge_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 19: REPRODUCIBILITY CONTRACT (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("19. Reproducibility Contract (Contrato de Reproducibilidad)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El siguiente cuello de botella del proyecto ya no es el laboratorio. Es la reproducibilidad. "
        "Cualquier persona debe poder ejecutar algo asi como 'git clone ... && python run_all.py' y "
        "obtener exactamente los mismos datasets, los mismos CSV, las mismas figuras, y el mismo "
        "registro de claims. Sin intervencion humana. Ese script termina siendo casi tan importante "
        "como el laboratorio. La reproducibilidad no es un lujo: es la diferencia entre un proyecto "
        "que produce resultados y un proyecto que produce conocimiento confiable."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("19.1 El Script run_all.py", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El script run_all.py es el contrato de reproducibilidad del proyecto. Debe ser capaz de "
        "reproducir todos los resultados del laboratorio desde cero, sin intervencion humana. "
        "Esto significa que el script debe: generar los datasets, ejecutar los experimentos, "
        "producir los CSV de resultados, generar las figuras, y actualizar el registro de claims. "
        "Si el script falla en alguna de estas tareas, el resultado correspondiente no es reproducible."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Componente", "Requisito", "Verificacion"],
        [
            ["Datasets", "Generados deterministicamente con seeds fijas", "Misma seed = mismo dataset"],
            ["Experimentos", "Ejecutados con la misma version del codigo (commit)", "Mismo commit = mismo resultado"],
            ["CSV de resultados", "Producidos automaticamente por los experimentos", "Mismo CSV byte por byte"],
            ["Figuras", "Generadas desde los CSV, sin intervencion humana", "Misma figura visualmente"],
            ["Claims", "Actualizados automaticamente desde los resultados", "Mismo estado de claims"],
            ["Evidence Ledger", "Poblado automaticamente por cada experimento", "Mismo ledger"],
        ],
        col_widths=[3.0, 5.5, 5.5],
    )

    h2 = doc.add_heading("19.2 Reglas del Contrato de Reproducibilidad", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    repro_rules = [
        "Todos los datasets deben ser generados deterministicamente. Si se usa una semilla fija, el dataset debe ser identico en cualquier maquina.",
        "Todos los experimentos deben ser deterministas dado el mismo commit, la misma seed y el mismo dataset. Si no lo son, el resultado no es reproducible.",
        "Los CSV de resultados deben ser producidos automaticamente. No se permiten copiar-y-pegar manual.",
        "Las figuras deben ser generadas desde los CSV. No se permiten ediciones manuales de figuras.",
        "El registro de claims debe ser actualizable automaticamente. No se permite actualizar claims manualmente sin pasar por el Evidence Gate.",
        "El Evidence Ledger debe ser poblado automaticamente por cada experimento. No se permiten entradas manuales.",
        "Si un resultado no es reproducible por run_all.py, no es un resultado del laboratorio. Es una observacion informal.",
    ]
    for i, rule in enumerate(repro_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    h2 = doc.add_heading("19.3 Verificacion de Reproducibilidad", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La reproducibilidad se verifica ejecutando run_all.py en una maquina limpia (sin estado "
        "previo) y comparando los resultados con los de referencia. Si los resultados no coinciden, "
        "hay un problema de reproducibilidad que debe resolverse antes de que el claim pueda "
        "avanzar en el Evidence Gate. La verificacion de reproducibilidad es un requisito para "
        "que un claim pase del nivel OBSERVED al nivel REPEATED (y eventualmente al nivel "
        "REPRODUCIBLE). Sin verificacion de reproducibilidad, ningun claim puede avanzar."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 20: REPRODUCIBLE CLAIMS RATIO (RCR) (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("20. Reproducible Claims Ratio (RCR) - KPI Primario", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El KPI que mas importa no es la cantidad de estrellas. Es otro. Se llama Reproducible "
        "Claims Ratio (RCR). Mide exactamente lo que uno quiere que aumente: cuantos claims "
        "sobreviven cuando cualquiera repite el experimento. No importa si hay veinte hipotesis. "
        "Importa cuantas sobreviven cuando cualquiera repite el experimento. Un proyecto con 3 "
        "claims reproducibles es mas cientificamente solido que un proyecto con 20 claims no "
        "reproducibles. El RCR mide la resiliencia del conocimiento producido por el laboratorio."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("20.1 Definicion", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("RCR = Claims Reproducibles / Claims Totales")
    run.font.size = Pt(16)
    run.font.color.rgb = c(P["blue"])
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Un claim es 'reproducible' cuando otra persona puede ejecutar el experimento asociado "
        "(usando la informacion del Evidence Ledger) y obtener resultados consistentes con los "
        "originales. 'Consistente' no significa identico: significa que las conclusiones del "
        "claim se mantienen. Si el claim dice 'MFT-First supera a Carving', la reproduccion "
        "debe confirmar que MFT-First supera a Carving, aunque los numeros exactos puedan "
        "diferir ligeramente."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("20.2 Estado Actual y Objetivos", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Metrica", "Actual", "Objetivo Fase A", "Objetivo Fase B"],
        [
            ["Claims totales", "5", "5-8", "8-12"],
            ["Claims reproducibles", "0", "3-5", "6-10"],
            ["RCR", "0%", ">= 60%", ">= 80%"],
            ["Claims con nivel REPRODUCIBLE+", "0", "3", "5"],
            ["Resultados con 3+ estrellas", "2/15", "5/15", "8/15"],
            ["Amenazas mitigadas", "4/19", "8/19", "12/19"],
            ["Umbral empirico calibrado", "No", "Si", "Si"],
            ["RVS calibrado con usuarios", "No", "Si", "Si"],
            ["Judge API version", "v1.0", "v1.0 (frozen)", "v1.1 (si necesario)"],
        ],
        col_widths=[4.5, 2.5, 3.5, 3.5],
    )

    h2 = doc.add_heading("20.3 Por que RCR es el KPI primario", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El RCR mide algo mas fundamental que la cantidad de estrellas. Las estrellas miden "
        "la calidad de la evidencia acumulada internamente. El RCR mide la resiliencia de esa "
        "evidencia frente a la verificacion externa. Un claim con 3 estrellas que no se puede "
        "reproducir es un claim con 3 estrellas fragiles. Un claim con 2 estrellas que se "
        "reproduce consistentemente es un claim con 2 estrellas solidas. El RCR responde a "
        "la pregunta mas importante del proyecto: 'que evidencia necesito para que otra persona "
        "crea lo que encontre?' Si la respuesta es 'ninguna', el RCR es 0%. Si la respuesta es "
        "'cualquiera puede verificarlo', el RCR es 100%. El objetivo del proyecto es acercarse "
        "a 100%."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 21: RVS CALIBRATION EXPERIMENT
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("21. Experimento de Calibracion de RVS", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Toda la tesis del RVS descansa en una idea: recuperar una tesis vale mas que recuperar "
        "miniaturas. Eso parece obvio. Pero todavia no fue medido. Este experimento convierte "
        "esa intuicion en datos. No para validar el software. Para validar el modelo. "
        "El experimento muestra pares de archivos y pide a los usuarios que elijan cual "
        "preferirian recuperar. Despues se ajusta el RVS con esos datos. Ese experimento "
        "puede terminar siendo tan importante como cualquier benchmark."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("21.1 Diseno del Experimento", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Parametro", "Valor"],
        [
            ["ID", "EXP-RVS-CAL"],
            ["Tipo", "HUMANO (no software)"],
            ["Metodo", "Bradley-Terry pairwise comparison"],
            ["Pares de archivos", "12 pares (tesis vs fotos, RAW vs MP4, SQLite vs DOCX, etc.)"],
            ["Poblaciones", "5: fotografos, juridicos, tecnologia, domesticos, estudiantes"],
            ["Respuestas minimas", "30 por poblacion (150 total)"],
            ["Pregunta", "Si solo pudieras recuperar uno de estos archivos, cual elegirias?"],
            ["Output", "Pesos RVS calibrados por poblacion"],
            ["Estado", "DISSENADO (sin datos reales todavia)"],
        ],
        col_widths=[4.0, 9.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Los resultados de este experimento podrian mostrar que los pesos de RVS necesitan "
        "ser especificos por poblacion: un fotografo valora RAW mas que un estudiante, un "
        "estudio juridico valora contratos mas que un usuario domestico. Si es asi, el RVS "
        "deberia ser parametrico, no universal. Este experimento es el que puede transformar "
        "al RVS de una metrica del laboratorio a una metrica calibrada con comportamiento real. "
        "Cuando se complete, probablemente requiera crear Judge API v1.1 (ver Seccion 18) "
        "y reejecutar todos los experimentos afectados."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 22: META-REGLA
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("22. Meta-Regla", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(
        "No agregar una sola caracteristica nueva\nsi no aumenta la calidad de la evidencia."
    )
    run.font.size = Pt(16)
    run.font.color.rgb = c(P["red"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta regla, por si sola, probablemente mantenga al proyecto en el camino cientifico "
        "que el protocolo intenta establecer. El proyecto ya tiene suficientes modulos: "
        "Dataset Builder, Corruptor, Judge, Carving, Motor MFT, Motor C, RVS, FQS, Protocol, "
        "Confidence Registry, Evidence Gate, Evidence Ledger. Eso ya es muchisimo. El tablero "
        "ya no mide lineas de codigo, features o modulos. Mide unicamente la calidad de la "
        "evidencia. Y la calidad de la evidencia se mide con el RCR (Seccion 20). Si el RCR "
        "no sube, el proyecto no avanza. Punto."
    )
    run.font.size = Pt(11)

    # ─── Save ──────────────────────────────────────────────────────────────
    output_path = "/home/z/my-project/download/Research_Protocol_v1.4.docx"
    doc.save(output_path)
    print(f"Research Protocol v1.4 saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build_research_protocol()
    print(f"Generated: {path}")
