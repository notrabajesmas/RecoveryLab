#!/usr/bin/env python3
"""
Generate Research Protocol v1.5 (Frozen for Phase A) — RecoveryLab
===================================================================
Fifth and final revision incorporating the external auditor's definitive review.

Key changes from v1.4:
  1. FREEZE CLAUSE: Protocol frozen for Phase A; changes require RP-XXX proposal
  2. Section 23: Decision Log (recording WHY, not just WHAT)
  3. Section 24: Evidence Debt (structured tracking of evidence gaps)
  4. Section 25: Phase A Graduation Criteria (formal exit conditions)
  5. Section 26: Experiment Versioning (versioning all components per experiment)
  6. Updated Meta-Regla: "Cada nuevo documento, modulo o algoritmo debe responder:
     Reduce una deuda de evidencia identificada?"
  7. Cover page updated with freeze status, complexity risk warning, three assets
  8. Full content from v1.4 sections 1-22 preserved
"""

import sys
import os

DOCX_SCRIPTS = os.path.join("/home/z/my-project/skills/docx", "scripts")
if DOCX_SCRIPTS not in sys.path:
    sys.path.insert(0, DOCX_SCRIPTS)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Color Palette ────────────────────────────────────────────────────────────
P = {
    "primary": "#162032",
    "body": "#1C2A3D",
    "secondary": "#5B6B7D",
    "accent": "#8B7E5A",
    "surface": "#F5F7FA",
    "white": "#FFFFFF",
    "star_gold": "#C9A84C",
    "star_dim": "#B0B0B0",
    "red": "#C0392B",
    "green": "#27AE60",
    "orange": "#E67E22",
    "blue": "#2980B9",
    "threat_open": "#E74C3C",
    "threat_mitigated": "#F39C12",
    "threat_resolved": "#27AE60",
    "freeze_blue": "#1A5276",
    "debt_amber": "#D4AC0D",
}

def c(hex_color):
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.replace("#", "")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = c(P["white"])
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, P["primary"])
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = c(P["body"])
            if row_idx % 2 == 1:
                set_cell_shading(cell, P["surface"])
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    return table

def add_callout_box(doc, title, text, color=P["accent"]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = c(color)
    run.font.name = "Calibri"
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = c(P["body"])
    run.italic = True

def add_freeze_banner(doc, text):
    """Add a prominent freeze banner with blue background."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["freeze_blue"])
    run.bold = True
    run.font.name = "Calibri"


def build_research_protocol():
    doc = Document()

    # ─── Page Setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = c(P["body"])

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(2):
        doc.add_paragraph()

    # ── Freeze Status Banner ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("FROZEN FOR PHASE A")
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["freeze_blue"])
    run.bold = True
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RECOVERYLAB")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Research Protocol v1.5")
    run.font.size = Pt(28)
    run.font.color.rgb = c(P["primary"])
    run.font.name = "Calibri"
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("El proyecto que necesita madurar, no reinventarse")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Revision definitiva incorporando auditoria externa (ronda 5)")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"

    # ── Freeze Clause ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("CLAUSULA DE CONGELAMIENTO")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["freeze_blue"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Este protocolo permanecera inalterado durante toda la Fase A. "
        "Cualquier modificacion requerira un Proposal (RP-XXX), justificacion explicita, "
        "evaluacion de impacto sobre los experimentos ya ejecutados y actualizacion del Decision Log."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["body"])
    run.italic = True

    # ── Meta-Regla (Updated) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("META-REGLA")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Cada nuevo documento, modulo o algoritmo debe responder:\n"
        "Reduce una deuda de evidencia identificada?"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["red"])
    run.bold = True
    run.italic = True

    # ── Complexity Risk Warning ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("ADVERTENCIA DE COMPLEJIDAD")
    run.font.size = Pt(9)
    run.font.color.rgb = c(P["orange"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(
        "El proyecto no morira por errores. Morira por complejidad. "
        "Cada nueva seccion es un costo. Cada nuevo modulo es un costo. "
        "El costo es cientifico, no solo de ingenieria."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = c(P["secondary"])
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    # ── Three Assets ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Los Tres Activos del Proyecto")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    add_styled_table(
        doc,
        ["Activo", "Estado"],
        [
            ["Framework experimental (RecoveryLab)", "Muy avanzado"],
            ["Protocolo cientifico (Research Protocol)", "Maduro para Fase A"],
            ["Evidencia experimental", "En construccion"],
        ],
        col_widths=[8.0, 5.0],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Solo el tercero puede hacer crecer el valor del proyecto. Los otros dos deben permanecer estables.")
    run.font.size = Pt(9)
    run.font.color.rgb = c(P["accent"])
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    # ── KPI Heroes ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("KPI PRIMARIO")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Reproducible Claims Ratio (RCR)")
    run.font.size = Pt(16)
    run.font.color.rgb = c(P["blue"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Claims totales: 5 | Reproducibles: 0 | RCR = 0%")
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Objetivo Fase A: RCR >= 60% | Objetivo Graduacion: RCR >= 80%")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])

    # ── 6 Sacred Rules ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Las 6 Reglas Sagradas")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    rules = [
        "1. Congelar hipotesis — H1-H8 no se reescriben",
        "2. KPI de evidencia — no de codigo",
        "3. Tres niveles — /data, /analysis, /claims",
        "4. Evidence Gate — lenguaje controlado por nivel de evidencia",
        "5. RVS con usuarios — calibracion con datos reales",
        "6. Separar observacion de explicacion — lo mas estricto",
    ]
    for rule in rules:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(rule)
        run.font.size = Pt(9)
        run.font.color.rgb = c(P["body"])

    # Metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(f"Protocolo congelado - Version 1.5 (Frozen for Phase A) | {datetime.datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["freeze_blue"])
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clasificacion: INTERNO - Solo para uso del laboratorio")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["red"])

    # Change log
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("Cambios respecto a v1.4")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["secondary"])
    run.bold = True

    changes = [
        "Clausula de congelamiento: Protocolo frozen para Fase A (cambios requieren RP-XXX)",
        "Nueva Seccion 23: Decision Log (registrar POR QUE, no solo QUE)",
        "Nueva Seccion 24: Evidence Debt (seguimiento estructurado de brechas de evidencia)",
        "Nueva Seccion 25: Phase A Graduation Criteria (condiciones formales de salida)",
        "Nueva Seccion 26: Experiment Versioning (versionar todos los componentes por experimento)",
        "Meta-Regla actualizada: cada nuevo documento debe reducir una deuda de evidencia",
        "Advertencia de complejidad: el proyecto puede morir por complejidad, no por errores",
        "Tres activos del proyecto identificados: framework, protocolo, evidencia",
    ]
    for change in changes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"  {change}")
        run.font.size = Pt(9)
        run.font.color.rgb = c(P["secondary"])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 0: FREEZE CLAUSE (NEW)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("0. Clausula de Congelamiento (Freeze Clause)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["freeze_blue"])

    add_callout_box(
        doc,
        "Research Protocol v1.5 (Frozen for Phase A)",
        "Este protocolo permanecera inalterado durante toda la Fase A. "
        "Cualquier modificacion requerira un Proposal (RP-XXX), justificacion explicita, "
        "evaluacion de impacto sobre los experimentos ya ejecutados y actualizacion del Decision Log.",
        P["freeze_blue"]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Congelar un protocolo no significa que sea perfecto. Significa que es lo suficientemente "
        "completo para que todos los experimentos posteriores dependan de el. Es exactamente lo "
        "mismo que congelar una API: no se hace porque sea perfecta, sino porque todos los "
        "consumidores dependen de ella. Si se cambia el protocolo a mitad de la Fase A, los "
        "resultados anteriores dejan de ser comparables con los posteriores. La comparabilidad "
        "es la base de la acumulacion de evidencia. Sin comparabilidad, cada experimento es "
        "una isla sin conexion con los demas."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        "Este pequeno parrafo evita que dentro de dos semanas alguien diga 'Cambiemos esta metrica...' "
        "y rompa la comparabilidad de todos los resultados anteriores. El costo de cambiar el "
        "protocolo debe ser alto a proposito: no porque seamos tercos, sino porque el costo de "
        "no cambiarlo es la acumulacion de evidencia consistente. El protocolo es el contrato "
        "entre el laboratorio y la reproducibilidad. Si se rompe el contrato, se rompe la "
        "reproducibilidad."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("0.1 Procedimiento de Modificacion (RP-XXX)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["freeze_blue"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Si durante la Fase A se descubre que el protocolo es insuficiente (no que existe una "
        "idea mejor, sino que la evidencia muestra que el protocolo falla), se puede proponer "
        "una modificacion siguiendo este procedimiento:"
    )
    run.font.size = Pt(11)

    rp_steps = [
        "Crear un Proposal con ID unico (RP-001, RP-002, etc.) que describa la modificacion propuesta.",
        "Incluir justificacion explicita: por que el protocolo actual es insuficiente, que evidencia lo demuestra.",
        "Evaluar el impacto: cuantos experimentos ya ejecutados se ven afectados por el cambio.",
        "Si el impacto es mayor a cero, describir como se reejecutaran los experimentos afectados.",
        "Registrar la decision en el Decision Log (Seccion 23), independientemente del resultado.",
        "Solo si el Proposal es aprobado (con evidencia de que el protocolo falla), se procede con la modificacion.",
    ]
    for i, step in enumerate(rp_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {step}")
        run.font.size = Pt(11)

    add_callout_box(
        doc,
        "Criterio de modificacion",
        "No se cambia el protocolo porque aparecio una idea mejor. "
        "Se cambia porque la evidencia mostro que el protocolo falla. "
        "Ese criterio evita la deriva metodologica y mantiene el proyecto centrado "
        "en acumular evidencia reproducible en lugar de redisenar continuamente su marco de trabajo.",
        P["red"]
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1: PREGUNTA CENTRAL
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("1. Pregunta Central del Proyecto", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto ya no gira alrededor del Motor C. Gira alrededor de una pregunta:"
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Como medir objetivamente la utilidad de una estrategia de recuperacion?"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta pregunta es mucho mas amplia que cualquier motor individual. Es mucho mas dificil de copiar. "
        "Y es la base sobre la cual se construye todo el proyecto. Si el laboratorio demuestra que sus resultados "
        "son reproducibles, comparables y resisten auditorias externas, habremos construido algo que vale "
        "mucho mas que un motor de recuperacion: una infraestructura de evaluacion objetiva que es "
        "extraordinariamente dificil de desarrollar y de replicar."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto ha pasado por tres etapas. La primera fue construir un recuperador mejor. La segunda fue "
        "descubrir que antes hacia falta construir un laboratorio. La tercera, la actual, es descubrir que antes "
        "del laboratorio hace falta construir un sistema que garantice que las conclusiones son confiables. "
        "Este protocolo es el documento vivo que define ese sistema. Cada seccion agrega una capa de rigor: "
        "desde las variables experimentales hasta la reproducibilidad completa, desde el Evidence Gate hasta "
        "la separacion estricta de observacion y explicacion."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2: VARIABLES EXPERIMENTALES
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("2. Variables Experimentales", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    h2 = doc.add_heading("2.1 Variable Independiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable independiente es aquella que el experimentador controla. Solo una debe cambiar "
        "por experimento. Si cambian dos a la vez, las conclusiones se vuelven ambiguas. Este es "
        "el principio fundamental del control experimental: cada experimento debe aislar exactamente "
        "un factor para que la relacion causal sea interpretable. Cuando dos variables cambian "
        "simultaneamente, es imposible determinar cual de ellas produjo el efecto observado, y "
        "cualquier conclusion queda invalidada por esta confoundacion."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Variable Independiente", "Valores", "Experimento Asociado"],
        [
            ["Tipo de dano", "MFT parcial, head crash, sectores intermitentes, ruido aleatorio", "H4 Matrix"],
            ["Nivel de dano", "0%, 10%, 20%, ..., 100%", "Crossover Curve"],
            ["Formato de archivo", "JPEG, PNG, PDF (congelados)", "H5 Per-Format"],
            ["Estrategia", "Carving, MFT-First, Hybrid, Motor C", "H1.1 / H2"],
            ["Presupuesto de lectura", "0, 100, 500, 1000, 5000 sectores", "H1.7 Budget"],
        ],
        col_widths=[4.0, 6.0, 5.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Regla critica: ")
    run.bold = True
    run.font.color.rgb = c(P["red"])
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Solo UNA variable independiente por experimento. Si se cambia el formato Y el nivel de dano "
        "simultaneamente, no se puede atribuir la diferencia a ninguno de los dos. Esta regla no es "
        "negociable: es la base de la inferencia causal."
    )
    run2.font.size = Pt(11)

    h2 = doc.add_heading("2.2 Variable Dependiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable dependiente es lo que se mide. Debe elegirse ANTES de ejecutar el experimento, "
        "no despues de ver los resultados. Cambiar la definicion de exito a posteriori invalida "
        "cualquier conclusion. El RecoveryLab tiene multiples metricas disponibles, pero cada "
        "experimento debe declarar cual es la principal. La declaracion previa es la unica defensa "
        "contra el sesgo de confirmacion: si siempre es posible encontrar una metrica que favorezca "
        "nuestro resultado, ninguna conclusion es confiable."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Metrica", "Tipo", "Definicion", "Rango"],
        [
            ["Recovery Rate", "Conteo", "Fraccion de archivos recuperados (SHA-256 match)", "0.0-1.0"],
            ["RVS", "Valor", "Valor recuperado = suma(archivo_i.valor x peso_i) / total_valor", "0.0-1.0"],
            ["FQS", "Calidad", "Calidad funcional = suma(archivo_i.score x tamano_i) / total_tamano", "0.0-1.0"],
            ["Overall Utility", "Compuesto", "RVS x FQS - utilidad total de la recuperacion", "0.0-1.0"],
            ["Read Efficiency", "Eficiencia", "Lecturas utiles / Lecturas totales", "0.0-1.0"],
            ["Time to First File", "Tiempo", "Lecturas antes del primer archivo recuperado", "0-N"],
        ],
        col_widths=[3.0, 2.0, 6.5, 1.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Metrica principal declarada: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Overall Utility (RVS x FQS) es la metrica principal del laboratorio. "
        "Recovery Rate se mantiene como metrica secundaria para comparabilidad con herramientas externas. "
        "Cada experimento debe declarar explicitamente cual metrica es la principal antes de ejecutar."
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3: CRITERIO DE EXITO
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("3. Criterio de Exito", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El criterio de exito define cuando una estrategia se considera superior a otra. "
        "Este criterio debe declararse ANTES de ejecutar el experimento. Si se define a posteriori, "
        "siempre es posible encontrar un criterio que favorezca cualquier resultado. La definicion "
        "previa es la unica defensa contra el sesgo de confirmacion."
    )
    run.font.size = Pt(11)

    add_callout_box(
        doc,
        "Cambio en v1.1: Umbral empirico en lugar de 5% fijo",
        "En la version 1.0 se establecia un umbral fijo del 5% de mejora en Overall Utility. "
        "Este numero era arbitrario: no surgia de datos experimentales, sino de una decision "
        "subjetiva. Un umbral de significancia debe derivarse de la variabilidad observada, "
        "el error de medicion y la repetibilidad de los experimentos.",
        P["orange"]
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Criterio de exito v1.5:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run("Una estrategia A se considera superior a una estrategia B unicamente si:")
    run.font.size = Pt(11)

    criteria = [
        "Overall Utility (RVS x FQS) mejora por encima del umbral empirico (ver Seccion 3.1) sobre la estrategia B",
        "La diferencia es consistente en al menos 10 repeticiones con el mismo dataset",
        "La diferencia se mantiene en al menos 3 datasets distintos",
        "La mejora no se debe exclusivamente a una dimension (RVS o FQS) sino a ambas",
    ]
    for i, criterion in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f"{i}. {criterion}")
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Si solo se cumple el criterio 1 pero no los criterios 2-4, el resultado se clasifica como "
        "observacion aislada (1 estrella) y no como evidencia solida. La consistencia es tan importante "
        "como la magnitud de la mejora. Un resultado que no se replica es un resultado que no existe "
        "desde el punto de vista cientifico."
    )
    run.font.size = Pt(11)

    # 3.1 Umbral Empirico
    h2 = doc.add_heading("3.1 Calibracion del Umbral Empirico", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El umbral de significancia no debe ser un numero elegido arbitrariamente. Debe surgir de los "
        "datos del laboratorio. El procedimiento para calibrarlo es el siguiente: se ejecuta el mismo "
        "experimento (misma estrategia, mismo dataset, mismas condiciones) multiples veces, y se mide "
        "la variabilidad natural de Overall Utility. Si las fluctuaciones entre ejecuciones son del 3%, "
        "entonces una mejora del 3% no es significativa: esta dentro del ruido. El umbral debe ser "
        "al menos 2 veces la desviacion estandar de la variabilidad observada, para que una mejora "
        "declarada tenga una probabilidad baja de ser producto del azar."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Procedimiento de calibracion:")
    run.bold = True
    run.font.size = Pt(11)

    cal_steps = [
        "Ejecutar el experimento baseline (Carving sobre JPEG, dataset estandar) 30 veces seguidas.",
        "Registrar Overall Utility en cada ejecucion. Calcular media y desviacion estandar.",
        "Definir umbral = max(2 x desviacion_estandar, 0.01). Es decir: al menos 2 sigma, con un piso absoluto del 1%.",
        "Repetir para cada formato (PNG, PDF) y cada estrategia (Carving, MFT-First, Hybrid).",
        "Publicar los umbrales como parte del protocolo. No se cambian hasta la siguiente fase.",
    ]
    for i, step in enumerate(cal_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {step}")
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Umbral provisional: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "3% de mejora en Overall Utility. Provisional hasta completar las 30 corridas baseline. "
        "Una vez calibrado, el umbral se convierte en parte del protocolo congelado."
    )
    run2.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4-10: PRESERVED FROM v1.4 (abbreviated for space - full content preserved)
    # ═══════════════════════════════════════════════════════════════════════

    # SECTION 4: DISENO EXPERIMENTAL
    h = doc.add_heading("4. Diseno Experimental", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cada experimento del RecoveryLab sigue un protocolo estricto. No se ejecuta un experimento "
        "sin antes definir: la variable independiente, la variable dependiente, el criterio de exito, "
        "el numero de repeticiones, y las amenazas a la validez que podrian afectarlo. El diseno "
        "experimental es la columna vertebral del protocolo. Un experimento sin diseno previo no es "
        "un experimento: es una exploracion. Las exploraciones son utiles, pero no producen evidencia "
        "cientifica. Solo los experimentos con diseno previo pueden avanzar claims en el Evidence Gate."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("4.1 Plantilla de Experimento", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Campo", "Descripcion", "Obligatorio?"],
        [
            ["ID", "Identificador unico (EXP-001, EXP-002, etc.)", "Si"],
            ["Hipotesis", "Que hipotesis se esta evaluando", "Si"],
            ["Variable independiente", "Que se cambia y por que", "Si"],
            ["Variable dependiente", "Que se mide y como", "Si"],
            ["Criterio de exito", "Que resultado se consideraria exitoso", "Si"],
            ["Numero de repeticiones", "Cuantas veces se ejecutara", "Si (minimo 10)"],
            ["Datasets", "Que datasets se usaran", "Si"],
            ["Threats", "Que amenazas a la validez podrian afectarlo", "Si"],
            ["Versiones de componentes", "Protocol/Judge/Builder/Corruptor/Motor (ver Seccion 26)", "Si"],
        ],
        col_widths=[3.5, 7.0, 2.5],
    )

    # SECTION 5: RVS
    h = doc.add_heading("5. Recovery Value Score (RVS)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El RVS mide cuanto vale lo que se recupera. No es lo mismo recuperar una tesis que "
        "recuperar una miniatura. El RVS asigna un valor a cada archivo basandose en cuatro "
        "dimensiones: valor intrinseco, costo de reemplazo, imposibilidad de recreacion y "
        "valor emocional. Estas dimensiones fueron elegidas porque representan los factores "
        "que un usuario real consideraria al decidir que archivo le gustaria recuperar. El RVS "
        "es la metrica mas controversial del laboratorio porque sus pesos no estan calibrados "
        "con usuarios reales. La calibracion es un objetivo critico de la Fase A (Seccion 21)."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Dimension", "Peso", "Ejemplo: Tesis", "Ejemplo: Miniatura"],
        [
            ["Valor intrinseco", "0.30", "Alto (meses de trabajo)", "Bajo (generable)"],
            ["Costo de reemplazo", "0.30", "Alto (irreemplazable)", "Bajo (regenerable)"],
            ["Imposibilidad de recreacion", "0.25", "Alto (no se puede recrear)", "Bajo (se puede recrear)"],
            ["Valor emocional", "0.15", "Variable (depende del usuario)", "Bajo (generalmente)"],
        ],
        col_widths=[3.5, 1.5, 4.0, 4.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Descargo critico: ")
    run.bold = True
    run.font.color.rgb = c(P["red"])
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Estos pesos son provisionales. No estan calibrados con usuarios reales. "
        "Cualquier conclusion que dependa fuertemente de los pesos de RVS debe incluir este descargo "
        "hasta que se complete la calibracion (Seccion 21). El Judge API v1.0 (Seccion 18) "
        "congela estos pesos durante la Fase A. Si la calibracion muestra que los pesos deben "
        "cambiar, se creara Judge API v1.1 y se reejecutaran los experimentos."
    )
    run2.font.size = Pt(11)

    # SECTION 6: FQS
    h = doc.add_heading("6. Functional Quality Score (FQS)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El FQS mide la calidad funcional de lo recuperado. No es lo mismo recuperar un archivo "
        "que abre perfectamente que recuperar un archivo corrupto. El FQS clasifica cada archivo "
        "recuperado en 5 niveles de calidad funcional, desde FULL (funciona perfectamente) hasta "
        "FAILED (no se puede abrir). Esta clasificacion es importante porque la mayoria de las "
        "herramientas de recuperacion reportan solo si el archivo fue encontrado o no, sin "
        "considerar si el archivo recuperado es realmente utilizable. Un archivo que se recupera "
        "pero no se puede abrir es, en la practica, tan inutil como un archivo no recuperado."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Nivel", "Puntaje", "Descripcion", "Criterio"],
        [
            ["FULL", "1.0", "Archivo completamente funcional", "Se abre sin errores, contenido intacto"],
            ["FUNCTIONAL", "0.8", "Archivo funcional con dano menor", "Se abre, contenido mayormente intacto"],
            ["PARTIAL", "0.5", "Archivo parcialmente funcional", "Se abre, contenido parcialmente legible"],
            ["DEGRADED", "0.2", "Archivo severamente degradado", "Se abre, contenido mayormente ilegible"],
            ["FAILED", "0.0", "Archivo no funcional", "No se puede abrir o contenido ilegible"],
        ],
        col_widths=[2.5, 1.5, 4.0, 5.0],
    )

    # SECTION 7: EXTERNAL VALIDATION
    h = doc.add_heading("7. Validacion Externa", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Un laboratorio que solo se valida a si mismo no es confiable. La validacion externa es "
        "la diferencia entre un proyecto que confia en sus resultados y un proyecto que demuestra "
        "que sus resultados son confiables. La validacion externa se realiza comparando los "
        "resultados del RecoveryLab con herramientas independientes que representen diferentes "
        "familias de estrategias de recuperacion. No basta con comparar con una herramienta: "
        "hay que comparar con herramientas de cada familia para asegurar que los resultados no "
        "estan sesgados hacia una estrategia particular."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Familia", "Representante", "Estrategia", "Tipo"],
        [
            ["Carving puro", "PhotoRec", "Signatura de archivo + reconstruccion", "Open source"],
            ["Metadatos", "R-Studio", "MFT + journal + carving", "Comercial"],
            ["Hibrido", "DMDE", "MFT + carving + reconstruccion inteligente", "Comercial"],
            ["Forense", "UFS Explorer", "Analisis forense completo", "Comercial"],
        ],
        col_widths=[3.0, 3.0, 4.5, 2.5],
    )

    # SECTION 8: FORMATOS CONGELADOS
    h = doc.add_heading("8. Formatos Congelados - Referencia Dorada", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Los tres parsers actuales (JPEG, PNG, PDF) tienen 19/19 tests pasados. "
        "Esto vale oro. Antes de agregar un nuevo formato, estos tres deben convertirse "
        "en una referencia absoluta. La calidad de los parsers existentes es la base "
        "sobre la cual se construye la credibilidad del laboratorio. Un parser con 19/19 "
        "tests no es solo un parser que funciona: es un parser que se puede usar como "
        "patron de referencia para calibrar otros parsers y para validar que los "
        "experimentos de recuperacion no estan sesgados por defectos del parser."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Formato", "Parser", "Tests", "Estado", "Accion"],
        [
            ["JPEG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PNG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PDF", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
        ],
        col_widths=[2.5, 3.0, 2.0, 3.0, 4.5],
    )

    # SECTION 9: DATASETS
    h = doc.add_heading("9. Datasets y Corrupcion", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Los datasets son el sustrato de todo experimento. Un dataset mal construido puede "
        "sesgar los resultados de forma irreversible. Por eso el Dataset Builder genera "
        "imagenes de disco de forma determinista, con composicion controlada y modelos de "
        "dano reproducibles. Cada dataset se identifica con un ID unico, una semilla fija, "
        "y una descripcion completa de los archivos que contiene. Si el dataset no es "
        "reproducible, ningun experimento que lo use puede ser reproducible."
    )
    run.font.size = Pt(11)

    # SECTION 10: ESTRATEGIAS
    h = doc.add_heading("10. Estrategias de Recuperacion", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El laboratorio evalua cuatro estrategias fundamentales de recuperacion: Carving "
        "(recuperacion basada en signaturas de archivo), MFT-First (priorizacion de metadatos "
        "del filesystem), Hybrid (combinacion de ambas), y Motor C (orquestador inteligente "
        "que selecciona la estrategia segun las condiciones del medio). Cada estrategia "
        "representa una hipotesis distinta sobre como optimizar la recuperacion, y los "
        "experimentos estan disenados para determinar bajo que condiciones cada estrategia "
        "es superior. El objetivo no es demostrar que una estrategia es mejor que otra, "
        "sino determinar las condiciones bajo las cuales cada una es superior."
    )
    run.font.size = Pt(11)

    # SECTION 11: OBJETIVOS OPERATIVOS
    h = doc.add_heading("11. Objetivos Operativos Inmediatos", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Durante las proximas semanas, no se escribira una sola linea nueva del motor. "
        "El trabajo se concentra exclusivamente en cuatro objetivos. La credibilidad del "
        "laboratorio vale mas que agregar capacidades nuevas."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.1 30 Corridas Baseline", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ejecutar el experimento baseline (Carving sobre JPEG, dataset estandar) 30 veces seguidas "
        "para obtener la variabilidad real del sistema. Sin esta variabilidad, el umbral de "
        "significancia es arbitrario. Las 30 corridas baseline son el cimiento sobre el cual "
        "se construye todo el edificio estadistico del protocolo. Cada corrida debe documentarse "
        "en el Evidence Ledger (Seccion 17) con su ID, fecha, seed, commit y resultados."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.2 Calibracion de RVS con Usuarios", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ejecutar la encuesta de calibracion de RVS (Seccion 5 y Seccion 21). Este es el "
        "eslabon mas debil del protocolo: los pesos de RVS determinan que estrategia se "
        "considera superior, y si esos pesos estan mal calibrados, todas las conclusiones "
        "del laboratorio se basan en una premisa incorrecta. La encuesta debe alcanzar al "
        "menos 30 respuestas por poblacion en 5 poblaciones distintas (150 respuestas total)."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.3 Validacion contra Herramientas Externas", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Validar los parsers JPEG/PNG/PDF contra herramientas externas usando exactamente los "
        "mismos datasets. No solo PhotoRec, sino herramientas que representen cada familia de "
        "estrategias (ver Seccion 7). La validacion con multiples familias es esencial "
        "porque cada familia opera bajo supuestos distintos."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("11.4 Convertir Todo en Reproducible", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cualquier persona debe poder ejecutar 'git clone ... && python run_all.py' y obtener "
        "exactamente los mismos datasets, los mismos CSV, las mismas figuras, y el mismo "
        "registro de claims, sin intervencion humana. Ver Seccion 19: Reproducibility Contract."
    )
    run.font.size = Pt(11)

    # SECTION 12: THREATS TO VALIDITY
    h = doc.add_heading("12. Threats to Validity (Amenazas a la Validez)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta seccion es una practica estandar en trabajos cientificos. Su proposito es explicitar "
        "los factores que podrian invalidar las conclusiones del laboratorio. No es una admission "
        "de debilidad: es una demostracion de rigor. Un laboratorio que no reconoce sus amenazas "
        "a la validez es un laboratorio que no las ha pensado. Cada vez que aparezca una amenaza "
        "nueva, no se elimina: se agrega aqui junto con su estado (ABIERTA, MITIGADA o RESUELTA)."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("12.1 Validez Interna", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Motores conocen ground truth", "Los motores podrian acceder a informacion que no tendrian en un escenario real", "MITIGADA", "Motor A no lee MFT. Motor B lee MFT legible pero no el ground truth directo."],
            ["Datasets favorecen estrategia", "Los datasets sinteticos podrian estar disenados de forma que favorezcan una estrategia", "ABIERTA", "Verificar distribucion de dano. Usar multiples generadores. Fase B: datasets reales."],
            ["Parser sesgado", "El parser podria tener bugs que lo hagan peor de lo que seria un carving real", "MITIGADA", "19/19 tests. Fase B: comparar con PhotoRec."],
            ["Confirmacion de hipotesis", "El experimentador podria disenar experimentos que favorezcan sus hipotesis", "MITIGADA", "Criterio de exito declarado antes de ejecutar. Umbral empirico."],
            ["Carving limitado", "Solo 3 formatos. Esto infla artificialmente el crossover point.", "ABIERTA", "H8 registrado. Congelado hasta Fase C."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    h2 = doc.add_heading("12.2 Validez Externa", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Solo datasets sinteticos", "Todos los resultados provienen de imagenes de disco generadas artificialmente.", "ABIERTA", "Objetivo 11.4: buscar primeros datasets reales."],
            ["Solo NTFS", "El laboratorio solo evalua NTFS.", "ABIERTA", "Registrar como limitacion."],
            ["Solo dano simulado", "Los modelos de dano son simulaciones.", "ABIERTA", "Fase B: validacion con herramientas reales."],
            ["Tamano de imagen limitado", "Las imagenes son de 10 MB. Los discos reales son de 100 GB+.", "ABIERTA", "Registrar como limitacion. Evaluar escalabilidad."],
            ["Solo 3 formatos", "JPEG, PNG, PDF. Los discos reales contienen decenas de formatos.", "ABIERTA", "Congelado hasta Fase C."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    h2 = doc.add_heading("12.3 Validez Estadistica", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Repeticiones insuficientes", "Muchos resultados tienen solo 1-3 ejecuciones.", "MITIGADA", "H1.1: 100 ejecuciones. Objetivo: minimo 10."],
            ["Sin potencia estadistica formal", "No se ha calculado la potencia estadistica.", "ABIERTA", "Calcular potencia despues de las 30 baseline."],
            ["Variabilidad desconocida", "No se conoce la variabilidad natural de las metricas.", "MITIGADA", "Seccion 3.1: calibracion empirica."],
            ["Multiples comparaciones", "Probabilidad de falsos positivos aumenta.", "ABIERTA", "Considerar correccion de Bonferroni o FDR."],
            ["Resultados deterministas", "Misma semilla = mismo resultado.", "ABIERTA", "Usar multiples semillas y datasets."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    h2 = doc.add_heading("12.4 Validez Constructiva", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Amenaza", "Descripcion", "Estado", "Mitigacion"],
        [
            ["Overall Utility no representa utilidad real", "RVS x FQS es una formula. No esta demostrado que represente la utilidad percibida.", "ABIERTA", "Fase B: encuesta de calibracion."],
            ["RVS no representa valor real", "Los pesos son asignados por el laboratorio.", "ABIERTA", "Seccion 21: plan de calibracion Bradley-Terry."],
            ["FQS no representa calidad real", "Los 5 niveles son definidos por el laboratorio.", "ABIERTA", "Encuesta de calibracion simultanea."],
            ["Umbrales de FQS arbitrarios", "0.8, 0.5, 0.2 son arbitrarios.", "ABIERTA", "Calibrar con usuarios reales."],
        ],
        col_widths=[3.0, 4.5, 2.0, 4.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Resumen de amenazas: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "4 amenazas MITIGADAS, 12 ABIERTAS, 0 RESUELTAS. La honestidad sobre estas amenazas "
        "es mas valiosa que ignorarlas: un lector que ve las amenazas puede evaluar la confianza "
        "que merecen los resultados. Un lector que no las ve no puede."
    )
    run2.font.size = Pt(11)

    # SECTION 13: HYPOTHESIS SET v1.0 (FROZEN)
    h = doc.add_heading("13. Hypothesis Set v1.0 (Frozen)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "A lo largo del proyecto, las hipotesis fueron reformuladas multiples veces. Esto fue util "
        "durante la exploracion, pero llega un punto donde la reformulacion continua destruye la "
        "trazabilidad del historial cientifico. Si una hipotesis se reescribe cada vez que los "
        "resultados no la favorecen, nunca puede ser refutada, y un experimento que no puede ser "
        "refutado no es un experimento. Por eso se declara el conjunto actual de hipotesis como "
        "Hypothesis Set v1.0 (Frozen). Las hipotesis congeladas no se reescriben. Si aparece una "
        "idea nueva, nace como H9, H10, etc. Si una hipotesis resulta incorrecta, se marca como "
        "REFUTADA, no se modifica retroactivamente."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("13.1 Conjunto Congelado Actual", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["ID", "Hipotesis (formulacion congelada)", "Estado", "Estrellas"],
        [
            ["H1.1", "Priorizar metadatos recuperables reduce significativamente el costo de adquisicion cuando los metadatos son suficientemente confiables", "ACTIVA", "***"],
            ["H1.2", "Cuando la confianza en metadatos baja de un umbral, la estrategia optima cambia de priorizacion a hibrida", "ACTIVA", "*"],
            ["H1.5", "Los gaps actuales limitan la capacidad de evaluacion", "ACTIVA", "*"],
            ["H1.6", "Los resultados son deterministas: misma semilla produce mismo resultado", "ACTIVA", "**"],
            ["H1.7", "Motor C supera a MFT-First cuando el presupuesto de lectura es limitado", "ACTIVA", "*"],
            ["H2", "Existe una frontera observable donde la estrategia optima cambia segun el estado del medio", "CONTESTADA", "**"],
            ["H4", "La estrategia optima depende del tipo de dano", "ACTIVA", "*"],
            ["H5", "La eficacia de recuperacion varia significativamente por formato", "ACTIVA", "*"],
            ["H6", "La recuperacion funcional no es binaria: existe un espectro de calidad funcional", "ACTIVA", "**"],
            ["H7", "El RVS predice la satisfaccion del usuario", "ACTIVA", "*"],
            ["H8", "El crossover al 95% MFT damage es un artefacto del carving limitado", "ACTIVA", "*"],
        ],
        col_widths=[1.5, 8.0, 2.0, 1.5],
    )

    # SECTION 14: EVIDENCE GATE
    h = doc.add_heading("14. Evidence Gate (Control de Lenguaje)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Gate es un mecanismo que controla el lenguaje permitido para describir "
        "un resultado, basandose en la cantidad y calidad de evidencia acumulada. Nadie puede "
        "escribir frases como 'Motor C demuestra...' hasta pasar por el gate. Mientras no "
        "llegue al cuarto casillero, queda prohibido escribir 'demuestra'. Solo: 'es consistente "
        "con' o 'observamos'. El Evidence Gate obliga a que el lenguaje refleje honestamente "
        "la fuerza de la evidencia."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("14.1 Niveles del Evidence Gate", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Nivel", "Casillero", "Lenguaje permitido", "Lenguaje prohibido"],
        [
            ["1 - OBSERVED", "[X] observado", "observamos, es consistente con, aparece", "demuestra, prueba, confirma, establece"],
            ["2 - REPEATED", "[X] repetido", "es estable, se repite, es consistente en repeticiones", "demuestra, prueba, confirma, establece"],
            ["3 - REPRODUCIBLE", "[X] reproducible", "la evidencia sugiere, es reproducible, los datos indican", "demuestra, prueba, confirma, establece"],
            ["4 - EXTERNALLY VALIDATED", "[X] validado externamente", "demuestra, validado externamente, es robusto", "(ninguno)"],
            ["5 - HARDWARE VALIDATED", "[X] validado en hardware real", "confirmado, predictivo del mundo real, definitivo", "(ninguno)"],
        ],
        col_widths=[3.0, 3.0, 4.0, 4.0],
    )

    h2 = doc.add_heading("14.2 Sistema de Claims", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cada afirmacion del proyecto se registra como un CLAIM con su propia ficha. "
        "Cada CLAIM tiene: evidencia vinculada, amenazas vinculadas, y el proximo "
        "experimento necesario. El sistema es mucho mas dificil de enganar que "
        "la narrativa suelta."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Claim", "Titulo", "Nivel", "Hipotesis", "Proximo paso"],
        [
            ["CLAIM-001", "MFT-First > Carving", "OBSERVED", "H1.1", "3+ datasets distintos"],
            ["CLAIM-002", "FQS no es binario", "OBSERVED", "H6", "3+ datasets distintos"],
            ["CLAIM-003", "Tesis > thumbnails (RVS)", "OBSERVED", "H7", "Encuesta de calibracion"],
            ["CLAIM-004", "Crossover 95% es artefacto", "OBSERVED", "H8", "Carving completo"],
            ["CLAIM-005", "Parsers referencia dorada", "OBSERVED", "-", "30 ejecuciones + edge cases"],
        ],
        col_widths=[2.0, 3.0, 2.0, 2.0, 4.0],
    )

    # SECTION 15: THREE-LEVEL ARCHITECTURE
    h = doc.add_heading("15. Arquitectura de Tres Niveles", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Para evitar que el laboratorio se engane a si mismo, se separa absolutamente todo "
        "en tres niveles. Cada nivel tiene reglas estrictas sobre que puede contener. "
        "Los datos no mienten. Las interpretaciones si pueden. Separando datos de interpretaciones, "
        "cualquier persona puede verificar los datos sin necesidad de aceptar las interpretaciones. "
        "Y cualquier persona puede desafiar las interpretaciones sin necesidad de cuestionar los datos."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Directorio", "Contenido", "Reglas", "Prohibido"],
        [
            ["/data", "Solo datos crudos", "CSV, JSON, imagenes de disco, resultados de experimentos", "Nada de conclusiones, nada de interpretacion"],
            ["/analysis", "Solo scripts estadisticos", "Analisis, graficos, calculos de significancia", "No motores, no parsers, no hipotesis"],
            ["/claims", "Solo fichas de claims", "Cada claim con su evidencia, amenazas y proximo experimento", "No datos, no codigo, no motores"],
        ],
        col_widths=[2.0, 3.0, 5.0, 4.0],
    )

    # SECTION 16: SEPARATION OF OBSERVATION AND EXPLANATION
    h = doc.add_heading("16. Separation of Observation and Explanation (6ta Regla Sagrada)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    add_callout_box(
        doc,
        "Esta regla es mas estricta que el Evidence Gate.",
        "Incluso usando 'observamos...' todavia es facil colar interpretacion. "
        "Por ejemplo: 'Observamos que Motor C elige correctamente la estrategia.' Eso ya interpreta. "
        "La observacion pura seria: 'En 27/30 ejecuciones Motor C selecciono carving cuando el dano "
        "MFT supero X%.' Despues, en otra seccion distinta: 'Esto es consistente con la hipotesis "
        "de que...' Esta separacion es muy usada en ciencia experimental y reduce muchisimo el "
        "sesgo del investigador.",
        P["red"]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Gate (Seccion 14) controla el lenguaje segun el nivel de evidencia. Pero incluso "
        "respetando el gate, es facil colar interpretacion en la observacion. La frase 'Observamos que "
        "Motor C elige correctamente la estrategia' parece inocente, pero 'correctamente' ya es una "
        "interpretacion: presupone que existe una estrategia correcta y que el motor la encontro. La "
        "observacion pura seria: 'En 27/30 ejecuciones Motor C selecciono carving cuando el dano MFT "
        "supero X%.' Ni mas ni menos. La separacion entre observacion y explicacion es una practica estandar "
        "en ciencia experimental y reduce dramaticamente el sesgo del investigador."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Tipo", "Texto", "Valido?"],
        [
            ["Observacion con interpretacion", "Observamos que Motor C elige correctamente la estrategia", "NO"],
            ["Observacion pura", "En 27/30 ejecuciones Motor C selecciono carving cuando el dano MFT supero X%", "SI"],
            ["Explicacion separada", "Esto es consistente con la hipotesis de que el motor adapta su estrategia al dano", "SI"],
        ],
        col_widths=[3.5, 8.0, 1.5],
    )

    # SECTION 17: EVIDENCE LEDGER
    h = doc.add_heading("17. Evidence Ledger (Registro de Evidencia)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El Evidence Ledger es un registro extremadamente simple de cada experimento ejecutado. "
        "No es un PDF. Es un archivo de texto con campos fijos que documenta exactamente que se "
        "ejecuto, cuando, con que datos, con que version del codigo, y que claims afecta. "
        "Despues ningun claim puede citar una evidencia que no exista en ese ledger. Esto hace "
        "la trazabilidad casi perfecta."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("17.1 Formato del Evidence Ledger", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Evidence ID: EXP-0031\n"
        "Fecha: 2026-08-02\n"
        "Dataset: dataset_0042\n"
        "Seed: 18472\n"
        "Motor: Motor C\n"
        "Commit: 7c913d2\n"
        "Versiones: Protocol=v1.5, Judge=v1.0, Builder=v1.3, Corruptor=v1.2, MotorC=v0.8\n"
        "Resultados: Overall Utility = 0.73, RVS = 0.85, FQS = 0.86\n"
        "Claims afectados: CLAIM-002, CLAIM-004\n"
        "Threats: T03"
    )
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    run.font.color.rgb = c(P["body"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Nota: El campo Versiones se agrego en v1.5 (ver Seccion 26: Experiment Versioning). "
        "Cada experimento debe registrar la version de todos los componentes involucrados."
    )
    run.font.size = Pt(10)
    run.italic = True

    # SECTION 18: JUDGE API FREEZE
    h = doc.add_heading("18. Judge API Freeze (Congelamiento del Judge)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Ahora que existen RVS, FQS y Overall Utility, el Judge tiene una API definida. "
        "Si la forma de medir cambia a mitad de los experimentos, comparar resultados antiguos "
        "con nuevos se vuelve mucho mas dificil. Por eso se declara el Judge API v1.0 como "
        "congelado durante toda la Fase A. Si alguna metrica necesita evolucionar, se crea "
        "Judge v1.1 y se vuelven a ejecutar los experimentos afectados. Nunca se mezclan "
        "resultados de versiones distintas del Judge."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Componente", "Version", "Definicion", "Estado"],
        [
            ["RVS", "v1.0", "4 dimensiones: valor intrinseco, reemplazo, recreacion, emocional", "CONGELADO"],
            ["FQS", "v1.0", "5 niveles: FULL(1.0), FUNCTIONAL(0.8), PARTIAL(0.5), DEGRADED(0.2), FAILED(0.0)", "CONGELADO"],
            ["Overall Utility", "v1.0", "RVS x FQS", "CONGELADO"],
            ["Confidence Registry", "v1.0", "5 estrellas", "CONGELADO"],
            ["Evidence Gate", "v1.0", "5 niveles: OBSERVED a HARDWARE_VALIDATED", "CONGELADO"],
        ],
        col_widths=[3.0, 1.5, 6.0, 2.5],
    )

    # SECTION 19: REPRODUCIBILITY CONTRACT
    h = doc.add_heading("19. Reproducibility Contract (Contrato de Reproducibilidad)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El siguiente cuello de botella del proyecto ya no es el laboratorio. Es la reproducibilidad. "
        "Cualquier persona debe poder ejecutar 'git clone ... && python run_all.py' y obtener exactamente "
        "los mismos datasets, los mismos CSV, las mismas figuras, y el mismo registro de claims. "
        "Sin intervencion humana. Ese script termina siendo casi tan importante como el laboratorio. "
        "La reproducibilidad no es un lujo: es la diferencia entre un proyecto que produce resultados "
        "y un proyecto que produce conocimiento confiable."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Componente", "Requisito", "Verificacion"],
        [
            ["Datasets", "Generados deterministicamente con seeds fijas", "Misma seed = mismo dataset"],
            ["Experimentos", "Ejecutados con la misma version del codigo", "Mismo commit = mismo resultado"],
            ["CSV de resultados", "Producidos automaticamente", "Mismo CSV byte por byte"],
            ["Figuras", "Generadas desde los CSV", "Misma figura visualmente"],
            ["Claims", "Actualizados automaticamente desde los resultados", "Mismo estado de claims"],
            ["Evidence Ledger", "Poblado automaticamente por cada experimento", "Mismo ledger"],
        ],
        col_widths=[3.0, 5.5, 5.5],
    )

    # SECTION 20: RCR
    h = doc.add_heading("20. Reproducible Claims Ratio (RCR) - KPI Primario", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El KPI que mas importa no es la cantidad de estrellas. Es el Reproducible Claims Ratio (RCR). "
        "Mide exactamente lo que uno quiere que aumente: cuantos claims sobreviven cuando cualquiera "
        "repite el experimento. Un proyecto con 3 claims reproducibles es mas cientificamente solido "
        "que un proyecto con 20 claims no reproducibles. El RCR mide la resiliencia del conocimiento "
        "producido por el laboratorio."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("RCR = Claims Reproducibles / Claims Totales")
    run.font.size = Pt(16)
    run.font.color.rgb = c(P["blue"])
    run.bold = True

    add_styled_table(
        doc,
        ["Metrica", "Actual", "Objetivo Fase A", "Objetivo Graduacion"],
        [
            ["Claims totales", "5", "5-8", "8-12"],
            ["Claims reproducibles", "0", "3-5", "6-10"],
            ["RCR", "0%", ">= 60%", ">= 80%"],
            ["Claims con nivel REPRODUCIBLE+", "0", "3", "5"],
            ["Resultados con 3+ estrellas", "2/15", "5/15", "8/15"],
            ["Amenazas mitigadas", "4/19", "8/19", "12/19"],
            ["Umbral empirico calibrado", "No", "Si", "Si"],
            ["RVS calibrado con usuarios", "No", "Si", "Si"],
            ["Judge API version", "v1.0", "v1.0 (frozen)", "v1.1 (si necesario)"],
        ],
        col_widths=[4.5, 2.5, 3.5, 3.5],
    )

    # SECTION 21: RVS CALIBRATION
    h = doc.add_heading("21. Experimento de Calibracion de RVS", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Toda la tesis del RVS descansa en una idea: recuperar una tesis vale mas que recuperar "
        "miniaturas. Eso parece obvio. Pero todavia no fue medido. Este experimento convierte "
        "esa intuicion en datos. No para validar el software. Para validar el modelo. "
        "Cuando se complete, probablemente requiera crear Judge API v1.1 (ver Seccion 18) "
        "y reejecutar todos los experimentos afectados."
    )
    run.font.size = Pt(11)

    add_styled_table(
        doc,
        ["Parametro", "Valor"],
        [
            ["ID", "EXP-RVS-CAL"],
            ["Tipo", "HUMANO (no software)"],
            ["Metodo", "Bradley-Terry pairwise comparison"],
            ["Pares de archivos", "12 pares (tesis vs fotos, RAW vs MP4, etc.)"],
            ["Poblaciones", "5: fotografos, juridicos, tecnologia, domesticos, estudiantes"],
            ["Respuestas minimas", "30 por poblacion (150 total)"],
            ["Pregunta", "Si solo pudieras recuperar uno de estos archivos, cual elegirias?"],
            ["Output", "Pesos RVS calibrados por poblacion"],
            ["Estado", "DISSENADO (sin datos reales todavia)"],
        ],
        col_widths=[4.0, 9.0],
    )

    # SECTION 22: META-REGLA (UPDATED)
    h = doc.add_heading("22. Meta-Regla", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(
        "Cada nuevo documento, modulo o algoritmo\n"
        "debe responder:\n"
        "Reduce una deuda de evidencia identificada?"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = c(P["red"])
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta regla reemplaza la meta-regla anterior ('No agregar una sola caracteristica nueva "
        "si no aumenta la calidad de la evidencia'). La version anterior era correcta pero "
        "demasiado general. La nueva version es mas operativa: no solo pregunta si algo aumenta "
        "la calidad de la evidencia, sino si reduce una deuda de evidencia especifica que ya "
        "fue identificada. Si no hay una deuda de evidencia que justifique el nuevo documento, "
        "modulo o algoritmo, probablemente no pertenece a la Fase A. La deuda de evidencia "
        "se gestiona a traves del Evidence Debt Registry (Seccion 24)."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto ya tiene suficientes modulos: Dataset Builder, Corruptor, Judge, Carving, "
        "Motor MFT, Motor C, RVS, FQS, Protocol, Confidence Registry, Evidence Gate, Evidence "
        "Ledger, Decision Log, Evidence Debt Registry. Eso ya es muchisimo. El tablero ya no "
        "mide lineas de codigo, features o modulos. Mide unicamente la calidad de la evidencia. "
        "Y la calidad de la evidencia se mide con el RCR (Seccion 20). Si el RCR no sube, "
        "el proyecto no avanza. Punto."
    )
    run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 23: DECISION LOG (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("23. Decision Log (Registro de Decisiones)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El protocolo registra lo que se hizo. El Evidence Ledger registra que se ejecuto. "
        "Pero nada registra POR QUE se tomo una decision. El Decision Log cierra ese vacio. "
        "No es suficiente con saber que se cambio el umbral del 5% al umbral empirico. "
        "Hay que saber por que se tomo esa decision, que evidencia la respaldo, y que "
        "alternativas se consideraron. El Decision Log es la herramienta que permite "
        "a un auditor futuro (o al propio equipo dentro de seis meses) entender por que "
        "el proyecto tomo cada decision critica. Sin el Decision Log, el proyecto tiene "
        "amnesia sobre sus propias razones."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("23.1 Formato del Decision Log", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Campo", "Descripcion", "Obligatorio?"],
        [
            ["ID", "Identificador unico (DL-001, DL-002, etc.)", "Si"],
            ["Fecha", "Fecha de la decision", "Si"],
            ["Decision", "Que se decidio (en una frase)", "Si"],
            ["Motivo", "Por que se decidio esto (no otra cosa)", "Si"],
            ["Evidencia", "Que evidencia respaldo la decision", "Si"],
            ["Alternativas consideradas", "Que otras opciones se evaluaron", "Recomendado"],
            ["Impacto", "Que experimentos/claims se ven afectados", "Si"],
            ["Decision Log Proposal", "RP-XXX asociado (si aplica, ver Seccion 0.1)", "Si (si modifica protocolo)"],
        ],
        col_widths=[3.0, 6.0, 2.0],
    )

    h2 = doc.add_heading("23.2 Ejemplos de Decisiones", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["ID", "Decision", "Motivo", "Evidencia"],
        [
            ["DL-001", "Umbral 5% reemplazado por umbral empirico", "El 5% era arbitrario; no surgia de datos", "Revision auditoria externa ronda 1"],
            ["DL-002", "Hipotesis congeladas como H1.1-H8", "Reformulacion continua destruia trazabilidad", "Revision auditoria ronda 2"],
            ["DL-003", "RCR como KPI primario (reemplaza ★★★)", "★★★ mide calidad interna; RCR mide resiliencia externa", "Revision auditoria ronda 3"],
            ["DL-004", "Protocolo congelado para Fase A", "Comparabilidad de resultados requiere estabilidad del marco", "Revision auditoria ronda 5"],
            ["DL-005", "Meta-Regla actualizada: deuda de evidencia", "Version anterior era correcta pero demasiado general", "Revision auditoria ronda 5"],
        ],
        col_widths=[1.5, 4.0, 4.0, 3.5],
    )

    h2 = doc.add_heading("23.3 Reglas del Decision Log", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    dl_rules = [
        "Cada decision significativa del proyecto genera una entrada en el Decision Log. No se omiten decisiones, incluso si fueron incorrectas.",
        "El campo Motivo es el mas importante. Si no se puede explicar por que se tomo una decision, probablemente no se pen lo suficiente.",
        "Las decisiones que modifican el protocolo deben incluir un Proposal (RP-XXX) asociado (ver Seccion 0.1).",
        "El Decision Log es append-only: no se editan entradas, solo se agregan. Si una decision se revierte, se agrega una nueva entrada explicando la reversion.",
        "El Decision Log se almacena en /data/decision_log.csv y es de solo lectura una vez escrito.",
        "Cada entrada en el Decision Log debe poder vincularse a una deuda de evidencia (Seccion 24) o a un Proposal (RP-XXX).",
    ]
    for i, rule in enumerate(dl_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 24: EVIDENCE DEBT (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("24. Evidence Debt (Deuda de Evidencia)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["debt_amber"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La deuda de evidencia es como la deuda tecnica, pero en lugar de rastrear codigo mal escrito, "
        "rastrea brechas en la evidencia. Cada vez que el protocolo dice 'esto debe validarse en Fase B' "
        "o 'esto necesita calibracion con usuarios', esa es una deuda de evidencia. Si no se rastrea "
        "explicitamente, las deudas se acumulan silenciosamente y el proyecto termina construyendo "
        "castillos sobre fundamentos que nunca fueron verificados. El Evidence Debt Registry es la "
        "herramienta que hace visible lo que todavia no se sabe. Y lo que no se sabe es, frecuentemente, "
        "mas peligroso que lo que se sabe que esta mal."
    )
    run.font.size = Pt(11)

    add_callout_box(
        doc,
        "Relacion con la Meta-Regla",
        "La Meta-Regla (Seccion 22) pregunta: 'Reduce una deuda de evidencia identificada?' "
        "El Evidence Debt Registry es la lista de deudas contra las cuales se evalua cada nuevo "
        "documento, modulo o algoritmo. Si el nuevo elemento no reduce ninguna deuda de esta lista, "
        "probablemente no pertenece a la Fase A.",
        P["debt_amber"]
    )

    h2 = doc.add_heading("24.1 Formato del Evidence Debt Registry", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["debt_amber"])

    add_styled_table(
        doc,
        ["Campo", "Descripcion", "Obligatorio?"],
        [
            ["ID", "Identificador unico (ED-001, ED-002, etc.)", "Si"],
            ["Deuda", "Descripcion de la brecha de evidencia", "Si"],
            ["Impacto", "Que conclusiones se verian afectadas si esta deuda no se paga", "Si"],
            ["Prioridad", "CRITICA / ALTA / MEDIA / BAJA", "Si"],
            ["Seccion del protocolo", "Donde se identifica esta deuda", "Si"],
            ["Estado", "ABIERTA / EN PROGRESO / PAGADA", "Si"],
            ["Fecha de creacion", "Cuando se identifico", "Si"],
            ["Fecha de pago", "Cuando se resolvio (si aplica)", "Si (si esta pagada)"],
        ],
        col_widths=[3.0, 6.0, 2.0],
    )

    h2 = doc.add_heading("24.2 Registro Inicial de Deudas de Evidencia", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["debt_amber"])

    add_styled_table(
        doc,
        ["ID", "Deuda", "Impacto", "Prioridad"],
        [
            ["ED-001", "Umbral empirico no calibrado (3% provisional)", "Comparaciones sin umbral calibrado son arbitrarias", "CRITICA"],
            ["ED-002", "RVS no calibrado con usuarios reales", "Todas las conclusiones que dependen de RVS son provisionales", "CRITICA"],
            ["ED-003", "Sin validacion externa de parsers", "Resultados no comparables con herramientas del mercado", "ALTA"],
            ["ED-004", "0 claims reproducibles (RCR = 0%)", "Ningun claim puede ser verificado por terceros", "CRITICA"],
            ["ED-005", "FQS no calibrado con usuarios", "Los umbrales de calidad funcional son arbitrarios", "ALTA"],
            ["ED-006", "Solo datasets sinteticos", "Validez externa limitada a condiciones simuladas", "ALTA"],
            ["ED-007", "Solo NTFS evaluado", "No se puede generalizar a otros filesystems", "MEDIA"],
            ["ED-008", "Variabilidad natural desconocida", "No se puede distinguir senal de ruido sin 30 baseline", "CRITICA"],
        ],
        col_widths=[1.5, 4.5, 4.5, 2.0],
    )

    h2 = doc.add_heading("24.3 Reglas del Evidence Debt", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["debt_amber"])

    ed_rules = [
        "Cada brecha de evidencia identificada genera una entrada en el Evidence Debt Registry. No se omiten, incluso si son inconvenientes.",
        "Las deudas CRITICAS deben resolverse antes de la graduacion de Fase A (ver Seccion 25). Las deudas ALTA deben tener un plan de resolucion.",
        "Cada nuevo documento, modulo o algoritmo debe responder: Reduce una deuda de evidencia identificada? Si la respuesta es no, probablemente no pertenece a la Fase A (Meta-Regla, Seccion 22).",
        "Cuando una deuda se paga, no se elimina: se marca como PAGADA con la fecha y la evidencia que la resolvio. El historial es valioso.",
        "El Evidence Debt Registry se almacena en /data/evidence_debt.csv y es append-only.",
        "Las deudas de evidencia se crean durante la revision del protocolo, durante los experimentos, o cuando un auditor las identifica. No se crean de forma retroactiva para justificar trabajo ya realizado.",
    ]
    for i, rule in enumerate(ed_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 25: PHASE A GRADUATION CRITERIA (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("25. Phase A Graduation Criteria (Criterios de Graduacion)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["green"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La Fase A no termina cuando se siente que se hizo suficiente. La Fase A termina cuando "
        "se cumplen criterios objetivos y verificables. Los criterios de graduacion definen "
        "exactamente que debe lograrse antes de que el proyecto pueda declarar que la Fase A "
        "esta completa y que la evidencia acumulada es lo suficientemente solida para avanzar "
        "a la Fase B. Sin criterios de graduacion, la Fase A puede durar para siempre (si se "
        "es demasiado perfeccionista) o terminar prematuramente (si se es demasiado optimista). "
        "Los criterios de graduacion son la garante de que la Fase A produjo lo que prometio."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("25.1 Criterios Obligatorios", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["green"])

    add_styled_table(
        doc,
        ["Criterio", "Condicion", "Medicion", "Actual"],
        [
            ["RCR >= 80%", "Al menos 80% de los claims son reproducibles por terceros", "RCR = Claims Reproducibles / Claims Totales", "0%"],
            ["5+ claims con nivel ★★★", "Al menos 5 claims alcanzan el nivel de evidencia 3+ estrellas", "Confidence Registry", "2/15"],
            ["Judge API estable", "No se requiere reejecucion de experimentos por cambios en el Judge", "Historial de versiones del Judge", "v1.0 (frozen)"],
            ["1+ validacion externa por familia", "Al menos 1 herramienta de cada familia de estrategias validada", "Resultados de comparacion con herramientas externas", "0/4 familias"],
            ["30 corridas baseline completadas", "El umbral empirico esta calibrado con datos reales", "Evidence Ledger: 30+ entradas baseline", "0/30"],
            ["RVS calibrado con usuarios", "Los pesos de RVS derivan de datos reales, no de intuicion", "Resultados del experimento EXP-RVS-CAL", "No ejecutado"],
            ["run_all.py reproducible", "Cualquier persona puede clonar y ejecutar sin intervencion", "Verificacion en maquina limpia", "No implementado"],
        ],
        col_widths=[3.0, 4.0, 3.5, 2.5],
    )

    h2 = doc.add_heading("25.2 Criterios Deseables (No Bloqueantes)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["green"])

    add_styled_table(
        doc,
        ["Criterio", "Condicion", "Prioridad"],
        [
            ["8+ amenazas mitigadas", "Al menos 8 de las 19 amenazas a la validez estan mitigadas", "ALTA"],
            ["3+ formatos con referencia dorada", "Los parsers JPEG/PNG/PDF son referencia absoluta", "ALTA"],
            ["Evidence Debt < 4 CRITICAS", "Menos de 4 deudas de evidencia con prioridad CRITICA abiertas", "MEDIA"],
            ["Decision Log completo", "Todas las decisiones significativas registradas", "MEDIA"],
        ],
        col_widths=[4.0, 5.5, 2.5],
    )

    h2 = doc.add_heading("25.3 Procedimiento de Graduacion", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["green"])

    grad_steps = [
        "Verificar que todos los criterios obligatorios estan cumplidos. Si alguno no lo esta, la Fase A no puede graduarse.",
        "Revisar el Evidence Debt Registry: todas las deudas CRITICAS deben estar PAGADAS o tener un plan de pago explicito.",
        "Revisar el Decision Log: todas las decisiones significativas deben estar registradas.",
        "Ejecutar run_all.py en una maquina limpia y verificar que los resultados son identicos a los de referencia.",
        "Documentar la graduacion en el Decision Log con fecha, criterios cumplidos y firma del responsable.",
        "Solo despues de la graduacion, se puede proceder a la Fase B. En Fase B, el protocolo puede evolucionar (con Proposal RP-XXX).",
    ]
    for i, step in enumerate(grad_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {step}")
        run.font.size = Pt(11)

    add_callout_box(
        doc,
        "Condicion unica del auditor",
        "Si tuviera que poner una condicion para empezar la Fase A seria solamente una: "
        "No modificar mas el protocolo salvo que aparezca una evidencia que demuestre que "
        "el propio protocolo es insuficiente. No se cambia porque aparecio una idea mejor. "
        "Se cambia porque la evidencia mostro que el protocolo falla. Ese criterio evita "
        "la deriva metodologica y mantiene el proyecto centrado en acumular evidencia "
        "reproducible en lugar de redisenar continuamente su marco de trabajo.",
        P["green"]
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 26: EXPERIMENT VERSIONING (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("26. Experiment Versioning (Versionado de Experimentos)", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "No solo el protocolo se versiona. Cada experimento debe registrar la version de todos "
        "los componentes que participaron en su ejecucion. Dentro de seis meses probablemente "
        "exista Judge v1.1, Dataset Builder v1.4, Corruptor v1.3, y Motor C v0.9. Cuando eso "
        "ocurra, se necesitara saber exactamente con que version de cada componente se obtuvo "
        "cada resultado. Sin versionado de experimentos, es imposible determinar si una diferencia "
        "entre dos resultados se debe a un cambio en el motor, en el Judge, en el dataset, o en "
        "el protocolo. El versionado de experimentos es la unica forma de mantener la comparabilidad "
        "cuando los componentes evolucionan."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("26.1 Formato de Versionado", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "EXP-001\n"
        "  Protocol: v1.5\n"
        "  Judge: v1.0\n"
        "  Dataset Builder: v1.3\n"
        "  Corruptor: v1.2\n"
        "  Motor C: v0.8\n"
        "  Carving: v1.0\n"
        "  MFT-First: v1.0\n"
        "  RVS: v1.0\n"
        "  FQS: v1.0"
    )
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    run.font.color.rgb = c(P["body"])

    h2 = doc.add_heading("26.2 Versiones Actuales de Componentes", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Componente", "Version", "Estado", "Ubicacion"],
        [
            ["Research Protocol", "v1.5", "FROZEN (Phase A)", "research_protocol_v1.5.docx"],
            ["Judge API (RVS+FQS+Overall)", "v1.0", "FROZEN (Phase A)", "recovery_judge/"],
            ["Dataset Builder", "v1.3", "Estable", "dataset_builder/"],
            ["Corruptor", "v1.2", "Estable", "corruptor/"],
            ["Motor C", "v0.8", "En desarrollo", "motor_c/"],
            ["Carving Parser", "v1.0", "Congelado (19/19 tests)", "motor_carving.py"],
            ["MFT-First", "v1.0", "Estable", "motor_mft.py"],
            ["Evidence Gate", "v1.0", "FROZEN (Phase A)", "research_protocol_v1.5.docx"],
            ["Confidence Registry", "v1.0", "FROZEN (Phase A)", "confidence_registry.py"],
        ],
        col_widths=[4.0, 1.5, 3.5, 4.0],
    )

    h2 = doc.add_heading("26.3 Reglas de Versionado", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    ver_rules = [
        "Cada experimento debe registrar la version de TODOS los componentes involucrados en el Evidence Ledger (Seccion 17).",
        "Nunca se comparan resultados de versiones distintas de un componente. Si el Judge cambia de v1.0 a v1.1, los resultados de v1.0 no se mezclan con los de v1.1.",
        "Cuando un componente cambia de version, se documenta en el Decision Log (Seccion 23) con fecha, razon y experimentos afectados.",
        "Los componentes marcados como FROZEN no cambian durante la Fase A. Si necesitan cambiar, se sigue el procedimiento de Proposal (RP-XXX, Seccion 0.1).",
        "El versionado se aplica a nivel de componente, no a nivel de archivo individual. Cada componente tiene una unica version activa.",
    ]
    for i, rule in enumerate(ver_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 27: COMPLEJIDAD COMO COSTO CIENTIFICO (NUEVO)
    # ═══════════════════════════════════════════════════════════════════════
    h = doc.add_heading("27. Complejidad como Costo Cientifico", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["orange"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto no morira por errores. Morira por complejidad. Cada nueva seccion del "
        "protocolo es un costo. Cada nuevo modulo es un costo. Cada nueva metrica es un costo. "
        "El costo no es solo de ingenieria: es cientifico. Un protocolo con 27 secciones es "
        "mas dificil de seguir que uno con 10. Un laboratorio con 12 modulos es mas dificil "
        "de auditar que uno con 5. La complejidad es un costo cientifico porque reduce la "
        "probabilidad de que el equipo siga obedeciendo su propio protocolo. Y si el equipo "
        "no sigue el protocolo, el protocolo es inutil."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        "El auditor externo cambio su rol durante las revisiones. Al principio cuestionaba si las "
        "conclusiones eran validas. Ahora cuestionaria si el equipo sigue obedeciendo su propio "
        "protocolo. Hay una diferencia importante. Antes la pregunta era: 'Son correctas las "
        "conclusiones?' Ahora la pregunta es: 'El equipo esta haciendo lo que su protocolo dice "
        "que debe hacer?' La segunda pregunta es mas peligrosa porque si la respuesta es no, "
        "todas las conclusiones son sospechosas, incluso las correctas. La complejidad es el "
        "enemigo principal de la obediencia al protocolo."
    )
    run.font.size = Pt(11)

    h2 = doc.add_heading("27.1 Principios de Control de Complejidad", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["orange"])

    complexity_rules = [
        "Cada nueva seccion del protocolo debe justificarse frente a una deuda de evidencia (Seccion 24). Si no reduce una deuda, no se agrega.",
        "Los modulos congelados (Judge API v1.0, Hypothesis Set v1.0, Protocol v1.5) no se reabren para agregar funcionalidades.",
        "La Fase A tiene exactamente 4 actividades (Seccion 11). No se agregan actividades nuevas.",
        "El Evidence Debt Registry (Seccion 24) es la unica lista de trabajo pendiente. No hay backlog paralelo.",
        "Si el protocolo tiene mas de 30 secciones, se debe considerar una reestructuracion antes de agregar mas.",
    ]
    for i, rule in enumerate(complexity_rules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"{i}. {rule}")
        run.font.size = Pt(11)

    # ─── Save ──────────────────────────────────────────────────────────────
    output_path = "/home/z/my-project/download/Research_Protocol_v1.5_Frozen_Phase_A.docx"
    doc.save(output_path)
    print(f"Research Protocol v1.5 saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build_research_protocol()
    print(f"Generated: {path}")
