#!/usr/bin/env python3
"""
Generate Research Protocol v1.0 — RecoveryLab
===============================================
This script generates the Research Protocol document as a .docx file.
The protocol audits the SCIENCE (not the code) of the RecoveryLab project.

Key questions answered:
  1. What is the independent variable?
  2. What is the dependent variable?
  3. What is the success criterion?
  4. Confidence registry (star-based)
  5. Metric decomposition: Overall Utility = RVS × FQS
"""

import sys
import os

# Add docx scripts to path
DOCX_SCRIPTS = os.path.join("/home/z/my-project/skills/docx", "scripts")
if DOCX_SCRIPTS not in sys.path:
    sys.path.insert(0, DOCX_SCRIPTS)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Color Palette — Deep Sea Academic ────────────────────────────────────────
P = {
    "primary": "#162032",
    "body": "#1C2A3D",
    "secondary": "#5B6B7D",
    "accent": "#8B7E5A",
    "surface": "#F5F7FA",
    "white": "#FFFFFF",
    "star_gold": "#C9A84C",
    "star_dim": "#B0B0B0",
    "red": "#C0392B",
    "green": "#27AE60",
    "orange": "#E67E22",
}

def c(hex_color):
    """Convert hex color to RGBColor."""
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.replace("#", "")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = c(P["white"])
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, P["primary"])

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = c(P["body"])
            if row_idx % 2 == 1:
                set_cell_shading(cell, P["surface"])

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    return table


def build_research_protocol():
    """Build the complete Research Protocol v1.0 document."""
    doc = Document()

    # ─── Page Setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # ─── Default Font ─────────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = c(P["body"])

    # ─── COVER PAGE ────────────────────────────────────────────────────────
    # Title block
    for _ in range(4):
        doc.add_paragraph()

    # Lab name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RECOVERYLAB")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"
    run.bold = True

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Research Protocol v1.0")
    run.font.size = Pt(28)
    run.font.color.rgb = c(P["primary"])
    run.font.name = "Calibri"
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Congelar la arquitectura, auditar la ciencia")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.italic = True

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = c(P["accent"])
    run.font.size = Pt(10)

    # Metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento vivo — Version 1.0")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(11)
    run.font.color.rgb = c(P["secondary"])
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clasificacion: INTERNO — Solo para uso del laboratorio")
    run.font.size = Pt(10)
    run.font.color.rgb = c(P["red"])
    run.font.name = "Calibri"

    # Page break
    doc.add_page_break()

    # ─── SECTION 1: PREGUNTA CENTRAL ──────────────────────────────────────
    h = doc.add_heading("1. Pregunta Central del Proyecto", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El proyecto ya no gira alrededor del Motor C. Gira alrededor de una pregunta:"
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # The central question — highlighted
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "\u00bfComo medir objetivamente la utilidad de una estrategia de recuperacion?"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta pregunta es mucho mas amplia que cualquier motor individual. Es mucho mas dificil de copiar. "
        "Y es la base sobre la cual se construye el RecoveryLab Benchmark Suite: una plataforma objetiva "
        "para evaluar estrategias de recuperacion de datos. Si el laboratorio demuestra que sus resultados "
        "son reproducibles, comparables y resisten auditorias externas, habremos construido algo que vale "
        "mucho mas que un motor de recuperacion: una infraestructura de evaluacion objetiva que es "
        "extraordinariamente dificil de desarrollar y de replicar."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ─── SECTION 2: VARIABLES ─────────────────────────────────────────────
    h = doc.add_heading("2. Variables Experimentales", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    # 2.1 Independent Variable
    h2 = doc.add_heading("2.1 Variable Independiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable independiente es aquella que el experimentador controla. Solo una debe cambiar "
        "por experimento. Si cambian dos a la vez, las conclusiones se vuelven ambiguas. Este es "
        "el principio fundamental del control experimental: cada experimento debe aislar exactamente "
        "un factor para que la relacion causal sea interpretable."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Independent variables table
    add_styled_table(
        doc,
        ["Variable Independiente", "Valores", "Experimento Asociado"],
        [
            ["Tipo de dano", "MFT parcial, head crash, sectores intermitentes, ruido aleatorio, etc.", "H4 Matrix"],
            ["Nivel de dano", "0%, 10%, 20%, ..., 100%", "Crossover Curve"],
            ["Formato de archivo", "JPEG, PNG, PDF, DOCX, MP4, SQLite, TXT", "H5 Per-Format"],
            ["Estrategia", "Carving, MFT-First, Hybrid, Motor C", "H1.1 / H2"],
            ["Presupuesto de lectura", "0, 100, 500, 1000, 5000 sectores", "H1.7 Budget"],
        ],
        col_widths=[4.0, 6.0, 5.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Regla critica: ")
    run.bold = True
    run.font.color.rgb = c(P["red"])
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Solo UNA variable independiente por experimento. Si se cambia el formato Y el nivel de dano "
        "simultaneamente, no se puede atribuir la diferencia a ninguno de los dos."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # 2.2 Dependent Variable
    h2 = doc.add_heading("2.2 Variable Dependiente", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La variable dependiente es lo que se mide. Debe elegirse ANTES de ejecutar el experimento, "
        "no despues de ver los resultados. Cambiar la definicion de exito a posteriori invalida "
        "cualquier conclusion. El RecoveryLab tiene multiples metricas disponibles, pero cada "
        "experimento debe declarar cual es la principal."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Metrica", "Tipo", "Definicion", "Rango"],
        [
            ["Recovery Rate", "Conteo", "Fraccion de archivos recuperados (SHA-256 match)", "0.0-1.0"],
            ["RVS", "Valor", "Valor recuperado = suma(archivo_i.valor × peso_i) / total_valor", "0.0-1.0"],
            ["FQS", "Calidad", "Calidad funcional = suma(archivo_i.score × tamaño_i) / total_tamaño", "0.0-1.0"],
            ["Overall Utility", "Compuesto", "RVS × FQS — utilidad total de la recuperacion", "0.0-1.0"],
            ["Read Efficiency", "Eficiencia", "Lecturas utiles / Lecturas totales", "0.0-1.0"],
            ["Time to First File", "Tiempo", "Lecturas antes del primer archivo recuperado", "0-N"],
        ],
        col_widths=[3.0, 2.0, 6.5, 1.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Metrica principal declarada: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Overall Utility (RVS × FQS) es la metrica principal del laboratorio. "
        "Recovery Rate se mantiene como metrica secundaria para comparabilidad con herramientas externas."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # ─── SECTION 3: CRITERIO DE EXITO ─────────────────────────────────────
    h = doc.add_heading("3. Criterio de Exito", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El criterio de exito define cuando una estrategia se considera superior a otra. "
        "Este criterio debe declararse ANTES de ejecutar el experimento. Si se define a posteriori, "
        "siempre es posible encontrar un criterio que favorezca cualquier resultado. La definicion "
        "previa es la unica defensa contra el sesgo de confirmacion."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Success criterion box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Criterio de exito v1.0:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = c(P["accent"])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(
        "Una estrategia A se considera superior a una estrategia B unicamente si:"
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    criteria = [
        "Overall Utility (RVS × FQS) mejora al menos un 5% sobre la estrategia B",
        "La diferencia es consistente en al menos 10 repeticiones con el mismo dataset",
        "La diferencia se mantiene en al menos 3 datasets distintos",
        "La mejora no se debe exclusivamente a una dimension (RVS o FQS) sino a ambas",
    ]
    for i, criterion in enumerate(criteria, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f"{i}. {criterion}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Si solo se cumple el criterio 1 pero no los criterios 2-4, el resultado se clasifica como "
        "observacion aislada (1 estrella) y no como evidencia solida. La consistencia es tan importante "
        "como la magnitud de la mejora. Un resultado que no se replica es un resultado que no existe "
        "desde el punto de vista cientifico."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ─── SECTION 4: REGISTRO DE CONFIANZA ──────────────────────────────────
    h = doc.add_heading("4. Registro de Confianza", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El registro de confianza es la forma mas transparente de comunicar la solidez cientifica "
        "de cada resultado. En lugar de porcentajes arbitrarios de confianza, utiliza un sistema "
        "de estrellas que mapea directamente al tipo de evidencia acumulada. Este sistema es "
        "mas honesto que cualquier numero: un lector puede ver inmediatamente que evidencia "
        "respalda un resultado y que tan lejos esta de ser validado externamente."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Star scale table
    add_styled_table(
        doc,
        ["Estrellas", "Evidencia Requerida", "Significado"],
        [
            ["*", "Observacion aislada (1 run, 1 dataset)", "Preliminar — no citar como conclusion"],
            ["**", "Repetido 10+ veces (determinista, mismo dataset)", "Estable — pero no generalizado"],
            ["***", "Repetido con datasets distintos", "Generalizable — pero no validado externamente"],
            ["****", "Validado con herramientas externas (PhotoRec, TestDisk)", "Robusto — comparable con el estado del arte"],
            ["*****", "Validado con hardware real", "Definitivo — predictivo del mundo real"],
        ],
        col_widths=[2.0, 6.0, 5.0],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Reglas del registro: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "Las estrellas solo pueden subir (acumulacion de evidencia), nunca bajar. "
        "Si aparece evidencia contradictoria, el resultado se marca como CONTESTADO pero "
        "no pierde estrellas. El lector debe decidir como interpretar la contradiccion."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # Current state table
    h2 = doc.add_heading("4.1 Estado Actual del Registro", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    add_styled_table(
        doc,
        ["Resultado", "Estrellas", "Resumen"],
        [
            ["H1.1", "*** (3/5)", "MFT-First supera a Carving consistentemente. 100/100 escenarios."],
            ["H1.2", "* (1/5)", "Umbral de cambio de estrategia no determinado."],
            ["H2", "** (2/5)", "Crossover existe pero es artefacto del carving limitado."],
            ["H4", "* (1/5)", "Matriz preliminar con 3 celdas de muchas."],
            ["H5", "* (1/5)", "Sin experimento sistematico por formato."],
            ["H6", "** (2/5)", "FunctionalValidator implementado, 19/19 tests."],
            ["H7", "* (1/5)", "RVS implementado, sin validacion de usuarios."],
            ["H8", "* (1/5)", "Crossover al 95% es artefacto, no descubrimiento."],
            ["RVS", "** (2/5)", "4 dimensiones implementadas, sin validacion externa."],
            ["FQS", "** (2/5)", "5 niveles funcionales, 19/19 tests."],
            ["WFS", "* (1/5)", "Concepto definido (RVS × FQS), aun no integrado en Judge."],
            ["BLOCKER-001", "*** (3/5)", "Resuelto: comparaciones previas invalidas."],
            ["H1.5", "* (1/5)", "Gaps: sin fragmentacion, sin INDX, sin jerarquia."],
            ["H1.6", "** (2/5)", "100 ejecuciones, resultados deterministas."],
            ["H1.7", "* (1/5)", "Motor C apenas supera MFT-First (5% support)."],
        ],
        col_widths=[2.5, 2.0, 10.5],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Resumen: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(
        "8 resultados a 1 estrella, 5 a 2 estrellas, 2 a 3 estrellas, 0 a 4 estrellas, "
        "0 a 5 estrellas. Esto es honesto: la mayoria de los resultados son preliminares. "
        "No hay resultados validados externamente ni con hardware real. El laboratorio "
        "esta en una fase temprana de acumulacion de evidencia."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # ─── SECTION 5: DECOMPOSICION DE METRICAS ──────────────────────────────
    h = doc.add_heading("5. Decomposicion de Metricas", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La metrica compuesta anterior (WFS) mezclaba dos dimensiones distintas: "
        "que archivos se recuperaron y que tan bien se recuperaron. La separacion "
        "es esencial porque un motor puede ganar por dos razones muy diferentes, "
        "y la interpretacion cambia completamente segun cual sea."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Decomposition formula
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Overall Utility = RVS (Valor) × FQS (Calidad)")
    run.font.size = Pt(14)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.bold = True

    # RVS definition
    h2 = doc.add_heading("5.1 RVS — Recovery Value Score (Que se recupero)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El RVS mide el valor de lo recuperado, ponderado por la importancia de cada archivo. "
        "No es lo mismo recuperar la tesis que 200 thumbnails. El RVS incorpora cuatro dimensiones: "
        "valor intrinseco del archivo, probabilidad de reemplazo, tiempo de recreacion, e impacto "
        "emocional de la perdida. Un motor que recupera la tesis pero pierde 200 thumbnails tiene "
        "mayor RVS que uno que hace lo contrario, porque la tesis tiene mas valor, es mas dificil "
        "de reemplazar, requiere mas tiempo para recrear, y su perdida tiene mayor impacto emocional."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Tipo de Archivo", "Valor", "Reemplazo", "Recreacion", "Emocional", "RVS Peso"],
        [
            ["tesis.docx", "100", "0.05", "500h", "100", "100"],
            ["sqlite.db", "95", "0.02", "200h", "50", "95"],
            ["foto familiar", "90", "0.00", "0h", "100", "90"],
            ["video vacaciones", "70", "0.00", "0h", "80", "70"],
            ["PSD proyecto", "65", "0.10", "40h", "30", "65"],
            ["RAW foto", "60", "0.00", "0h", "60", "60"],
            ["MP4 descargado", "20", "0.90", "1h", "5", "20"],
            ["Linux ISO", "2", "1.00", "0.5h", "0", "2"],
            ["thumbnail", "1", "1.00", "0h", "0", "1"],
        ],
        col_widths=[3.0, 1.5, 1.5, 2.0, 2.0, 1.5],
    )

    # FQS definition
    h2 = doc.add_heading("5.2 FQS — Functional Quality Score (Que tan bien se recupero)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El FQS mide la calidad funcional de la recuperacion. No es lo mismo recuperar un "
        "JPEG bit-perfect que uno con 2 pixeles corruptos. El FQS reemplaza el binario "
        "SHA-256 coincide/no coincide con un espectro de calidad funcional que refleja "
        "mejor la experiencia real del usuario. Un archivo que abre y funciona tiene valor "
        "incluso si su checksum no coincide exactamente."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Nivel", "Score", "Definicion", "Ejemplo"],
        [
            ["FULL", "1.0", "SHA-256 coincide exactamente", "Copia bit-perfect del original"],
            ["FUNCTIONAL", "0.8", "Archivo funciona con dano menor", "JPEG con 2 pixeles corruptos"],
            ["PARTIAL", "0.5", "Funciona parcialmente, perdida de datos", "MP4 que se reproduce hasta la mitad"],
            ["DEGRADED", "0.2", "Contenido accesible pero daniado", "DOCX que abre pero perdio imagenes"],
            ["FAILED", "0.0", "Completamente inutilizable", "Datos aleatorios sin estructura"],
        ],
        col_widths=[2.5, 1.5, 5.0, 5.0],
    )

    # Diagnostic interpretation
    h2 = doc.add_heading("5.3 Diagnostico: Por que gano un motor?", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "La separacion RVS × FQS permite diagnosticar por que un motor gano. "
        "Un motor puede ganar por dos razones fundamentalmente distintas, y la "
        "interpretacion cambia completamente segun cual sea la dominante."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["RVS", "FQS", "Diagnostico", "Interpretacion"],
        [
            ["Alto", "Alto", "STRONG", "Recupero archivos importantes con buena calidad"],
            ["Alto", "Bajo", "VALUE-DRIVEN", "Recupero archivos importantes pero con mala calidad"],
            ["Bajo", "Alto", "QUALITY-DRIVEN", "Recupero archivos bien pero no eran importantes"],
            ["Bajo", "Bajo", "WEAK", "Fallo en ambas dimensiones"],
        ],
        col_widths=[2.0, 2.0, 3.0, 7.0],
    )

    # ─── SECTION 6: AUDITORIA DE HIPOTESIS ─────────────────────────────────
    h = doc.add_heading("6. Auditoria de Hipotesis", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cada hipotesis debe declarar explicitamente su variable independiente, variable dependiente, "
        "y criterio de exito antes de ejecutar el experimento. Si estos no estan definidos, "
        "la hipotesis no es testeable y debe ser reformulada."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Hipotesis", "Variable Independiente", "Variable Dependiente", "Criterio de Exito", "Confiabilidad"],
        [
            ["H1.1", "Estrategia (Carving vs MFT-First)", "Overall Utility", "MFT-First > Carving por 5%+ en 3+ datasets", "***"],
            ["H1.2", "Confianza en metadatos (0-100%)", "Estrategia optima", "Cambio de estrategia en umbral definido", "*"],
            ["H2", "Nivel de dano MFT (0-100%)", "Recovery Rate", "Crossover observable con carving completo", "**"],
            ["H4", "Tipo de dano", "Mejor estrategia", "Matriz completa con >80% celdas llenas", "*"],
            ["H5", "Formato de archivo", "Recovery Rate por formato", "Diferencia >10% entre formatos", "*"],
            ["H6", "Nivel de dano en archivo", "FQS", "Espectro no-binario con >2 niveles", "**"],
            ["H7", "Tipo de archivo recuperado", "RVS", "Tesis > 200 thumbnails (RVS)", "*"],
            ["H8", "Motor de carving (limitado vs completo)", "Crossover point", "Cambio significativo en crossover", "*"],
        ],
        col_widths=[1.5, 3.5, 2.5, 4.0, 1.5],
    )

    # ─── SECTION 7: FASES ─────────────────────────────────────────────────
    h = doc.add_heading("7. Hoja de Ruta por Fases", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    # Phase A
    h2 = doc.add_heading("7.1 Fase A — Ahora Mismo (Congelar)", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    phase_a_items = [
        "Congelar nuevas funcionalidades. No agregar nuevos formatos (MP4, DOCX, SQLite) al motor de carving.",
        "Consolidar el protocolo experimental. Este documento es el paso 1.",
        "Ejecutar experimentos por formato con JPEG, PNG y PDF unicamente.",
        "Repetir los experimentos suficientes veces para verificar estabilidad (minimo 10 repeticiones, 3 datasets).",
        "Completar la integracion de Overall Utility = RVS × FQS en el Judge.",
        "Exigir ratio 1:1 de codigo de recuperacion vs codigo de validacion.",
    ]
    for item in phase_a_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Phase B
    h2 = doc.add_heading("7.2 Fase B — Validacion Externa", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Validar el laboratorio contra herramientas reales. Comparar resultados con PhotoRec o TestDisk "
        "usando exactamente los mismos datasets. Si las conclusiones del laboratorio reflejan el "
        "comportamiento de herramientas existentes, los resultados ganan la cuarta estrella. "
        "Si no coinciden, el laboratorio tiene un problema de validez que debe resolverse antes "
        "de continuar. Esta fase es critica porque es la primera verificacion externa del "
        "laboratorio como plataforma de evaluacion."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    phase_b_items = [
        "Ejecutar PhotoRec sobre los mismos datasets de JPEG/PNG/PDF.",
        "Comparar Recovery Rate, RVS y FQS de PhotoRec vs los motores del laboratorio.",
        "Si los resultados no son consistentes, investigar y documentar las discrepancias.",
        "Publicar los resultados de la comparacion como parte del Benchmark Suite.",
    ]
    for item in phase_b_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Phase C
    h2 = doc.add_heading("7.3 Fase C — Expansion Controlada", level=2)
    for run in h2.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Solo entonces incorporar un formato nuevo (por ejemplo MP4) y exigirle el mismo "
        "estandar de calidad y pruebas que ya se alcanzo con JPEG, PNG y PDF. Cada nuevo "
        "formato debe pasar por el mismo proceso: parser impecable, 19+ tests de validacion, "
        "experimentos por formato, y estabilidad verificada. La regla es simple: un parser "
        "excelente de JPEG tiene muchisimo mas valor cientifico que diez parsers aceptables."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ─── SECTION 8: FORMATOS CONGELADOS ────────────────────────────────────
    h = doc.add_heading("8. Formatos Congelados — Referencia Dorada", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "Los tres parsers actuales (JPEG, PNG, PDF) tienen 19/19 tests pasados. "
        "Esto vale oro. Antes de agregar un nuevo formato, estos tres deben convertirse "
        "en una referencia absoluta. La calidad de los parsers existentes es la base "
        "sobre la cual se construye la credibilidad del laboratorio."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    add_styled_table(
        doc,
        ["Formato", "Parser", "Tests", "Estado", "Accion"],
        [
            ["JPEG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PNG", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["PDF", "motor_carving.py", "19/19", "IMPECABLE", "Convertir en referencia dorada"],
            ["MP4", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["DOCX", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["SQLite", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["CR2/NEF", "No implementado", "0/0", "CONGELADO", "No implementar hasta Fase C"],
            ["TXT", "Imposible", "0/0", "IMPOSIBLE", "Sin firma ni footer — no es carvable"],
        ],
        col_widths=[2.0, 3.0, 1.5, 2.5, 5.0],
    )

    # ─── SECTION 9: REGLA DE ORO ───────────────────────────────────────────
    h = doc.add_heading("9. Regla de Oro", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    # Highlighted box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Por cada 500 lineas nuevas de codigo de recuperacion, "
        "agregar al menos 500 lineas de validacion, pruebas y metricas."
    )
    run.font.size = Pt(13)
    run.font.color.rgb = c(P["accent"])
    run.font.name = "Calibri"
    run.bold = True
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta regla no es una sugerencia. Es la unica defensa contra la tentacion de agregar "
        "funcionalidades sin validarlas. La calidad del laboratorio es su activo mas valioso. "
        "Un motor con 3 parsers impecables y 19/19 tests es mas cientificamente valioso que "
        "un motor con 17 parsers mediocres y tests insuficientes. La presion para agregar "
        "formatos es real, pero ceder a ella destruye la base sobre la cual se construye "
        "la credibilidad del proyecto."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ─── SECTION 10: PRODUCTO ──────────────────────────────────────────────
    h = doc.add_heading("10. Producto del Proyecto", level=1)
    for run in h.runs:
        run.font.color.rgb = c(P["primary"])

    p = doc.add_paragraph()
    run = p.add_run(
        "El producto ya no es un motor de recuperacion que compite con R-Studio. "
        "El producto es el RecoveryLab Benchmark Suite: una plataforma objetiva para "
        "evaluar estrategias de recuperacion de datos. Si el laboratorio demuestra que "
        "sus resultados son reproducibles, comparables y resisten auditorias externas, "
        "habremos construido algo que vale mucho mas que un motor individual: una "
        "infraestructura de evaluacion objetiva que es extraordinariamente dificil de "
        "desarrollar y de replicar. Este cambio de producto es fundamental y debe "
        "guiar todas las decisiones futuras."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # ─── Save ──────────────────────────────────────────────────────────────
    output_path = "/home/z/my-project/download/Research_Protocol_v1.0.docx"
    doc.save(output_path)
    print(f"Research Protocol saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build_research_protocol()
    print(f"Generated: {path}")
